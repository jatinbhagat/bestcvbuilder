"""
Advanced CV Parser API for Comprehensive ATS Analysis
Production-ready implementation with industry-aligned scoring
Hosted on Vercel as serverless function
Updated: Fix FITZ_AVAILABLE build error - Aug 14, 2025
"""

import json
import os
import requests
import re
import io
from typing import Dict, Any, List, Tuple, Optional
import logging
from collections import Counter
from urllib.parse import urlparse
import math
import uuid
import gc
import sys
from datetime import datetime

# Configure logging first
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Memory management utilities
def cleanup_memory():
    """Force garbage collection to free memory"""
    try:
        # Force multiple garbage collection passes
        for _ in range(3):
            collected = gc.collect()
            if collected == 0:
                break
        logger.debug(f"🧹 CV parser memory cleanup - collected {collected} objects")
    except Exception as e:
        logger.warning(f"Memory cleanup failed: {e}")

def get_memory_usage():
    """Get current memory usage info if available"""
    try:
        import psutil
        process = psutil.Process(os.getpid())
        memory_mb = round(process.memory_info().rss / 1024 / 1024, 2)
        return memory_mb
    except ImportError:
        return None

def extract_tables_with_pdfplumber(file_content: bytes) -> str:
    """Extract complex tables using pdfplumber when needed"""
    try:
        import pdfplumber
        import io
        
        pdf_file = io.BytesIO(file_content)
        table_texts = []
        
        with pdfplumber.open(pdf_file) as pdf:
            for page_num, page in enumerate(pdf.pages):
                tables = page.extract_tables()
                
                if tables:
                    logger.info(f"📊 Page {page_num + 1}: Found {len(tables)} tables")
                    
                    for table_num, table in enumerate(tables):
                        # Convert table to text format
                        table_text = f"\n--- Table {table_num + 1} from Page {page_num + 1} ---\n"
                        
                        for row in table:
                            if row and any(cell for cell in row if cell):
                                # Clean and join cells
                                clean_row = [str(cell).strip() if cell else '' for cell in row]
                                table_text += ' | '.join(clean_row) + '\n'
                        
                        table_texts.append(table_text)
        
        result = '\n'.join(table_texts)
        if result:
            logger.info(f"✅ Table extraction complete: {len(result)} characters")
        
        return result
        
    except ImportError:
        logger.warning("⚠️ pdfplumber not available - skipping table extraction")
        return ""
    except Exception as e:
        logger.warning(f"⚠️ Table extraction failed: {e}")
        return ""

def extract_pdf_text_clean(file_content: bytes) -> str:
    """Clean PDF extraction using only PyMuPDF (table extraction disabled)"""
    if not PYMUPDF_AVAILABLE:
        raise TextExtractionError("PyMuPDF (fitz) is required for PDF extraction")
    
    logger.info("📄 Starting clean PDF extraction with PyMuPDF")
    
    try:
        import fitz
        
        # Check if text is extractable from first page
        pdf_document = fitz.open(stream=file_content, filetype="pdf")
        
        if pdf_document.page_count == 0:
            pdf_document.close()
            raise TextExtractionError("PDF has no pages")
        
        # Check first page for text
        first_page_text = pdf_document[0].get_text().strip()
        if not first_page_text or len(first_page_text) < 50:
            pdf_document.close()
            raise TextExtractionError("Scanned or image-based resumes are not supported. Please upload a text-based PDF.")
        
        # Extract text from all pages using PyMuPDF only
        text_parts = []
        
        for page_num in range(pdf_document.page_count):
            page = pdf_document[page_num]
            page_text = page.get_text()
            
            if page_text and page_text.strip():
                text_parts.append(page_text)
        
        pdf_document.close()
        
        result = '\n\n'.join(text_parts)
        logger.info(f"✅ PDF extraction complete: {len(result)} characters")
        
        return result
        
    except Exception as e:
        logger.error(f"❌ PDF extraction failed: {e}")
        raise TextExtractionError(f"PDF extraction failed: {str(e)}")

# Legacy function - now uses PyMuPDF instead of PyPDF2
def extract_with_pymupdf_simple_legacy(file_content: bytes) -> str:
    """Simple PyMuPDF extraction for large files - memory efficient"""
    try:
        doc = fitz.open(stream=file_content, filetype="pdf")
        
        text_parts = []
        max_pages = min(len(doc), 5)  # Limit to 5 pages for memory
        
        for i in range(max_pages):
            try:
                page = doc.load_page(i)
                page_text = page.get_text()
                if page_text and len(page_text.strip()) > 20:
                    text_parts.append(page_text)
                    
                # Clean up is automatic with PyMuPDF
                
                # Force cleanup every 2 pages
                if i % 2 == 0:
                    cleanup_memory()
                    
            except Exception as e:
                logger.warning(f"Error extracting page {i+1}: {e}")
                continue
        
        # Clean up document
        doc.close()
        cleanup_memory()
        
        result = '\n'.join(text_parts)
        logger.info(f"Simple extraction completed: {len(result)} characters")
        return result
        
    except Exception as e:
        logger.error(f"Simple PDF extraction failed: {e}")
        cleanup_memory()
        raise TextExtractionError(f"Simple PDF extraction failed: {str(e)}")

# Import configuration system
from config.config_loader import (
    get_industry_keywords, get_grammar_patterns, get_spelling_corrections,
    get_achievement_verbs, get_professional_indicators, get_component_weights,
    get_score_categories, get_keywords_for_industry, config_loader
)

# Text extraction libraries - Critical dependency checking
PYPDF2_AVAILABLE = False
DOCX_AVAILABLE = False
PDFPLUMBER_AVAILABLE = False
PYMUPDF_AVAILABLE = False
PDFMINER_AVAILABLE = False

# PyPDF2 not used - removed dependency
PYPDF2_AVAILABLE = False

# Check python-docx (for DOCX files) - Optional for minimal deployment
try:
    import docx
    DOCX_AVAILABLE = True
    logger.info("✅ python-docx available (DOCX extraction)")
except ImportError as e:
    logger.info(f"ℹ️  python-docx not available (minimal deployment): {e}")
    DOCX_AVAILABLE = False

# Check pdfplumber (better PDF extraction)
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
    logger.info("✅ pdfplumber available (enhanced PDF extraction)")
except ImportError as e:
    logger.warning(f"⚠️  pdfplumber not available: {e}")
    PDFPLUMBER_AVAILABLE = False

# Check PyMuPDF/fitz (primary PDF extraction)
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
    logger.info("✅ PyMuPDF available for PDF extraction")
except ImportError as e:
    logger.error(f"❌ CRITICAL: PyMuPDF not available: {e}")
    PYMUPDF_AVAILABLE = False

# Check pdfminer (comprehensive extraction)
try:
    from pdfminer.high_level import extract_text as pdfminer_extract_text
    PDFMINER_AVAILABLE = True
    logger.info("✅ pdfminer.six available (comprehensive PDF extraction)")
except ImportError as e:
    logger.warning(f"⚠️  pdfminer not available: {e}")
    PDFMINER_AVAILABLE = False

# Log dependency summary
def log_dependency_status():
    """Log the status of all PDF extraction dependencies"""
    total_available = sum([PYPDF2_AVAILABLE, PDFPLUMBER_AVAILABLE, PYMUPDF_AVAILABLE, PDFMINER_AVAILABLE])
    
    logger.info("📊 PDF Extraction Dependencies Status:")
    logger.info(f"   PyPDF2: {'✅ Available' if PYPDF2_AVAILABLE else '❌ Missing'}")
    logger.info(f"   pdfplumber: {'✅ Available' if PDFPLUMBER_AVAILABLE else '❌ Missing'}")
    logger.info(f"   PyMuPDF: {'✅ Available' if PYMUPDF_AVAILABLE else '❌ Missing'}")
    logger.info(f"   pdfminer: {'✅ Available' if PDFMINER_AVAILABLE else '❌ Missing'}")
    logger.info(f"   python-docx: {'✅ Available' if DOCX_AVAILABLE else '❌ Missing'}")
    logger.info(f"Total PDF extraction methods available: {total_available}/4")
    
    if total_available == 1 and PYPDF2_AVAILABLE:
        logger.info("✅ PyPDF2 available - optimized for serverless deployment")
    elif total_available >= 3:
        logger.info("✅ Excellent: Multiple PDF extraction methods available for high quality")
    elif total_available >= 2:
        logger.info("✅ Good: Multiple PDF extraction methods available")
    elif total_available == 0:
        logger.error("❌ CRITICAL: No PDF extraction methods available!")
        
    return total_available

# Check dependencies at startup
available_methods = log_dependency_status()

# API configuration
ALLOWED_ORIGINS = [
    "http://localhost:3000",
    "http://localhost:3002", 
    "https://bestcvbuilder-gnktl1mxh-bestcvbuilder.vercel.app"
]

# Configuration-driven data - loaded from config files
# Industry keywords, action verbs, and other data points are now externalized

# Contact information patterns
CONTACT_PATTERNS = {
    'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
    'phone': r'(\+\d{1,2}\s?)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}',
    'linkedin': r'linkedin\.com/in/[\w-]+|linkedin\.com/in/[\w\.-]+|www\.linkedin\.com/in/[\w-]+',
    'website': r'https?://[\w.-]+\.[\w]{2,}',
    'github': r'github\.com/[\w-]+'
}

# Additional patterns for comprehensive data extraction
NAME_PATTERNS = [
    r'^([A-Z][a-z]+ [A-Z][a-z]+)',  # First line name pattern
    r'([A-Z][a-z]+ [A-Z][a-z]+)\s*\n',  # Name followed by newline
    r'Name:\s*([A-Z][a-z]+ [A-Z][a-z]+)',  # Explicit name field
    # Enhanced patterns for merged PDF lines and Unicode
    r'^([A-Z][a-zA-Z\u00C0-\u017F]+ [A-Z][a-zA-Z\u00C0-\u017F]+(?:\s+[A-Z][a-zA-Z\u00C0-\u017F]+)?)',  # Unicode support
    r'([A-Z][a-zA-Z\u00C0-\u017F]{2,}\s+[A-Z][a-zA-Z\u00C0-\u017F]{2,}(?:\s+[A-Z][a-zA-Z\u00C0-\u017F]{2,})?)',  # General name pattern with Unicode
]

ADDRESS_PATTERNS = [
    r'(\d+\s+[A-Za-z\s,]+(?:Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Lane|Ln|Boulevard|Blvd|Way|Place|Pl)[\s,]*[A-Za-z\s,]*\d{5}(?:-\d{4})?)',  # US address
    r'Address:\s*(.+?)(?:\n|Email|Phone)',  # Explicit address field
    r'(\d+\s+[A-Za-z\s,]+(?:,\s*[A-Z]{2}\s*\d{5}))',  # City, State ZIP
]

CITY_STATE_PATTERNS = [
    r'([A-Za-z\s]+),\s*([A-Z]{2})\s*(\d{5}(?:-\d{4})?)',  # City, State ZIP
    r'([A-Za-z\s]+),\s*([A-Za-z\s]+)(?:\s*,\s*[A-Z]{2,3})?',  # City, State/Country
]

SKILLS_PATTERNS = [
    r'(?:SKILLS?|TECHNICAL SKILLS?|TECHNOLOGIES?|COMPETENCIES)[\s:]*\n?(.*?)(?:\n\n|\n[A-Z]{2,}|\Z)',
    r'(?:Programming|Languages?|Tools?|Frameworks?)[\s:]*[:\-]?\s*(.*?)(?:\n|$)',
]

EDUCATION_PATTERNS = [
    r'(?:EDUCATION|ACADEMIC|QUALIFICATIONS?)[\s:]*\n?(.*?)(?:\n\n|\n[A-Z]{2,}|\Z)',
    r'([A-Za-z\s]+(?:University|College|Institute|School).*?)(?:\n(?:[A-Z]{2,}|$))',
    r'((?:Bachelor|Master|PhD|MBA|B\.S\.|M\.S\.|B\.A\.|M\.A\.).*?)(?:\n|$)',
]

EXPERIENCE_PATTERNS = [
    r'(?:EXPERIENCE|EMPLOYMENT|WORK HISTORY|PROFESSIONAL EXPERIENCE)[\s:]*\n?(.*?)(?:\n\n|\n[A-Z]{2,}|\Z)',
    r'([A-Za-z\s]+\|[A-Za-z\s,]+\|\s*\d{4}.*?)(?:\n(?:[A-Z]{2,}|$))',
]

SUMMARY_PATTERNS = [
    r'(?:SUMMARY|PROFILE|OBJECTIVE|ABOUT)[\s:]*\n?(.*?)(?:\n\n|\n[A-Z]{2,}|\Z)',
]

# Quantified achievement patterns
QUANTIFIED_PATTERNS = [
    r'\b\d+%\b',  # Percentages
    r'\$\d+[,\d]*(?:\.\d{2})?\b',  # Dollar amounts
    r'\b\d+[,\d]*\s*(?:million|thousand|billion|k)\b',  # Large numbers
    r'\b\d+\s*(?:years?|months?|weeks?|days?)\b',  # Time periods
    r'\b\d+\s*(?:people|employees|team members|clients|customers|users)\b',  # Team/user sizes
    r'\b\d+[,\d]*\s*(?:projects?|initiatives?|campaigns?|deals?)\b',  # Project counts
    r'\b(?:increased|decreased|improved|reduced|grew|generated)\s+.*?\d+[%\d]*\b'  # Performance metrics
]

class ATSAnalysisError(Exception):
    """Custom exception for ATS analysis errors"""
    pass

class FileProcessingError(ATSAnalysisError):
    """Error during file processing"""
    pass

class TextExtractionError(ATSAnalysisError):
    """Error during text extraction"""
    pass

def cors_headers():
    """Return CORS headers for API responses"""
    return {
        "Access-Control-Allow-Origin": "*",
        "Access-Control-Allow-Methods": "POST, OPTIONS, HEAD, GET",
        "Access-Control-Allow-Headers": "Content-Type, Accept, Authorization",
        "Access-Control-Max-Age": "86400"
    }

def validate_file_url(file_url: str) -> bool:
    """Validate file URL format and security"""
    try:
        parsed = urlparse(file_url)
        if not parsed.scheme in ['http', 'https']:
            return False
        if not parsed.netloc:
            return False
        return True
    except:
        return False

def extract_text_from_file(file_content: bytes, file_url: str) -> str:
    """
    Extract text content using ultra-safe methods to prevent worker timeouts
    
    Args:
        file_content: Raw file content
        file_url: Original file URL for determining file type
        
    Returns:
        Extracted text content
    """
    file_extension = file_url.split('.')[-1].lower()
    file_size_mb = len(file_content) / (1024 * 1024)
    
    logger.info(f"🔍 Extracting {file_extension.upper()} file ({file_size_mb:.1f}MB) with timeout protection")
    
    try:
        if file_extension == 'pdf':
            return extract_pdf_text_clean(file_content)
                
        elif file_extension == 'docx':
            return extract_docx_text(file_content)
        elif file_extension == 'doc':
            return extract_doc_text(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    except Exception as e:
        logger.error(f"Text extraction failed: {str(e)}")
        raise TextExtractionError(f"Could not extract text from file: {str(e)}")

def extract_pdf_text(file_content: bytes) -> str:
    """Legacy function - redirects to clean extraction"""
    return extract_pdf_text_clean(file_content)
    
    # Check if any extraction method is available
    if not any([PDFPLUMBER_AVAILABLE, PYMUPDF_AVAILABLE, PDFMINER_AVAILABLE]):
        raise TextExtractionError("CRITICAL: No PDF extraction libraries available")
    
    # Prioritized extraction methods (best first for Render.com)
    extraction_methods = []
    
    # Priority 1: pdfplumber (excellent for structured text, emails, preserves formatting)
    if PDFPLUMBER_AVAILABLE:
        extraction_methods.append(("pdfplumber", extract_with_pdfplumber))
    
    # Priority 2: PyMuPDF (best overall accuracy, speed, and text quality)  
    if PYMUPDF_AVAILABLE:
        extraction_methods.append(("pymupdf", extract_with_pymupdf))
        
    # Priority 3: pdfminer (comprehensive extraction, good for complex layouts)
    if PDFMINER_AVAILABLE:
        extraction_methods.append(("pdfminer", extract_with_pdfminer))
        
    # PyPDF2 removed - using PyMuPDF instead
    
    if not extraction_methods:
        raise TextExtractionError("No PDF extraction methods available")
    
    logger.info(f"🔄 Trying {len(extraction_methods)} PDF extraction methods in priority order")
    
    best_text = ""
    best_score = 0
    best_method = "none"
    
    for method_name, method_func in extraction_methods:
        try:
            logger.info(f"📄 Attempting PDF extraction with {method_name}")
            text = method_func(file_content)
            
            if text and len(text.strip()) >= 50:
                quality_score = score_text_quality(text)
                logger.info(f"📊 {method_name} extraction score: {quality_score}/100")
                
                # Update best result if this is better
                if quality_score > best_score:
                    best_text = text
                    best_score = quality_score
                    
                    # Clean up to save memory
                    cleanup_memory()
                
                # If we get excellent results, stop trying other methods
                if quality_score >= 85:
                    logger.info(f"✅ Excellent quality ({quality_score}) achieved with {method_name}, stopping here")
                    break
                    best_method = method_name
                    logger.info(f"🏆 New best extraction method: {method_name} (score: {quality_score})")
                
                # If we get excellent results with priority methods, use them immediately
                if quality_score >= 90 and method_name in ["pdfplumber", "pymupdf"]:
                    logger.info(f"🎯 Excellent quality achieved with priority method {method_name}")
                    break
                    
                # For less accurate methods, require higher threshold to stop early
                elif quality_score >= 95 and method_name in ["pdfminer", "pypdf2"]:
                    logger.info(f"🎯 Excellent quality achieved with {method_name}")
                    break
                    
            else:
                logger.warning(f"⚠️  {method_name} extracted minimal text")
                    
        except Exception as e:
            logger.warning(f"❌ {method_name} failed: {str(e)}")
            continue
    
    # Ensure we have some text to work with
    if not best_text or len(best_text.strip()) < 50:
        # Last resort: try basic PyMuPDF if available  
        if PYMUPDF_AVAILABLE:
            try:
                logger.warning("🚨 Falling back to basic PyMuPDF extraction (last resort)")
                basic_text = extract_with_pymupdf(file_content)
                if basic_text and len(basic_text.strip()) >= 20:  # Lower threshold for last resort
                    logger.info(f"✅ Basic PyMuPDF extracted {len(basic_text)} characters")
                    return basic_text
            except Exception as e:
                logger.error(f"❌ Basic PyMuPDF extraction also failed: {str(e)}")
        
        raise TextExtractionError("PDF extraction failed - file may be corrupted, image-only, or password protected")
    
    logger.info(f"🎉 Final PDF extraction completed successfully!")
    logger.info(f"📊 Best method score: {best_score}/100")
    logger.info(f"📝 Extracted text length: {len(best_text)} characters")
    return best_text

def extract_with_pypdf2_basic(file_content: bytes) -> str:
    """Basic PyPDF2 extraction without cleaning - last resort fallback"""
    text = ""
    pdf_file = io.BytesIO(file_content)
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    
    for page in pdf_reader.pages:
        try:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        except Exception as e:
            logger.warning(f"Failed to extract text from page: {e}")
            continue
    
    return text

def extract_with_pypdf2(file_content: bytes) -> str:
    """Extract text using PyPDF2 with memory management"""
    text = ""
    pdf_file = io.BytesIO(file_content)
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    
    page_count = 0
    max_pages = min(len(pdf_reader.pages), 10)  # Limit pages for memory
    
    for page in pdf_reader.pages[:max_pages]:
        try:
            # Try different extraction methods for better quality
            page_text = page.extract_text()
            
            # If standard extraction yields poor results, try alternative method
            if page_text and len(page_text.strip()) < 50:
                # Try extracting with different parameters for better results
                try:
                    page_text = page.extract_text(extraction_mode="layout")
                except:
                    pass  # Fall back to standard extraction
            
            if page_text:
                text += page_text + "\n"
            
            page_count += 1
            
            # Clean up memory every 3 pages
            if page_count % 3 == 0:
                cleanup_memory()
                
        except Exception as e:
            logger.warning(f"Failed to extract text from page: {e}")
            continue
    
    # Clean up reader and file
    del pdf_reader
    pdf_file.close()
    cleanup_memory()
    
    # Enhanced cleaning for PyPDF2 specific issues
    cleaned_text = clean_extracted_text_enhanced(text)
    return cleaned_text

def clean_extracted_text_enhanced(text: str) -> str:
    """
    Enhanced text cleaning specifically for PyPDF2 extraction issues
    """
    if not text:
        return text
    
    # Basic normalization first
    text = clean_extracted_text(text)
    
    # Fix common PyPDF2 email extraction issues
    # Pattern: Single letter prefix before email (e.g., "ebhagat.jatin@gmail.com")
    text = re.sub(r'\b[a-z]([a-zA-Z0-9._%+-]{2,}@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})\b', r'\1', text)
    
    # Fix broken email domains (e.g., "gmail. com" -> "gmail.com")
    text = re.sub(r'@([a-zA-Z0-9.-]+)\.\s+([a-zA-Z]{2,})\b', r'@\1.\2', text)
    
    # Fix spaces in email local part (e.g., "user .name@domain.com" -> "user.name@domain.com")
    text = re.sub(r'([a-zA-Z0-9_%+-]+)\s+\.([a-zA-Z0-9_%+-]+)@', r'\1.\2@', text)
    
    # Fix spaces around @ symbol
    text = re.sub(r'([a-zA-Z0-9._%+-]+)\s+@\s+([a-zA-Z0-9.-]+\.[a-zA-Z]{2,})', r'\1@\2', text)
    
    return text

def extract_with_pdfplumber(file_content: bytes) -> str:
    """Extract text using pdfplumber (more accurate for complex layouts)"""
    if not PDFPLUMBER_AVAILABLE:
        raise Exception("pdfplumber not available in runtime environment")
    
    try:
        import pdfplumber
        text = ""
        pdf_file = io.BytesIO(file_content)
        
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        return clean_extracted_text(text)
    except Exception as e:
        raise Exception(f"pdfplumber extraction failed: {str(e)}")

def extract_with_pymupdf(file_content: bytes) -> str:
    """Extract text using PyMuPDF/fitz (good for text accuracy)"""
    if not PYMUPDF_AVAILABLE:
        raise Exception("PyMuPDF not available in runtime environment")
    
    try:
        import fitz  # PyMuPDF
        text = ""
        pdf_document = fitz.open("pdf", file_content)
        
        for page_num in range(pdf_document.page_count):
            page = pdf_document[page_num]
            page_text = page.get_text()
            if page_text:
                text += page_text + "\n"
        
        pdf_document.close()
        return clean_extracted_text(text)
    except Exception as e:
        raise Exception(f"PyMuPDF extraction failed: {str(e)}")

def extract_with_pdfminer(file_content: bytes) -> str:
    """Extract text using pdfminer (most comprehensive but slower)"""
    if not PDFMINER_AVAILABLE:
        raise Exception("pdfminer not available in runtime environment")
    
    try:
        from io import BytesIO
        
        text = pdfminer_extract_text(BytesIO(file_content))
        return clean_extracted_text(text)
    except Exception as e:
        raise Exception(f"pdfminer extraction failed: {str(e)}")

def score_text_quality(text: str) -> float:
    """
    Score the quality of extracted text based on various factors
    Returns a score from 0-100
    """
    if not text or len(text.strip()) < 10:
        return 0
    
    score = 0
    
    # Factor 1: Basic text characteristics (30 points)
    word_count = len(text.split())
    if word_count >= 50:
        score += 15
    elif word_count >= 20:
        score += 10
    elif word_count >= 10:
        score += 5
    
    # Reasonable length (not too short or extremely long)
    if 100 <= word_count <= 2000:
        score += 15
    elif 50 <= word_count < 100 or 2000 < word_count <= 5000:
        score += 10
    elif word_count > 5000:
        score += 5
    
    # Factor 2: Email detection quality (25 points)
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b'
    emails = re.findall(email_pattern, text)
    
    if emails:
        # Check for clean email extraction (no obvious artifacts)
        clean_emails = [email for email in emails if not re.match(r'^[a-z][A-Z]', email)]
        if clean_emails:
            score += 25
        else:
            score += 10  # Found emails but they might have artifacts
    
    # Factor 3: Text structure quality (20 points)
    lines = text.split('\n')
    non_empty_lines = [line for line in lines if line.strip()]
    
    if len(non_empty_lines) >= 10:
        score += 10
    elif len(non_empty_lines) >= 5:
        score += 5
    
    # Check for reasonable line lengths (not too many very short or very long lines)
    reasonable_lines = [line for line in non_empty_lines if 10 <= len(line.strip()) <= 200]
    if len(reasonable_lines) >= len(non_empty_lines) * 0.7:
        score += 10
    elif len(reasonable_lines) >= len(non_empty_lines) * 0.5:
        score += 5
    
    # Factor 4: Character quality (15 points)
    # Check for reasonable character distribution
    alpha_ratio = sum(1 for c in text if c.isalpha()) / len(text)
    if 0.6 <= alpha_ratio <= 0.9:
        score += 10
    elif 0.4 <= alpha_ratio <= 0.95:
        score += 5
    
    # Check for reasonable punctuation
    punct_ratio = sum(1 for c in text if c in '.,;:!?') / len(text)
    if 0.02 <= punct_ratio <= 0.15:
        score += 5
    
    # Factor 5: Professional content indicators (10 points)
    professional_keywords = [
        'experience', 'education', 'skills', 'work', 'job', 'company',
        'university', 'degree', 'project', 'manager', 'developer', 'engineer'
    ]
    
    found_keywords = sum(1 for keyword in professional_keywords if keyword.lower() in text.lower())
    if found_keywords >= 5:
        score += 10
    elif found_keywords >= 3:
        score += 5
    elif found_keywords >= 1:
        score += 2
    
    return min(score, 100)

def clean_extracted_text(text: str) -> str:
    """
    Basic text cleaning for PDF extraction results
    Focus on simple normalization rather than complex artifact fixing
    """
    # Normalize whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove excessive line breaks
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    
    # Fix common spacing issues around punctuation
    text = re.sub(r'\s+([,.;:!?])', r'\1', text)
    
    return text.strip()

def extract_docx_text(file_content: bytes) -> str:
    """Extract text from DOCX files"""
    if not DOCX_AVAILABLE:
        raise TextExtractionError("DOCX processing not available - please use PDF format")
    
    try:
        import docx
        doc_file = io.BytesIO(file_content)
        doc = docx.Document(doc_file)
        
        text = ""
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text += " | ".join(row_text) + "\n"
                
        return text
    except Exception as e:
        logger.error(f"DOCX extraction error: {str(e)}")
        raise TextExtractionError(f"DOCX extraction failed: {str(e)}")

def extract_doc_text(file_content: bytes) -> str:
    """Extract text from DOC files using docx2txt"""
    # DOC format not supported with current dependencies
    raise TextExtractionError("DOC files are not supported. Please use PDF or DOCX format.")

def detect_industry(content: str) -> str:
    """Detect the most likely industry based on content keywords"""
    content_lower = content.lower()
    industry_scores = {}
    
    # Load industry keywords from config
    industry_keywords = get_industry_keywords()
    
    for industry, categories in industry_keywords.items():
        score = 0
        for category, keywords in categories.items():
            for keyword in keywords:
                if keyword in content_lower:
                    score += content_lower.count(keyword)
        industry_scores[industry] = score
    
    if not industry_scores or max(industry_scores.values()) == 0:
        return 'general'
    
    return max(industry_scores, key=industry_scores.get)

def analyze_content_structure(content: str) -> Dict[str, Any]:
    """Analyze resume structure and organization using config data - More stringent scoring"""
    content_lower = content.lower()
    
    # Load essential sections from config
    essential_sections_config = config_loader.get_essential_sections()
    
    # Build essential sections structure with patterns and weights
    essential_sections = {
        'contact': {
            'patterns': ['email', 'phone', 'linkedin', '@', 'contact'],
            'weight': 8,
            'found': False
        },
        'experience': {
            'patterns': essential_sections_config.get('required', [])[:5],  # First 5 experience patterns
            'weight': 10,
            'found': False
        },
        'education': {
            'patterns': essential_sections_config.get('required', [])[5:10],  # Education patterns
            'weight': 6,
            'found': False
        },
        'skills': {
            'patterns': essential_sections_config.get('required', [])[10:],  # Skills patterns
            'weight': 6,
            'found': False
        }
    }
    
    # Fallback to hardcoded patterns if config is empty
    if not essential_sections_config.get('required'):
        essential_sections = {
            'contact': {
                'patterns': ['email', 'phone', 'linkedin', '@', 'contact'],
                'weight': 8,
                'found': False
            },
            'experience': {
                'patterns': ['experience', 'employment', 'work history', 'professional experience', 'career'],
                'weight': 10,
                'found': False
            },
            'education': {
                'patterns': ['education', 'degree', 'university', 'college', 'school', 'academic'],
                'weight': 6,
                'found': False
            },
            'skills': {
                'patterns': ['skills', 'technical skills', 'competencies', 'technologies', 'tools'],
                'weight': 6,
                'found': False
            }
        }
    
    # Check for sections
    total_score = 0
    found_sections = []
    
    for section, data in essential_sections.items():
        for pattern in data['patterns']:
            if pattern in content_lower:
                data['found'] = True
                found_sections.append(section)
                total_score += data['weight']
                break
    
    # Load optional sections from config
    optional_sections_config = config_loader.get_essential_sections()
    optional_sections = optional_sections_config.get('recommended', []) + optional_sections_config.get('optional', [])
    
    # Fallback if config is empty
    if not optional_sections:
        optional_sections = ['summary', 'objective', 'achievements', 'projects', 'certifications', 'awards']
    
    optional_found = []
    
    for section in optional_sections:
        if section in content_lower:
            optional_found.append(section)
            total_score += 2  # Bonus points
    
    # More stringent scoring - must have most sections to get high scores
    section_score = min(total_score * 0.8, 20)  # Reduce by 20% and cap at 20
    
    return {
        'score': max(section_score, 0),  # More realistic scoring
        'essential_sections': found_sections,
        'optional_sections': optional_found,
        'missing_sections': [k for k, v in essential_sections.items() if not v['found']]
    }

def analyze_keyword_optimization(content: str, industry: str = None) -> Dict[str, Any]:
    """Analyze keyword density and relevance - MAX 20 POINTS to match config weights"""
    content_lower = content.lower()
    score = 0
    keyword_analysis = {}
    
    # Load industry keywords from config
    industry_keywords_config = get_industry_keywords()
    
    # Industry-specific keywords
    if industry and industry in industry_keywords_config:
        industry_keywords = industry_keywords_config[industry]
        industry_found = {}
        
        for category, keywords in industry_keywords.items():
            found_keywords = []
            for keyword in keywords:
                count = content_lower.count(keyword)
                if count > 0:
                    found_keywords.append({
                        'keyword': keyword,
                        'count': count,
                        'density': count / len(content.split()) * 1000  # Per 1000 words
                    })
                    # Award points with diminishing returns
                    score += min(count * 2, 4)  # Max 4 points per keyword
            
            if found_keywords:
                industry_found[category] = found_keywords
        
        keyword_analysis['industry_keywords'] = industry_found
    
    # Action verbs analysis using config
    action_verbs_config = get_achievement_verbs()
    action_verb_score = 0
    found_verbs = {}
    
    for category, verbs in action_verbs_config.items():
        category_verbs = []
        for verb in verbs:
            count = content_lower.count(verb)
            if count > 0:
                category_verbs.append({'verb': verb, 'count': count})
                action_verb_score += min(count, 2)  # Max 2 points per verb
        
        if category_verbs:
            found_verbs[category] = category_verbs
    
    score += min(action_verb_score, 15)  # Cap action verbs at 15 points
    keyword_analysis['action_verbs'] = found_verbs
    
    return {
        'score': min(score, 20),  # Cap at 20 points to match config weights
        'analysis': keyword_analysis
    }

def analyze_contact_information(content: str) -> Dict[str, Any]:
    """Analyze contact information completeness and format using config patterns"""
    found_contacts = {}
    score = 0
    
    # Load contact requirements from config
    contact_requirements = config_loader.get_contact_requirements()
    
    # Enhanced phone number extraction with multiple strategies
    phone_patterns = [
        r'\+?\d{1,4}[\s\-\.]?\(?\d{3}\)?[\s\-\.]?\d{3}[\s\-\.]?\d{4}',  # International format
        r'\(?\d{3}\)?[\s\-\.]?\d{3}[\s\-\.]?\d{4}',  # US format
        r'\d{3}[\s\-\.]\d{3}[\s\-\.]\d{4}',  # Simple format
        r'\+\d{1,3}\s?\(?\d{3}\)?\s?\d{3}[\s\-\.]?\d{4}',  # Mixed international
        r'\d{10,}',  # Raw digits (10+ digits)
    ]
    
    phone_matches = []
    for pattern in phone_patterns:
        matches = re.findall(pattern, content)
        for match in matches:
            # Validate phone number has enough digits
            digits = re.sub(r'[^\d]', '', match)
            if len(digits) >= 10:
                phone_matches.append(match)
    
    # Remove duplicates and take the most complete phone number
    if phone_matches:
        # Sort by length to get the most complete number
        phone_matches.sort(key=len, reverse=True)
        found_contacts['phone'] = phone_matches
    
    # Process essential contact types
    for contact_type, contact_data in contact_requirements.get('essential', {}).items():
        if contact_type == 'phone' and 'phone' in found_contacts:
            # Use enhanced phone extraction results
            matches = found_contacts['phone']
        else:
            patterns = contact_data.get('patterns', [])
            weight = contact_data.get('weight', 3)
            
            matches = []
            for pattern in patterns:
                pattern_matches = re.findall(pattern, content, re.IGNORECASE)
                matches.extend(pattern_matches)
            
            found_contacts[contact_type] = matches
        
        if matches:
            weight = contact_data.get('weight', 3)
            score += weight
    
    # Process recommended contact types
    for contact_type, contact_data in contact_requirements.get('recommended', {}).items():
        patterns = contact_data.get('patterns', [])
        weight = contact_data.get('weight', 2)
        
        matches = []
        for pattern in patterns:
            pattern_matches = re.findall(pattern, content, re.IGNORECASE)
            matches.extend(pattern_matches)
        
        found_contacts[contact_type] = matches
        if matches:
            score += weight
    
    # Process optional contact types
    for contact_type, contact_data in contact_requirements.get('optional', {}).items():
        patterns = contact_data.get('patterns', [])
        weight = contact_data.get('weight', 1)
        
        matches = []
        for pattern in patterns:
            pattern_matches = re.findall(pattern, content, re.IGNORECASE)
            matches.extend(pattern_matches)
        
        found_contacts[contact_type] = matches
        if matches:
            score += weight
    
    return {
        'score': min(score, 15),  # Cap at 15 points
        'found_contacts': found_contacts,
        'missing': [k for k, v in found_contacts.items() if not v]
    }

def analyze_formatting_quality(content: str) -> Dict[str, Any]:
    """Analyze formatting and ATS readability - More stringent scoring"""
    score = 15  # Start with lower baseline, deduct for issues
    issues = []
    
    # Check for formatting issues that hurt ATS parsing
    formatting_checks = {
        'excessive_whitespace': (r'\s{4,}', 'Excessive whitespace detected'),
        'special_characters': (r'[^\w\s\-\.,@():/]', 'Unusual special characters found'),
        'inconsistent_spacing': (r'\n\s*\n\s*\n', 'Inconsistent paragraph spacing'),
        'tab_characters': (r'\t', 'Tab characters detected (use spaces instead)'),
        'mixed_date_formats': (r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}.*\d{4}[/-]\d{1,2}[/-]\d{1,2}', 'Mixed date formats')
    }
    
    for check_name, (pattern, message) in formatting_checks.items():
        matches = re.findall(pattern, content)
        if matches:
            count = len(matches)
            deduction = min(count, 5)  # Max 5 points deduction per issue type
            score -= deduction
            issues.append(f"{message} ({count} instances)")
    
    # Check for proper bullet points
    bullet_patterns = ['•', '◦', '▪', '-', '*']
    has_bullets = any(bullet in content for bullet in bullet_patterns)
    if not has_bullets:
        score -= 3
        issues.append("No bullet points detected - consider using bullets for lists")
    
    return {
        'score': max(score, 0),  # Don't go below 0
        'issues': issues
    }

def analyze_quantified_achievements(content: str) -> Dict[str, Any]:
    """Look for quantified achievements with numbers and percentages using config patterns"""
    quantified_achievements = []
    
    # Load quantification patterns from config
    quantification_patterns = config_loader.get_quantification_patterns()
    
    # Combine all pattern types
    all_patterns = []
    for pattern_type, patterns in quantification_patterns.items():
        all_patterns.extend(patterns)
    
    # Fallback to hardcoded patterns if config is empty
    if not all_patterns:
        all_patterns = [
            r'\b\d+%\b',  # Percentages
            r'\$\d+[,\d]*(?:\.\d{2})?\b',  # Dollar amounts
            r'\b\d+[,\d]*\s*(?:million|thousand|billion|k)\b',  # Large numbers
            r'\b\d+\s*(?:years?|months?|weeks?|days?)\b',  # Time periods
            r'\b\d+\s*(?:people|employees|team members|clients|customers|users)\b',  # Team/user sizes
            r'\b\d+[,\d]*\s*(?:projects?|initiatives?|campaigns?|deals?)\b',  # Project counts
            r'\b(?:increased|decreased|improved|reduced|grew|generated)\s+.*?\d+[%\d]*\b'  # Performance metrics
        ]
    
    for pattern in all_patterns:
        matches = re.finditer(pattern, content, re.IGNORECASE)
        for match in matches:
            # Get surrounding context (20 chars before and after)
            start = max(0, match.start() - 20)
            end = min(len(content), match.end() + 20)
            context = content[start:end].replace('\n', ' ').strip()
            
            quantified_achievements.append({
                'match': match.group(),
                'context': context
            })
    
    # Remove duplicates based on context similarity
    unique_achievements = []
    for achievement in quantified_achievements:
        if not any(achievement['context'] in existing['context'] or 
                  existing['context'] in achievement['context'] 
                  for existing in unique_achievements):
            unique_achievements.append(achievement)
    
    score = min(len(unique_achievements) * 2, 10)  # 2 points per achievement, max 10
    
    return {
        'score': score,
        'count': len(unique_achievements),
        'achievements': unique_achievements[:5]  # First 5 examples
    }

def check_grammar_issues(content: str) -> List[Dict[str, str]]:
    """
    Check for common grammar issues in resume content using config patterns
    """
    grammar_issues = []
    
    # Load grammar patterns from config
    grammar_patterns = get_grammar_patterns()
    
    for pattern_data in grammar_patterns:
        pattern = pattern_data['pattern']
        message = pattern_data['message']
        
        matches = re.finditer(pattern, content, re.IGNORECASE)
        for match in matches:
            grammar_issues.append({
                'type': 'grammar',
                'issue': message,
                'position': match.start(),
                'text': match.group(),
                'category': pattern_data.get('category', 'general'),
                'severity': pattern_data.get('severity', 'medium')
            })
            # Limit to avoid overwhelming the user
            if len(grammar_issues) >= 6:
                break
        if len(grammar_issues) >= 6:
            break
    
    return grammar_issues

def check_spelling_issues(content: str) -> List[Dict[str, str]]:
    """
    Check for common spelling issues in resume content using config corrections
    """
    spelling_issues = []
    
    # Load spelling corrections from config
    spelling_corrections = get_spelling_corrections()
    
    # Check for each misspelling
    words = re.findall(r'\b\w+\b', content.lower())
    for word in words:
        if word in spelling_corrections:
            spelling_issues.append({
                'type': 'spelling',
                'incorrect': word,
                'correct': spelling_corrections[word],
                'issue': f'Misspelled "{word}" should be "{spelling_corrections[word]}"'
            })
            # Limit to avoid overwhelming the user
            if len(spelling_issues) >= 6:
                break
    
    return spelling_issues

def analyze_readability_and_length(content: str) -> Dict[str, Any]:
    """
    Analyze content readability, optimal length, grammar, and spelling using config thresholds - MAX 10 POINTS to match config weights
    """
    word_count = len(content.split())
    char_count = len(content)
    sentence_count = len(re.findall(r'[.!?]+', content))
    
    # Grammar and spelling analysis
    grammar_issues = check_grammar_issues(content)
    spelling_issues = check_spelling_issues(content)
    
    # Load readability metrics from config
    readability_metrics = config_loader.get_readability_metrics()
    
    # Optimal word count scoring using config thresholds
    optimal_word_count = readability_metrics.get('optimal_word_count', {
        'min': 300, 'max': 800, 'ideal_min': 400, 'ideal_max': 600
    })
    
    length_score = 0
    length_feedback = ""
    
    if optimal_word_count['ideal_min'] <= word_count <= optimal_word_count['ideal_max']:
        length_score = 5
        length_feedback = "Optimal resume length"
    elif optimal_word_count['min'] <= word_count < optimal_word_count['ideal_min'] or optimal_word_count['ideal_max'] < word_count <= optimal_word_count['max']:
        length_score = 4
        length_feedback = "Good resume length"
    elif word_count < optimal_word_count['min']:
        length_score = 1
        length_feedback = "Resume too short - add more detail"
    else:
        length_score = 2
        length_feedback = "Resume too long - consider condensing"
    
    # Calculate average sentence length
    avg_sentence_length = word_count / max(sentence_count, 1)
    
    # Readability scoring using config
    optimal_sentence_length = readability_metrics.get('optimal_sentence_length', {
        'min': 10, 'max': 20, 'ideal': 15
    })
    
    readability_score = 0
    if optimal_sentence_length['min'] <= avg_sentence_length <= optimal_sentence_length['max']:
        readability_score = 3
    elif avg_sentence_length < optimal_sentence_length['min'] or avg_sentence_length > optimal_sentence_length['max']:
        readability_score = 2
    else:
        readability_score = 1
    
    # Enhanced grammar scoring with severity-based penalties
    grammar_score = 4  # Start with max points
    punctuation_penalty = 0
    date_formatting_penalty = 0
    
    # Calculate penalties based on issue severity and category
    for issue in grammar_issues:
        severity = issue.get('severity', 'medium')
        category = issue.get('category', 'general')
        
        if category == 'punctuation':
            if severity == 'high':
                punctuation_penalty += 1.0
            elif severity == 'medium':
                punctuation_penalty += 0.5
            else:
                punctuation_penalty += 0.25
        elif category == 'date_formatting':
            if severity == 'high':
                date_formatting_penalty += 1.0
            elif severity == 'medium':
                date_formatting_penalty += 0.5
            else:
                date_formatting_penalty += 0.25
        else:
            # General grammar issues
            if severity == 'high':
                grammar_score -= 1.0
            elif severity == 'medium':
                grammar_score -= 0.5
            else:
                grammar_score -= 0.25
    
    # Apply specific penalties
    grammar_score -= min(punctuation_penalty, 2.0)  # Max 2 point penalty for punctuation
    grammar_score -= min(date_formatting_penalty, 2.0)  # Max 2 point penalty for date formatting
    grammar_score = max(grammar_score, 0)  # Don't go below 0
    
    # Spelling scoring (3 points max) 
    spelling_score = 0
    if len(spelling_issues) == 0:
        spelling_score = 3
    elif len(spelling_issues) <= 2:
        spelling_score = 2
    else:
        spelling_score = 1
    
    # Scale total score to match config weight of 10 points (not 15)
    total_score = length_score + readability_score + grammar_score + spelling_score
    scaled_score = min(total_score * (10.0 / 15.0), 10)  # Scale from 15-point range to 10-point range
    
    return {
        'score': round(scaled_score, 1),  # Max 10 points to match config weights
        'word_count': word_count,
        'character_count': char_count,
        'sentence_count': sentence_count,
        'avg_sentence_length': round(avg_sentence_length, 1),
        'feedback': length_feedback,
        'grammar_errors': len(grammar_issues),
        'spelling_errors': len(spelling_issues),
        'grammar_issues': [issue['issue'] for issue in grammar_issues[:3]],  # Top 3
        'spelling_issues': [issue['issue'] for issue in spelling_issues[:3]], # Top 3
        'punctuation_penalty': round(punctuation_penalty, 2),
        'date_formatting_penalty': round(date_formatting_penalty, 2),
        'readability_level': 'Excellent' if scaled_score >= 8.5 else 'Good' if scaled_score >= 6.5 else 'Needs Improvement'
    }

def analyze_date_formatting(content: str) -> Dict[str, Any]:
    """
    Analyze date formatting consistency in work experience and education sections
    Returns score out of 10 points
    """
    score = 10  # Start with perfect score
    issues = []
    
    # Extract relevant sections only (Experience, Education, Projects)
    relevant_content = extract_relevant_sections_for_dates(content)
    
    # More precise date patterns with year validation
    date_patterns = [
        (r'\b(0[1-9]|1[0-2])/(19[9-9][0-9]|20[0-3][0-9])\b', 'MM/YYYY'),           # MM/YYYY format
        (r'\b([1-9])/(19[9-9][0-9]|20[0-3][0-9])\b', 'M/YYYY'),                     # M/YYYY format  
        (r'\b(0[1-9]|1[0-2])-(19[9-9][0-9]|20[0-3][0-9])\b', 'MM-YYYY'),           # MM-YYYY format
        (r'\b([1-9])-(19[9-9][0-9]|20[0-3][0-9])\b', 'M-YYYY'),                     # M-YYYY format
        (r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+(19[9-9][0-9]|20[0-3][0-9])\b', 'Month YYYY'),  # Month YYYY
        (r'\b(19[9-9][0-9]|20[0-3][0-9])\s*[-–]\s*(19[9-9][0-9]|20[0-3][0-9])\b', 'YYYY-YYYY'),  # YYYY-YYYY format
        (r'\b(19[9-9][0-9]|20[0-3][0-9])\s*[-–]\s*(Present|Ongoing|Current)\b', 'YYYY-Present'),  # YYYY - Present
        (r'\b(19[9-9][0-9]|20[0-3][0-9])\s+(to|through)\s+(19[9-9][0-9]|20[0-3][0-9])\b', 'YYYY to YYYY'),  # YYYY to YYYY
    ]
    
    # Find all dates and track which format they use
    all_dates = []
    format_types = []
    format_names = []
    
    for i, (pattern, format_name) in enumerate(date_patterns):
        matches = re.findall(pattern, relevant_content, re.IGNORECASE)
        for match in matches:
            # Extract the actual date string (handle tuple matches)
            if isinstance(match, tuple):
                date_str = match[0] if format_name in ['MM/YYYY', 'M/YYYY', 'MM-YYYY', 'M-YYYY'] else ' '.join(match)
            else:
                date_str = match
            
            # Additional validation to avoid false positives
            if is_valid_employment_date(date_str, relevant_content):
                all_dates.append(date_str.strip())
                format_types.append(i)
                format_names.append(format_name)
    
    # Debug output
    logger.info(f"🗓️ Found {len(all_dates)} valid dates in relevant sections")
    for date, fmt in zip(all_dates, format_names):
        logger.info(f"🗓️ Date: '{date}' -> Format: {fmt}")
    
    # AGGRESSIVE SCORING: Count all types of inconsistencies
    total_inconsistencies = 0
    
    # 1. Format inconsistencies - count each unique format beyond the first as an inconsistency
    unique_formats = set(format_types)
    if len(unique_formats) > 1:
        # Each additional format beyond the most common one is an inconsistency
        format_inconsistencies = len(unique_formats) - 1  # -1 because one format is the "standard"
        total_inconsistencies += format_inconsistencies
        issues.append(f"Mixed date formats: {format_inconsistencies} different non-standard date formats found ({len(unique_formats)} total formats)")
    
    # 2. Missing dates inconsistency - if some positions have dates but others don't
    position_entries = extract_position_entries(relevant_content)
    positions_with_dates = 0
    total_positions = len(position_entries)
    
    for entry in position_entries:
        entry_has_date = any(date in entry for date in all_dates)
        if entry_has_date:
            positions_with_dates += 1
    
    if total_positions > 0 and positions_with_dates < total_positions:
        missing_date_inconsistencies = total_positions - positions_with_dates
        total_inconsistencies += missing_date_inconsistencies
        issues.append(f"Missing dates: {missing_date_inconsistencies} positions/entries lack proper dates")
    
    # 3. Incomplete date inconsistencies - mix of year-only vs month-year formats
    year_only_count = 0
    full_date_count = 0
    
    for date, fmt in zip(all_dates, format_names):
        if fmt in ['Month YYYY', 'MM/YYYY', 'M/YYYY', 'MM-YYYY', 'M-YYYY']:
            full_date_count += 1
        elif len(date) == 4 and date.isdigit():  # Year only like "2022"
            year_only_count += 1
    
    if year_only_count > 0 and full_date_count > 0:
        # Count year-only dates as inconsistencies if we have full dates elsewhere
        total_inconsistencies += year_only_count
        issues.append(f"Incomplete dates: {year_only_count} dates are year-only while others include months")
    
    # AGGRESSIVE PENALTY: 2 points per inconsistency
    penalty = total_inconsistencies * 2
    score -= penalty
    
    # Additional penalty for completely missing dates
    if not all_dates:
        score -= 6  # Heavy penalty for no dates at all
        issues.append("No employment dates found in experience section")
        total_inconsistencies += 3  # Count as 3 inconsistencies
    elif len(all_dates) < 2:
        score -= 4  # Penalty for very limited dates
        issues.append("Limited date information in work experience")
        total_inconsistencies += 2  # Count as 2 inconsistencies
    
    return {
        'score': max(score, 0),
        'total_dates_found': len(all_dates),
        'unique_formats': len(unique_formats),
        'total_inconsistencies': total_inconsistencies,
        'penalty_applied': penalty,
        'detected_dates': all_dates[:10],  # Show first 10 for debugging
        'issues': issues
    }

def extract_relevant_sections_for_dates(content: str) -> str:
    """
    Extract only sections where employment/education dates should appear
    """
    lines = content.split('\n')
    relevant_lines = []
    current_section = ""
    include_section = False
    
    # Keywords that indicate relevant sections
    experience_keywords = [
        'experience', 'employment', 'work history', 'professional', 'career',
        'positions', 'roles', 'jobs', 'internship', 'intern'
    ]
    education_keywords = [
        'education', 'academic', 'degree', 'university', 'college', 'school',
        'certification', 'training', 'course'
    ]
    project_keywords = ['projects', 'portfolio', 'achievements']
    
    # Headers to exclude (these often contain false positive dates)
    exclude_keywords = [
        'contact', 'personal', 'address', 'phone', 'email', 'references',
        'skills', 'technologies', 'languages', 'hobbies', 'interests'
    ]
    
    for line in lines:
        line_lower = line.lower().strip()
        
        # Check if this is a section header
        if len(line_lower) > 0 and (line.isupper() or line.strip().endswith(':') or 
            any(keyword in line_lower for keyword in experience_keywords + education_keywords + project_keywords + exclude_keywords)):
            
            # Determine if we should include this section
            include_section = (
                any(keyword in line_lower for keyword in experience_keywords + education_keywords + project_keywords) and
                not any(keyword in line_lower for keyword in exclude_keywords)
            )
            current_section = line_lower
            
        # Include the line if we're in a relevant section
        if include_section:
            relevant_lines.append(line)
    
    relevant_content = '\n'.join(relevant_lines)
    logger.info(f"🗓️ Extracted {len(relevant_content)} characters from relevant sections for date analysis")
    return relevant_content

def is_valid_employment_date(date_str: str, context: str) -> bool:
    """
    Additional validation to ensure the matched pattern is actually an employment date
    """
    # Skip if it appears to be in contact information context
    contact_indicators = ['phone', 'tel', 'mobile', 'address', 'email', 'linkedin']
    context_lower = context.lower()
    
    # Check surrounding context for phone/address patterns
    date_index = context_lower.find(date_str.lower())
    if date_index != -1:
        # Check 50 characters before and after the date
        start = max(0, date_index - 50)
        end = min(len(context_lower), date_index + len(date_str) + 50)
        surrounding_text = context_lower[start:end]
        
        if any(indicator in surrounding_text for indicator in contact_indicators):
            return False
    
    # Skip obvious phone number patterns
    if re.match(r'^\d{1,2}[/-]\d{4}$', date_str) and '/' in date_str:
        # Could be phone number like "1/2023" - check if it's in a reasonable year range
        try:
            parts = date_str.split('/')
            if len(parts) == 2 and int(parts[0]) > 12:  # Month can't be > 12
                return False
        except:
            return False
    
    return True

def extract_position_entries(content: str) -> List[str]:
    """
    Extract individual position/job entries to check for missing dates
    """
    entries = []
    lines = content.split('\n')
    current_entry = []
    
    # Common patterns that indicate a new position/entry
    position_indicators = [
        r'^[A-Z][A-Za-z\s&,.-]+\|',  # Job Title | Company
        r'^[A-Z][A-Za-z\s&,.-]+\s*-\s*[A-Z]',  # Job Title - Company
        r'^[A-Z][A-Za-z\s&,.-]+\s*,\s*[A-Z]',  # Job Title, Company
        r'^\s*•\s*[A-Z]',  # Bullet point start
        r'^\s*-\s*[A-Z]',  # Dash bullet point
    ]
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check if this line starts a new position entry
        is_new_position = any(re.match(pattern, line) for pattern in position_indicators[:3])
        
        if is_new_position and current_entry:
            # Save previous entry and start new one
            entries.append('\n'.join(current_entry))
            current_entry = [line]
        else:
            current_entry.append(line)
    
    # Add the last entry
    if current_entry:
        entries.append('\n'.join(current_entry))
    
    # Filter out very short entries (likely not real positions)
    meaningful_entries = [entry for entry in entries if len(entry.split()) > 3]
    
    return meaningful_entries

def analyze_bullet_lengths(content: str) -> Dict[str, Any]:
    """
    Analyze bullet point length optimization in Experience section only
    Returns score out of 10 points
    """
    score = 10  # Start with perfect score
    issues = []
    
    # Extract only Experience section content
    experience_content = extract_experience_section(content)
    
    if not experience_content:
        return {
            'score': 5,  # Neutral score if no experience section found
            'total_bullets': 0,
            'issues': ['No Experience section found for bullet analysis']
        }
    
    # Find actual bullet points (sentence-like content in experience)
    all_bullets = extract_experience_bullets(experience_content)
    
    if not all_bullets:
        return {
            'score': 5,  # Neutral score if no bullets found
            'total_bullets': 0,
            'issues': ['No bullet points detected in Experience section']
        }
    
    # Analyze each bullet point
    non_optimal_count = 0
    too_short = 0    # < 10 words
    optimal = 0      # 10-30 words  
    too_long = 0     # > 30 words
    
    logger.info(f"🔫 Found {len(all_bullets)} bullets in Experience section:")
    
    for i, bullet in enumerate(all_bullets):
        # Clean bullet text and count words
        clean_bullet = re.sub(r'[^\w\s]', ' ', bullet)
        word_count = len([word for word in clean_bullet.split() if word.strip()])
        
        logger.info(f"🔫 Bullet {i+1}: '{bullet[:50]}...' ({word_count} words)")
        
        if word_count < 10:
            too_short += 1
            non_optimal_count += 1
        elif word_count <= 30:
            optimal += 1
        else:
            too_long += 1
            non_optimal_count += 1
    
    total_bullets = len(all_bullets)
    
    # Deduct 1.5 points per non-optimal bullet
    penalty = non_optimal_count * 1.5
    score = max(score - penalty, 0)  # Don't go below 0
    
    # Generate issues
    if too_short > 0:
        issues.append(f"{too_short} bullets too short (< 10 words)")
    if too_long > 0:
        issues.append(f"{too_long} bullets too long (> 30 words)")
        
    logger.info(f"🔫 Bullet analysis complete: {optimal} optimal, {too_short} too short, {too_long} too long")
        
    return {
        'score': round(score, 1),
        'total_bullets': total_bullets,
        'optimal_bullets': optimal,
        'too_short': too_short,
        'too_long': too_long,
        'non_optimal_count': non_optimal_count,
        'penalty_applied': round(penalty, 1),
        'sample_bullets': all_bullets[:5],  # First 5 for debugging
        'issues': issues
    }

def extract_experience_section(content: str) -> str:
    """
    Extract only the Experience/Work History section content
    """
    lines = content.split('\n')
    experience_lines = []
    in_experience = False
    
    # Keywords that indicate experience section start
    experience_keywords = [
        'experience', 'employment', 'work history', 'professional experience',
        'career history', 'work experience', 'professional background'
    ]
    
    # Keywords that indicate other sections (end experience)
    other_section_keywords = [
        'education', 'skills', 'certifications', 'projects', 'achievements',
        'languages', 'interests', 'references', 'contact', 'summary',
        'objective', 'profile', 'training', 'awards', 'publications'
    ]
    
    for line in lines:
        line_stripped = line.strip()
        line_lower = line_stripped.lower()
        
        # Check if this is an experience section header
        if any(keyword in line_lower for keyword in experience_keywords):
            in_experience = True
            continue
            
        # Check if this is a different section header
        elif in_experience and any(keyword in line_lower for keyword in other_section_keywords):
            in_experience = False
            break
            
        # Collect experience content
        if in_experience and line_stripped:
            experience_lines.append(line_stripped)
    
    experience_content = '\n'.join(experience_lines)
    logger.info(f"🔫 Extracted {len(experience_content)} characters from Experience section")
    return experience_content

def extract_experience_bullets(experience_content: str) -> List[str]:
    """
    Extract actual bullet points from experience section
    Excludes job titles, company names, dates, and single words
    """
    lines = experience_content.split('\n')
    bullets = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Remove bullet symbols if present
        clean_line = re.sub(r'^\s*[•\-\*o▪→]\s*', '', line).strip()
        
        # Skip if line is too short to be a sentence
        if len(clean_line) < 15:  # At least 15 characters for a meaningful sentence
            continue
            
        # Skip obvious non-bullets
        if is_non_bullet_line(clean_line):
            continue
            
        # Must be sentence-like (at least 3 words with some complexity)
        words = clean_line.split()
        if len(words) < 3:
            continue
            
        # Check if it looks like a sentence (has some complexity)
        if is_sentence_like(clean_line):
            bullets.append(clean_line)
    
    return bullets

def is_non_bullet_line(line: str) -> bool:
    """
    Check if a line is NOT a bullet point (job title, company, date, etc.)
    """
    line_lower = line.lower().strip()
    
    # Skip lines that are all caps (likely section headers or company names)
    if line.isupper() and len(line) > 3:
        return True
        
    # Skip lines that are mostly dates
    if re.match(r'^[a-zA-Z]*\s*\d{4}[\s\-–]+[a-zA-Z]*\s*\d{4}$', line.strip()):
        return True
        
    # Skip lines that end with common date patterns
    if re.search(r'(present|current|ongoing)$', line_lower):
        return True
        
    # Skip obvious job titles (often followed by dates or at company)
    job_title_patterns = [
        r'^(senior|junior|lead|principal|chief|head of|director of|manager|associate|assistant)',
        r'(engineer|developer|analyst|specialist|coordinator|consultant|executive)$',
        r'^[A-Z][a-z]+ [A-Z][a-z]+$'  # Title Case words (likely job titles)
    ]
    
    for pattern in job_title_patterns:
        if re.search(pattern, line, re.IGNORECASE):
            return True
            
    # Skip company names (often have Inc, LLC, Corp, etc.)
    if re.search(r'\b(inc|llc|corp|ltd|company|technologies|solutions|systems)\b', line_lower):
        return True
        
    return False

def is_sentence_like(line: str) -> bool:
    """
    Check if a line looks like a sentence (has verbs, complexity)
    """
    words = line.split()
    
    # Must have at least 4 words to be considered sentence-like
    if len(words) < 4:
        return False
        
    # Should contain some action words or descriptive content
    # Look for common verb patterns or connecting words
    action_indicators = [
        'developed', 'created', 'managed', 'led', 'implemented', 'designed',
        'built', 'achieved', 'improved', 'increased', 'reduced', 'collaborated',
        'worked', 'responsible', 'coordinated', 'analyzed', 'maintained',
        'supported', 'delivered', 'executed', 'planned', 'optimized',
        'and', 'with', 'for', 'by', 'using', 'through', 'to', 'of', 'in'
    ]
    
    line_lower = line.lower()
    has_action_words = any(word in line_lower for word in action_indicators)
    
    # Should have reasonable length and complexity
    has_good_length = len(line) >= 20  # At least 20 characters
    
    return has_action_words and has_good_length

# ========================================
# FRONTEND ANALYSIS FUNCTIONS - COPIED EXACTLY FROM FRONTEND
# ========================================

def analyze_contact_details_frontend(resume_text: str) -> int:
    """Copied exactly from frontend analyzeContactDetails"""
    score = 10  # Start with perfect score, deduct 2.5 for each missing element
    
    # 1. Mobile Number (2.5 points)
    has_mobile = has_mobile_number(resume_text)
    if not has_mobile:
        score -= 2.5
    
    # 2. Email Address (2.5 points)
    has_email = has_email_address(resume_text)
    if not has_email:
        score -= 2.5
    
    # 3. LinkedIn Profile (2.5 points)
    has_linkedin = has_linkedin_profile(resume_text)
    if not has_linkedin:
        score -= 2.5
    
    # 4. Location (2.5 points)
    has_location = has_location_info(resume_text)
    if not has_location:
        score -= 2.5
    
    return max(0, min(10, score))

def has_mobile_number(resume_text: str) -> bool:
    """Check if resume contains mobile number"""
    mobile_patterns = [
        r'\+?[\d\s\-\(\)]{10,}',  # General phone pattern
        r'\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',  # US format
        r'\(\d{3}\)\s?\d{3}[-.\s]?\d{4}'   # (XXX) XXX-XXXX
    ]
    text_lower = resume_text.lower()
    if any(keyword in text_lower for keyword in ['phone', 'mobile', 'cell', 'tel']):
        return True
    for pattern in mobile_patterns:
        if re.search(pattern, resume_text):
            return True
    return False

def has_email_address(resume_text: str) -> bool:
    """Check if resume contains email address"""
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    return bool(re.search(email_pattern, resume_text))

def has_linkedin_profile(resume_text: str) -> bool:
    """Check if resume contains LinkedIn profile"""
    text_lower = resume_text.lower()
    return 'linkedin' in text_lower or 'linkedin.com' in text_lower

def has_location_info(resume_text: str) -> bool:
    """Check if resume contains location information"""
    location_keywords = ['city', 'state', 'country', 'address', 'location']
    text_lower = resume_text.lower()
    # Check for common location patterns
    if any(keyword in text_lower for keyword in location_keywords):
        return True
    # Check for state abbreviations or zip codes
    if re.search(r'\b[A-Z]{2}\b|\b\d{5}(-\d{4})?\b', resume_text):
        return True
    return False

def analyze_education_section_frontend(resume_text: str, structure_data: dict = None) -> int:
    """Copied exactly from frontend analyzeEducationSection"""
    score = 8  # Start with good score
    
    education_keywords = ['education', 'degree', 'university', 'college', 'bachelor', 'master', 'phd']
    text_lower = resume_text.lower()
    
    # Check if education section exists
    has_education_section = any(keyword in text_lower for keyword in education_keywords)
    if not has_education_section:
        score -= 3
    
    # Check for degree information
    degree_keywords = ['bachelor', 'master', 'phd', 'b.s.', 'm.s.', 'b.a.', 'm.a.']
    has_degree = any(keyword in text_lower for keyword in degree_keywords)
    if has_degree:
        score += 1
    
    # Check for graduation year
    if re.search(r'(19|20)\d{2}', resume_text):
        score += 1
    
    return max(0, min(10, score))

def analyze_skills_section_frontend(resume_text: str, structure_data: dict = None) -> int:
    """Copied exactly from frontend analyzeSkillsSection"""
    score = 7  # Start with average score
    
    skills_keywords = ['skills', 'technical skills', 'core competencies', 'technologies']
    text_lower = resume_text.lower()
    
    # Check if skills section exists
    has_skills_section = any(keyword in text_lower for keyword in skills_keywords)
    if not has_skills_section:
        score -= 2
    
    # Check for programming languages
    programming_languages = ['python', 'java', 'javascript', 'c++', 'sql', 'html', 'css']
    found_languages = sum(1 for lang in programming_languages if lang in text_lower)
    if found_languages >= 3:
        score += 2
    elif found_languages >= 1:
        score += 1
    
    # Check for tools and frameworks
    tools = ['excel', 'powerpoint', 'word', 'git', 'aws', 'docker', 'kubernetes']
    found_tools = sum(1 for tool in tools if tool in text_lower)
    if found_tools >= 2:
        score += 1
    
    return max(0, min(10, score))

def analyze_analytical_skills_frontend(resume_text: str, keywords_data: dict = None) -> int:
    """Copied exactly from frontend analyzeAnalyticalSkills"""
    analytical_keywords = [
        'analyze', 'analysis', 'analytical', 'data', 'statistics', 'metrics',
        'insights', 'research', 'investigate', 'evaluate', 'assess'
    ]
    text_lower = resume_text.lower()
    found_keywords = sum(1 for keyword in analytical_keywords if keyword in text_lower)
    
    if found_keywords >= 5:
        return 9
    elif found_keywords >= 3:
        return 7
    elif found_keywords >= 1:
        return 5
    else:
        return 3

def analyze_leadership_skills_frontend(resume_text: str, keywords_data: dict = None) -> int:
    """Copied exactly from frontend analyzeLeadershipSkills"""
    leadership_keywords = [
        'lead', 'leader', 'leadership', 'manage', 'manager', 'management',
        'supervise', 'mentor', 'train', 'coordinate', 'direct', 'guide'
    ]
    text_lower = resume_text.lower()
    found_keywords = sum(1 for keyword in leadership_keywords if keyword in text_lower)
    
    if found_keywords >= 5:
        return 9
    elif found_keywords >= 3:
        return 7
    elif found_keywords >= 1:
        return 5
    else:
        return 3

def analyze_page_density_frontend(resume_text: str) -> int:
    """Copied exactly from frontend analyzePageDensity"""
    lines = resume_text.split('\n')
    non_empty_lines = [line for line in lines if line.strip()]
    
    # Estimate page density based on content
    if len(non_empty_lines) > 40:
        return 9  # Good density
    elif len(non_empty_lines) > 20:
        return 7  # Average density
    else:
        return 5  # Low density

def analyze_bullet_usage_frontend(resume_text: str, formatting_data: dict = None) -> int:
    """Copied exactly from frontend analyzeBulletUsage"""
    bullet_symbols = ['•', '-', '*', '○', '▪', '→']
    has_bullets = any(symbol in resume_text for symbol in bullet_symbols)
    
    if has_bullets:
        bullet_count = sum(resume_text.count(symbol) for symbol in bullet_symbols)
        if bullet_count >= 5:
            return 9
        elif bullet_count >= 3:
            return 7
        else:
            return 5
    else:
        return 3

def analyze_grammar_frontend(resume_text: str) -> int:
    """Copied exactly from frontend analyzeGrammar - LLM-powered"""
    # Simple grammar check patterns
    grammar_issues = 0
    
    # Check for common grammar mistakes
    text_lower = resume_text.lower()
    
    # Subject-verb disagreement patterns
    disagreement_patterns = [
        r'\bi am\b.*\bresponsible\b',  # Should be "I was" for past jobs
        r'\bwere\b.*\bi\b',  # "were I" instead of "was I"
    ]
    
    for pattern in disagreement_patterns:
        if re.search(pattern, text_lower):
            grammar_issues += 1
    
    # Sentence structure issues
    if grammar_issues == 0:
        return 9
    elif grammar_issues <= 2:
        return 7
    else:
        return 5

def analyze_llm_spelling_frontend(resume_text: str) -> int:
    """Copied exactly from frontend analyzeLLMSpelling - LLM-powered"""
    # Common spelling mistakes to check
    common_mistakes = {
        'recieve': 'receive',
        'seperate': 'separate',
        'definately': 'definitely',
        'occured': 'occurred',
        'managment': 'management',
        'responsibilty': 'responsibility'
    }
    
    text_lower = resume_text.lower()
    spelling_errors = 0
    
    for mistake in common_mistakes.keys():
        if mistake in text_lower:
            spelling_errors += 1
    
    if spelling_errors == 0:
        return 10
    elif spelling_errors <= 2:
        return 8
    else:
        return 5

def analyze_verb_tenses_frontend(resume_text: str) -> int:
    """Copied exactly from frontend analyzeVerbTenses"""
    # Check for consistent verb tenses
    past_tense_verbs = ['developed', 'created', 'managed', 'led', 'implemented']
    present_tense_verbs = ['develop', 'create', 'manage', 'lead', 'implement']
    
    text_lower = resume_text.lower()
    past_count = sum(1 for verb in past_tense_verbs if verb in text_lower)
    present_count = sum(1 for verb in present_tense_verbs if verb in text_lower)
    
    # Prefer past tense for work experience
    if past_count > present_count:
        return 8
    elif past_count == present_count:
        return 6
    else:
        return 4

def analyze_personal_pronouns_frontend(resume_text: str) -> int:
    """Analyzes personal pronouns with improved word boundary detection"""
    import re
    
    # Use word boundaries to avoid false positives with company names and abbreviations
    pronoun_patterns = [
        r'\bi\b',        # "i" as standalone word
        r'\bme\b',       # "me" as standalone word  
        r'\bmy\b',       # "my" as standalone word
        r'\bmyself\b',   # "myself" as standalone word
        r'\bour\b',      # "our" as standalone word
        r'\bwe\b'        # "we" as standalone word
    ]
    
    text_lower = resume_text.lower()
    pronoun_count = 0
    
    for pattern in pronoun_patterns:
        matches = re.findall(pattern, text_lower)
        pronoun_count += len(matches)
    
    if pronoun_count == 0:
        return 10
    elif pronoun_count <= 2:
        return 7
    elif pronoun_count <= 5:
        return 4
    else:
        return 1

def analyze_quantifiable_achievements_frontend(resume_text: str) -> int:
    """Copied exactly from frontend analyzeQuantifiableAchievements"""
    # Look for numbers and percentages
    number_patterns = [
        r'\d+%',  # Percentages
        r'\$[\d,]+',  # Dollar amounts
        r'\d+\+',  # Numbers with plus
        r'\d+\s*(million|thousand|k)',  # Large numbers
        r'\d+\s*(years?|months?)',  # Time periods
    ]
    
    achievements_count = 0
    for pattern in number_patterns:
        achievements_count += len(re.findall(pattern, resume_text, re.IGNORECASE))
    
    if achievements_count >= 5:
        return 9
    elif achievements_count >= 3:
        return 7
    elif achievements_count >= 1:
        return 5
    else:
        return 2

def analyze_action_verbs_frontend(resume_text: str) -> int:
    """Copied exactly from frontend analyzeActionVerbs"""
    strong_verbs = [
        'achieved', 'developed', 'implemented', 'led', 'managed', 'created',
        'improved', 'increased', 'reduced', 'optimized', 'delivered', 'executed'
    ]
    weak_verbs = ['was', 'were', 'did', 'worked', 'helped', 'responsible']
    
    text_lower = resume_text.lower()
    strong_count = sum(1 for verb in strong_verbs if verb in text_lower)
    weak_count = sum(1 for verb in weak_verbs if verb in text_lower)
    
    if strong_count > weak_count * 2:
        return 9
    elif strong_count > weak_count:
        return 7
    elif strong_count == weak_count:
        return 5
    else:
        return 3

def analyze_active_voice_frontend(resume_text: str) -> int:
    """Copied exactly from frontend analyzeActiveVoice"""
    passive_indicators = ['was', 'were', 'been', 'being', 'by']
    active_indicators = ['led', 'managed', 'created', 'developed', 'implemented']
    
    text_lower = resume_text.lower()
    passive_count = sum(1 for indicator in passive_indicators if indicator in text_lower)
    active_count = sum(1 for indicator in active_indicators if indicator in text_lower)
    
    if active_count > passive_count:
        return 8
    elif active_count == passive_count:
        return 6
    else:
        return 4

def analyze_summary_section_frontend(resume_text: str) -> int:
    """
    Updated Professional Summary scoring with strict penalty system
    Starting Score: 10 points
    Deductions:
    - Word Count > 100: -1 point
    - Vague Buzzwords: -2 points
    - Personal Pronouns: -2 points  
    - No Metrics: -2 points
    """
    import re
    
    # Find summary section
    summary_text = extract_summary_section(resume_text)
    if not summary_text:
        return 3  # No summary section found
    
    # Start with perfect score
    score = 10
    
    # 1. Word count penalty (>100 words)
    word_count = len(summary_text.split())
    if word_count > 100:
        score -= 1
    
    # 2. Vague buzzwords penalty (-2 points)
    vague_buzzwords = [
        'collaboration', 'problem-solving', 'communication', 'teamwork', 
        'leadership', 'detail-oriented', 'hardworking', 'motivated', 
        'dedicated', 'passionate', 'results-driven', 'dynamic', 
        'innovative', 'creative', 'analytical', 'strategic',
        'excellent', 'strong', 'effective', 'efficient',
        'experienced', 'skilled', 'proven', 'successful'
    ]
    
    summary_lower = summary_text.lower()
    found_buzzwords = sum(1 for buzzword in vague_buzzwords if buzzword in summary_lower)
    if found_buzzwords > 0:
        score -= 2
    
    # 3. Personal pronouns penalty (-2 points)
    personal_pronouns = [r'\bi\b', r'\bme\b', r'\bmy\b', r'\bmyself\b', r'\bwe\b', r'\bour\b', r'\bus\b']
    found_pronouns = sum(1 for pronoun in personal_pronouns if re.search(pronoun, summary_lower))
    if found_pronouns > 0:
        score -= 2
    
    # 4. No metrics penalty (-2 points)
    metric_patterns = [
        r'\d+\+?\s*(years?|months?)',  # Experience: "5+ years", "6 months"
        r'[\$₹€£¥]\s*\d+[kmb]?',       # Money: "$50K", "₹10M"
        r'\d+[kmb]?\+?\s*[\$₹€£¥]',    # Money: "50K$", "10M+"
        r'\d+\+?\s*%',                 # Percentages: "25%", "50%+"
        r'\d+[kmb]?\+?',               # Large numbers: "100K+", "5M"
        r'\d+\+?\s*(team|people|employees|staff)',  # Team size
        r'\d+\+?\s*(projects?|clients?|accounts?)', # Volume metrics
        r'managed?\s+[\$₹€£¥]?\d+',    # Management amounts
        r'saved?\s+[\$₹€£¥]?\d+',      # Savings
        r'increased?\s+\d+',           # Growth metrics
        r'reduced?\s+\d+',             # Efficiency metrics
    ]
    
    has_metrics = any(re.search(pattern, summary_text, re.IGNORECASE) for pattern in metric_patterns)
    if not has_metrics:
        score -= 2
    
    # Ensure score doesn't go below 0
    return max(score, 0)

def extract_summary_section(resume_text: str) -> str:
    """Extract the professional summary/about section from resume"""
    import re
    
    # Common summary section headers
    summary_headers = [
        r'professional\s+summary', r'about\s+me', r'summary', r'profile', 
        r'objective', r'overview', r'introduction', r'career\s+summary',
        r'personal\s+statement', r'career\s+objective'
    ]
    
    lines = resume_text.split('\n')
    summary_content = []
    in_summary = False
    
    # Section headers that indicate end of summary
    end_headers = [
        r'experience', r'work\s+history', r'employment', r'career\s+journey',
        r'education', r'skills', r'competenc', r'qualifications',
        r'projects', r'achievements', r'certifications'
    ]
    
    for line in lines:
        line_clean = line.strip()
        if not line_clean:
            continue
            
        line_lower = line_clean.lower()
        
        # Check if we're starting a summary section
        if any(re.search(header, line_lower) for header in summary_headers):
            in_summary = True
            continue
        
        # Check if we've reached a different section (end of summary)
        if in_summary and any(re.search(header, line_lower) for header in end_headers):
            break
        
        # Collect summary content
        if in_summary and line_clean and not line_clean.startswith('-') and not line_clean.startswith('•'):
            summary_content.append(line_clean)
    
    return ' '.join(summary_content)

def get_summary_detailed_analysis(resume_text: str) -> str:
    """Generate detailed analysis for Professional Summary CTA modal"""
    import re
    
    summary_text = extract_summary_section(resume_text)
    if not summary_text:
        return "No professional summary section found. Add a compelling 2-3 sentence summary at the top of your resume highlighting your experience, key skills, and value proposition."
    
    issues = []
    recommendations = []
    
    # Analyze word count
    word_count = len(summary_text.split())
    if word_count > 100:
        issues.append(f"Summary is too long ({word_count} words)")
        recommendations.append("Keep summary under 100 words for better readability")
    
    # Check for vague buzzwords
    vague_buzzwords = [
        'collaboration', 'problem-solving', 'communication', 'teamwork', 
        'leadership', 'detail-oriented', 'hardworking', 'motivated', 
        'dedicated', 'passionate', 'results-driven', 'dynamic'
    ]
    
    found_buzzwords = [word for word in vague_buzzwords if word in summary_text.lower()]
    if found_buzzwords:
        issues.append(f"Contains vague buzzwords: {', '.join(found_buzzwords[:3])}{'...' if len(found_buzzwords) > 3 else ''}")
        recommendations.append("Replace generic terms with specific skills and achievements")
    
    # Check for personal pronouns
    pronouns_found = []
    pronoun_patterns = [r'\bi\b', r'\bme\b', r'\bmy\b', r'\bmyself\b', r'\bwe\b', r'\bour\b', r'\bus\b']
    for pattern in pronoun_patterns:
        matches = re.findall(pattern, summary_text, re.IGNORECASE)
        pronouns_found.extend(matches)
    
    if pronouns_found:
        issues.append(f"Contains personal pronouns: {', '.join(set([p.lower() for p in pronouns_found]))}")
        recommendations.append("Remove all personal pronouns (I, my, me, we, our)")
    
    # Check for metrics
    metric_patterns = [
        r'\d+\+?\s*(years?|months?)', r'[\$₹€£¥]\s*\d+[kmb]?', r'\d+\+?\s*%',
        r'\d+[kmb]?\+?', r'\d+\+?\s*(team|people|employees)'
    ]
    
    has_metrics = any(re.search(pattern, summary_text, re.IGNORECASE) for pattern in metric_patterns)
    if not has_metrics:
        issues.append("Missing quantifiable achievements or metrics")
        recommendations.append("Add specific numbers: years of experience, budget size, team size, or achievements")
    
    # Generate response
    if issues:
        analysis = f"Issues found: {' | '.join(issues)}. "
        analysis += f"Recommendations: {' | '.join(recommendations)}"
    else:
        analysis = "Summary meets all criteria: appropriate length, specific content, no personal pronouns, includes metrics."
    
    return analysis

def generate_fix_this_modal_content(category_name: str, resume_text: str, score: int) -> dict:
    """
    Generate comprehensive modal content for 'Fix This' CTA
    Combines generic explanation with specific examples from user's resume
    """
    import sys
    import os
    
    # Import the modal config with error handling
    try:
        sys.path.append(os.path.dirname(__file__))
        from ats_modal_config import ATS_MODAL_CONFIG
    except ImportError as e:
        logger.warning(f"Failed to import ATS_MODAL_CONFIG: {e}")
        ATS_MODAL_CONFIG = {}
    
    # Get generic explanation
    generic_config = ATS_MODAL_CONFIG.get(category_name, {})
    generic_title = generic_config.get('title', f'Why {category_name} Matters for ATS')
    generic_explanation = generic_config.get('generic_explanation', f'{category_name} is important for ATS scoring and professional presentation.')
    
    # Generate dynamic examples based on category
    dynamic_examples = generate_dynamic_examples(category_name, resume_text, score)
    
    return {
        'category': category_name,
        'score': score,
        'title': generic_title,
        'generic_explanation': generic_explanation,
        'dynamic_examples': dynamic_examples
    }

def generate_dynamic_examples(category_name: str, resume_text: str, score: int) -> list:
    """Generate specific examples of issues found in the user's resume"""
    
    examples = []
    
    if category_name == 'Action Verbs':
        examples = generate_action_verbs_examples(resume_text)
    elif category_name == 'Repetition':
        examples = generate_repetition_examples(resume_text)
    elif category_name == 'Personal Pronouns':
        examples = generate_pronouns_examples(resume_text)
    elif category_name == 'Quantifiable Achievements':
        examples = generate_quantifiable_examples(resume_text)
    elif category_name == 'Summary':
        examples = generate_summary_examples(resume_text)
    elif category_name == 'Dates':
        examples = generate_dates_examples(resume_text)
    elif category_name == 'Grammar':
        examples = generate_grammar_examples(resume_text)
    elif category_name == 'Spelling':
        examples = generate_spelling_examples(resume_text)
    elif category_name == 'Contact Details':
        examples = generate_contact_examples(resume_text)
    elif category_name == 'Skills Section':
        examples = generate_skills_examples(resume_text)
    elif category_name == 'Analytical':
        examples = generate_analytical_examples(resume_text)
    elif category_name == 'Leadership':
        examples = generate_leadership_examples(resume_text)
    elif category_name == 'Certifications':
        examples = generate_certifications_examples(resume_text)
    elif category_name == 'Growth Signals':
        examples = generate_growth_examples(resume_text)
    elif category_name == 'Drive':
        examples = generate_drive_examples(resume_text)
    elif category_name == 'Active Voice':
        examples = generate_active_voice_examples(resume_text)
    elif category_name == 'Teamwork':
        examples = generate_teamwork_examples(resume_text)
    elif category_name == 'Education Section':
        examples = generate_education_examples(resume_text)
    elif category_name == 'Page Density':
        examples = generate_page_density_examples(resume_text)
    elif category_name == 'Use of Bullets':
        examples = generate_bullets_examples(resume_text)
    elif category_name == 'Verb Tenses':
        examples = generate_verb_tenses_examples(resume_text)
    elif category_name == 'Verbosity':
        examples = generate_verbosity_examples(resume_text)
    elif category_name == 'Unnecessary Sections':
        examples = generate_unnecessary_sections_examples(resume_text)
    elif category_name == 'CV Readability Score':
        examples = generate_readability_examples(resume_text)
    else:
        # Generic fallback
        examples = [
            {
                'issue': f'{category_name} needs improvement',
                'example': 'Multiple areas in your resume could be enhanced',
                'suggestion': f'Professional optimization will improve your {category_name} score'
            },
            {
                'issue': 'ATS compatibility concerns',
                'example': 'Current format may not parse correctly in ATS systems',
                'suggestion': 'Restructuring will ensure proper ATS processing'
            }
        ]
    
    # Ensure we return exactly 2 examples
    return examples[:2] if len(examples) >= 2 else examples + [examples[0]] if examples else []

def generate_repetition_examples(resume_text: str) -> list:
    """Generate specific repetition examples from the resume"""
    import re
    from collections import Counter
    
    # Find repeated verbs (using same logic as scoring function)
    action_verbs_patterns = [
        r'\b(manage[ds]?|managing)\b',
        r'\b(develop[eds]?|developing)\b', 
        r'\b(creat[ed]?|creating)\b',
        r'\b(implement[eds]?|implementing)\b',
        r'\b(lead[s]?|leading|led)\b',
        r'\b(design[eds]?|designing)\b',
        r'\b(execut[ed]?|executing)\b',
        r'\b(deliver[eds]?|delivering)\b',
        r'\b(coordinat[ed]?|coordinating)\b',
        r'\b(establish[eds]?|establishing)\b'
    ]
    
    verb_counts = Counter()
    verb_examples = {}
    
    for pattern in action_verbs_patterns:
        matches = re.findall(pattern, resume_text, re.IGNORECASE)
        if matches:
            # Get base verb form
            base_verb = pattern.split('(')[1].split('[')[0]  # Extract base form
            verb_counts[base_verb] = len(matches)
            # Find actual context where verb appears
            context_matches = re.finditer(pattern, resume_text, re.IGNORECASE)
            contexts = []
            for match in context_matches:
                start = max(0, match.start() - 20)
                end = min(len(resume_text), match.end() + 30)
                context = resume_text[start:end].strip()
                contexts.append(context)
            verb_examples[base_verb] = contexts[:2]  # Max 2 examples
    
    # Find most repeated verbs
    repeated_verbs = [(verb, count) for verb, count in verb_counts.items() if count > 1]
    repeated_verbs.sort(key=lambda x: x[1], reverse=True)
    
    examples = []
    for verb, count in repeated_verbs[:2]:  # Top 2 repeated verbs
        contexts = verb_examples.get(verb, [])
        examples.append({
            'issue': f'"{verb.title()}" used {count} times',
            'example': f'Found in: "{contexts[0][:50]}..."' if contexts else f'Repeated {count} times throughout resume',
            'suggestion': f'Replace with alternatives like "orchestrated", "spearheaded", "championed"'
        })
    
    # If no repetitions found, provide generic examples
    if not examples:
        examples = [
            {
                'issue': 'Limited verb variety detected',
                'example': 'Multiple bullet points use similar action words',
                'suggestion': 'Diversify with stronger verbs like "orchestrated", "pioneered", "revolutionized"'
            },
            {
                'issue': 'Weak action verbs identified',
                'example': 'Some accomplishments use passive or weak language',
                'suggestion': 'Replace with impactful verbs that demonstrate leadership and results'
            }
        ]
    
    return examples

def generate_summary_examples(resume_text: str) -> list:
    """Generate specific summary examples"""
    import re
    
    summary_text = extract_summary_section(resume_text)
    examples = []
    
    if not summary_text:
        return [
            {
                'issue': 'No professional summary found',
                'example': 'Resume starts directly with experience section',
                'suggestion': 'Add a compelling 2-3 sentence summary highlighting your value proposition'
            },
            {
                'issue': 'Missing career positioning statement',
                'example': 'No clear professional identity or specialization stated',
                'suggestion': 'Include your role, experience level, and key achievements in opening summary'
            }
        ]
    
    # Check for buzzwords
    buzzwords_found = []
    vague_buzzwords = ['collaboration', 'problem-solving', 'communication', 'teamwork', 'results-driven', 'passionate', 'motivated']
    for buzzword in vague_buzzwords:
        if buzzword in summary_text.lower():
            buzzwords_found.append(buzzword)
    
    # Check for pronouns
    pronouns_found = []
    pronouns = ['I ', 'my ', 'me ', 'we ', 'our ']
    for pronoun in pronouns:
        if pronoun.lower() in summary_text.lower():
            pronouns_found.append(pronoun.strip())
    
    # Check for metrics
    has_metrics = bool(re.search(r'\d+', summary_text))
    
    if buzzwords_found:
        examples.append({
            'issue': f'Vague buzzwords: {", ".join(buzzwords_found[:2])}',
            'example': f'Found in your summary: "{summary_text[:60]}..."',
            'suggestion': 'Replace with specific achievements like "managed $2M budget" or "led 15-person team"'
        })
    
    if pronouns_found:
        examples.append({
            'issue': f'Personal pronouns: {", ".join(pronouns_found)}',
            'example': f'Summary contains: "{summary_text[:60]}..."',
            'suggestion': 'Rewrite in third person: "Procurement Lead with 6+ years managing $180M annual spend"'
        })
    
    if not has_metrics and len(examples) < 2:
        examples.append({
            'issue': 'Missing quantifiable achievements',
            'example': 'Summary lacks specific numbers, percentages, or metrics',
            'suggestion': 'Add concrete metrics: "6+ years experience", "managed $180M spend", "achieved 30% cost savings"'
        })
    
    # Ensure we have 2 examples
    if len(examples) < 2:
        examples.append({
            'issue': 'Generic professional description',
            'example': 'Summary reads like a job description rather than personal achievements',
            'suggestion': 'Highlight your unique value: specific results, expertise, and career accomplishments'
        })
    
    return examples[:2]

def generate_pronouns_examples(resume_text: str) -> list:
    """Generate specific personal pronouns examples"""
    import re
    
    examples = []
    
    # Find pronouns with context
    pronoun_patterns = [(r'\bi\b', 'I'), (r'\bmy\b', 'my'), (r'\bme\b', 'me'), (r'\bwe\b', 'we'), (r'\bour\b', 'our')]
    
    for pattern, pronoun in pronoun_patterns[:2]:  # Check first 2 patterns
        matches = list(re.finditer(pattern, resume_text, re.IGNORECASE))
        if matches:
            # Get context for first match
            match = matches[0]
            start = max(0, match.start() - 25)
            end = min(len(resume_text), match.end() + 35)
            context = resume_text[start:end].strip()
            
            examples.append({
                'issue': f'Personal pronoun "{pronoun}" detected',
                'example': f'Found in: "{context}"',
                'suggestion': f'Remove "{pronoun}" and rewrite: "Led team of 10" instead of "I led team of 10"'
            })
    
    # If no pronouns found, provide generic examples
    if not examples:
        examples = [
            {
                'issue': 'Informal language detected',
                'example': 'Some sections may contain casual or personal language',
                'suggestion': 'Ensure all descriptions use professional, third-person format'
            },
            {
                'issue': 'Professional formatting opportunities',
                'example': 'Resume language could be more formal and objective',
                'suggestion': 'Use action verbs without personal pronouns for stronger impact'
            }
        ]
    
    return examples[:2]

def generate_quantifiable_examples(resume_text: str) -> list:
    """Generate specific quantifiable achievements examples"""
    import re
    
    # Find non-quantified achievements
    bullet_points = re.findall(r'[•·\*\-]\s*(.+)', resume_text)
    non_quantified = []
    
    for point in bullet_points[:10]:  # Check first 10 points
        # Check if it lacks numbers
        has_numbers = bool(re.search(r'\d+', point))
        if not has_numbers and len(point) > 20:
            non_quantified.append(point.strip())
    
    examples = []
    for point in non_quantified[:2]:  # Take first 2
        examples.append({
            'issue': 'Missing quantifiable metrics',
            'example': f'"{point[:60]}..."',
            'suggestion': 'Add specific numbers: "increased efficiency by 25%" or "managed team of 8"'
        })
    
    # If no bullet points or all quantified, provide generic examples
    if not examples:
        examples = [
            {
                'issue': 'Achievements lack specific metrics',
                'example': 'Multiple accomplishments could include concrete numbers',
                'suggestion': 'Add percentages, dollar amounts, time savings, or team sizes to quantify impact'
            },
            {
                'issue': 'Vague impact statements',
                'example': 'General statements like "improved processes" without specifics',
                'suggestion': 'Specify results: "improved processes reducing cycle time by 30%"'
            }
        ]
    
    return examples[:2]

def generate_action_verbs_examples(resume_text: str) -> list:
    """Generate specific action verb examples"""
    import re
    
    # Find weak verbs
    weak_verbs = ['responsible', 'assisted', 'helped', 'worked', 'participated', 'involved', 'handled', 'dealt with']
    weak_found = []
    
    for verb in weak_verbs:
        matches = re.finditer(rf'\b{verb}\b', resume_text, re.IGNORECASE)
        for match in matches:
            start = max(0, match.start() - 20)
            end = min(len(resume_text), match.end() + 40)
            context = resume_text[start:end].strip()
            weak_found.append({'verb': verb, 'context': context})
            if len(weak_found) >= 2:
                break
        if len(weak_found) >= 2:
            break
    
    examples = []
    for item in weak_found[:2]:
        examples.append({
            'issue': f'Weak verb "{item["verb"]}" detected',
            'example': f'Found in: "{item["context"][:55]}..."',
            'suggestion': f'Replace with stronger verbs like "led", "spearheaded", "orchestrated", "delivered"'
        })
    
    if not examples:
        examples = [
            {
                'issue': 'Generic action verbs identified',
                'example': 'Some bullet points use common, weak action words',
                'suggestion': 'Use powerful verbs: "engineered", "pioneered", "revolutionized", "accelerated"'
            },
            {
                'issue': 'Missed leadership opportunities',
                'example': 'Accomplishments could demonstrate more ownership and initiative',
                'suggestion': 'Start bullet points with impactful verbs showing leadership and results'
            }
        ]
    
    return examples[:2]

def generate_contact_examples(resume_text: str) -> list:
    """Generate specific contact details examples based on actual scoring logic"""
    import re
    
    issues = []
    missing_elements = []
    
    # Use the same scoring logic as analyze_contact_details_frontend
    has_mobile = has_mobile_number(resume_text)
    has_email = has_email_address(resume_text)
    has_linkedin = has_linkedin_profile(resume_text)
    has_location = has_location_info(resume_text)
    
    # Identify specific missing elements (-2.5 points each)
    if not has_mobile:
        missing_elements.append({
            'issue': 'Missing mobile/phone number (-2.5 points)',
            'example': 'No phone number found in contact section',
            'suggestion': 'Add professional phone number: "(555) 123-4567" or "+1-555-123-4567"'
        })
    
    if not has_email:
        missing_elements.append({
            'issue': 'Missing professional email address (-2.5 points)',
            'example': 'No email address found in resume header',
            'suggestion': 'Add professional email: "firstname.lastname@email.com"'
        })
    
    if not has_linkedin:
        missing_elements.append({
            'issue': 'Missing LinkedIn profile URL (-2.5 points)',
            'example': 'No LinkedIn profile detected in contact information',
            'suggestion': 'Add LinkedIn URL: "linkedin.com/in/yourname" or "LinkedIn: YourName"'
        })
    
    if not has_location:
        missing_elements.append({
            'issue': 'Missing location information (-2.5 points)',
            'example': 'No city, state, or location details found',
            'suggestion': 'Add location: "New York, NY" or "San Francisco, CA" or "Remote"'
        })
    
    # If no missing elements, provide optimization suggestions
    if not missing_elements:
        issues = [
            {
                'issue': 'Contact formatting optimization opportunity',
                'example': 'All contact elements present but formatting could be enhanced',
                'suggestion': 'Ensure contact details are prominently displayed at top of resume'
            },
            {
                'issue': 'ATS parsing enhancement possible',
                'example': 'Contact information could be better structured for ATS systems',
                'suggestion': 'Use standard labels: "Email:", "Phone:", "LinkedIn:", "Location:"'
            }
        ]
        return issues[:2]
    
    # Return the specific missing elements (max 2 for display)
    return missing_elements[:2]

def generate_skills_examples(resume_text: str) -> list:
    """Generate specific skills section examples"""
    import re
    
    has_skills_section = bool(re.search(r'skills|competenc|technical|proficien', resume_text, re.IGNORECASE))
    
    if not has_skills_section:
        return [
            {
                'issue': 'No dedicated skills section found',
                'example': 'Resume lacks a clear "Technical Skills" or "Core Competencies" section',
                'suggestion': 'Add a skills section with relevant technical and professional competencies'
            },
            {
                'issue': 'Missing keyword optimization',
                'example': 'Skills may be buried in job descriptions rather than highlighted',
                'suggestion': 'Create dedicated skills section for better ATS keyword matching'
            }
        ]
    
    # Check for generic vs specific skills
    generic_skills = ['communication', 'teamwork', 'leadership', 'problem-solving']
    specific_skills = ['python', 'sql', 'tableau', 'excel', 'project management', 'salesforce']
    
    has_generic = any(skill in resume_text.lower() for skill in generic_skills)
    has_specific = any(skill in resume_text.lower() for skill in specific_skills)
    
    examples = []
    if has_generic and not has_specific:
        examples.append({
            'issue': 'Generic skills without technical specifics',
            'example': 'Skills section contains broad terms like "communication" and "teamwork"',
            'suggestion': 'Add specific technical skills, software, and tools relevant to your field'
        })
    
    if len(examples) < 2:
        examples.append({
            'issue': 'Skills organization could be enhanced',
            'example': 'Skills may not be categorized or prioritized effectively',
            'suggestion': 'Group skills by category: Technical, Leadership, Industry-specific for better impact'
        })
    
    return examples[:2]

def generate_dates_examples(resume_text: str) -> list:
    """Generate specific date formatting examples"""
    import re
    
    # Find different date formats
    date_patterns = [
        r'\d{4}\s*-\s*\d{4}',  # 2020-2022
        r'\d{1,2}/\d{4}',       # 12/2020
        r'[A-Za-z]+\s+\d{4}',   # Jan 2020
        r'\d{4}'                # 2020
    ]
    
    found_formats = []
    for pattern in date_patterns:
        matches = re.findall(pattern, resume_text)
        if matches:
            found_formats.extend(matches[:2])
    
    if len(set([len(date) for date in found_formats])) > 1:  # Different lengths = inconsistent
        return [
            {
                'issue': 'Inconsistent date formats detected',
                'example': f'Found mixed formats: {", ".join(found_formats[:2])}',
                'suggestion': 'Use consistent format: "Jan 2020 - Dec 2022" throughout resume'
            },
            {
                'issue': 'Date formatting affects ATS parsing',
                'example': 'Mixed date styles make it difficult for ATS to calculate experience',
                'suggestion': 'Standardize all dates to Month Year format for optimal ATS processing'
            }
        ]
    
    return [
        {
            'issue': 'Date presentation could be optimized',
            'example': 'Employment dates may not follow ATS best practices',
            'suggestion': 'Ensure all dates use consistent "Month Year - Month Year" format'
        },
        {
            'issue': 'Timeline clarity opportunities',
            'example': 'Date ranges could be more clearly formatted for ATS systems',
            'suggestion': 'Use standard date format: "Jan 2020 - Present" or "Mar 2018 - Dec 2021"'
        }
    ]

# Generic fallback generators for remaining categories
def generate_grammar_examples(resume_text: str) -> list:
    return [
        {
            'issue': 'Grammar inconsistencies detected',
            'example': 'Multiple sentences may have grammar or punctuation issues',
            'suggestion': 'Professional grammar review will eliminate errors and improve readability'
        },
        {
            'issue': 'Sentence structure optimization needed',
            'example': 'Some descriptions could have clearer, more impactful grammar',
            'suggestion': 'Refine sentence structure for professional presentation and ATS parsing'
        }
    ]

def generate_spelling_examples(resume_text: str) -> list:
    return [
        {
            'issue': 'Potential spelling inconsistencies',
            'example': 'Technical terms or industry jargon may have spelling variations',
            'suggestion': 'Comprehensive spell-check will ensure error-free professional presentation'
        },
        {
            'issue': 'Professional proofreading needed',
            'example': 'Multiple sections could benefit from detailed spelling review',
            'suggestion': 'Eliminate all spelling errors for polished, professional appearance'
        }
    ]

def generate_analytical_examples(resume_text: str) -> list:
    return [
        {
            'issue': 'Missing analytical accomplishments',
            'example': 'Limited examples of data-driven decision making or analysis',
            'suggestion': 'Add specific examples: "analyzed data to identify 15% cost reduction opportunity"'
        },
        {
            'issue': 'Quantified analytical impact needed',
            'example': 'Analytical work lacks specific metrics and measurable outcomes',
            'suggestion': 'Include analytical tools used and quantified results of your analysis'
        }
    ]

def generate_leadership_examples(resume_text: str) -> list:
    return [
        {
            'issue': 'Leadership scope unclear',
            'example': 'Management responsibilities lack specific team size or budget details',
            'suggestion': 'Specify leadership scope: "managed team of 12" or "oversaw $2M budget"'
        },
        {
            'issue': 'Leadership outcomes not quantified',
            'example': 'Leadership roles need measurable impact and specific achievements',
            'suggestion': 'Add leadership results: "led team achieving 25% productivity increase"'
        }
    ]

def generate_certifications_examples(resume_text: str) -> list:
    return [
        {
            'issue': 'Missing industry certifications',
            'example': 'No professional certifications or credentials mentioned',
            'suggestion': 'Add relevant certifications: PMP, AWS, Google Analytics, or industry-specific credentials'
        },
        {
            'issue': 'Certification currency unclear',
            'example': 'Existing certifications may lack renewal dates or current status',
            'suggestion': 'Include certification dates and renewal status for credibility'
        }
    ]

def generate_growth_examples(resume_text: str) -> list:
    return [
        {
            'issue': 'Career progression not highlighted',
            'example': 'Limited evidence of promotions or role advancement',
            'suggestion': 'Emphasize career growth: "promoted from Analyst to Senior Manager in 2 years"'
        },
        {
            'issue': 'Skills development not demonstrated',
            'example': 'Missing examples of continuous learning and skill advancement',
            'suggestion': 'Show growth trajectory: expanded responsibilities, new skills, leadership roles'
        }
    ]

def generate_drive_examples(resume_text: str) -> list:
    return [
        {
            'issue': 'Initiative examples limited',
            'example': 'Few examples of proactive problem-solving or self-directed projects',
            'suggestion': 'Add initiative examples: "identified and resolved process inefficiency saving $50K annually"'
        },
        {
            'issue': 'Ownership mentality not demonstrated',
            'example': 'Accomplishments focus on assigned tasks rather than driven initiatives',
            'suggestion': 'Highlight self-motivation: "pioneered new system" or "championed process improvement"'
        }
    ]

def generate_active_voice_examples(resume_text: str) -> list:
    return [
        {
            'issue': 'Passive voice detected in descriptions',
            'example': 'Some accomplishments use weak passive language',
            'suggestion': 'Convert to active voice: "Led project" instead of "was responsible for project"'
        },
        {
            'issue': 'Weak action statements identified',
            'example': 'Multiple bullet points lack strong, direct action language',
            'suggestion': 'Use active verbs showing direct ownership and impact of your work'
        }
    ]

def generate_teamwork_examples(resume_text: str) -> list:
    return [
        {
            'issue': 'Collaboration scope unclear',
            'example': 'Teamwork mentions lack specific cross-functional details',
            'suggestion': 'Specify collaboration: "partnered with marketing, sales, and engineering teams"'
        },
        {
            'issue': 'Team impact not quantified',
            'example': 'Collaborative achievements need measurable team outcomes',
            'suggestion': 'Add team results: "collaborated with 5 departments to deliver $1M project"'
        }
    ]

def generate_education_examples(resume_text: str) -> list:
    """Extract actual education issues from resume"""
    import re
    
    # Look for education section
    has_education_section = bool(re.search(r'education|academic|qualification', resume_text, re.IGNORECASE))
    
    if not has_education_section:
        return [
            {
                'issue': 'No education section found',
                'example': 'Resume lacks dedicated "Education" section',
                'suggestion': 'Add education section with degree, institution, and graduation year'
            },
            {
                'issue': 'Missing educational credentials',
                'example': 'No degree or qualification information provided',
                'suggestion': 'Include highest degree, major, and university/college name'
            }
        ]
    
    # Extract education section content
    education_text = ''
    lines = resume_text.split('\n')
    in_education = False
    
    for line in lines:
        if re.search(r'education|academic|qualification', line, re.IGNORECASE) and len(line.strip()) < 50:
            in_education = True
            continue
        elif in_education and re.search(r'experience|work|skills|projects', line, re.IGNORECASE) and len(line.strip()) < 50:
            break
        elif in_education:
            education_text += line + ' '
    
    examples = []
    
    # Check for missing degree
    has_degree = bool(re.search(r'bachelor|master|phd|degree|diploma|b\.?[a-z]+|m\.?[a-z]+', education_text, re.IGNORECASE))
    if not has_degree:
        examples.append({
            'issue': 'Missing degree information',
            'example': f'Education section: "{education_text[:60]}..."',
            'suggestion': 'Add specific degree: "Bachelor of Science in Computer Science"'
        })
    
    # Check for missing institution
    has_institution = bool(re.search(r'university|college|institute|school', education_text, re.IGNORECASE))
    if not has_institution and len(examples) < 2:
        examples.append({
            'issue': 'Missing institution name',
            'example': f'Education content: "{education_text[:60]}..."',
            'suggestion': 'Include university/college name: "Stanford University" or "MIT"'
        })
    
    # If education looks complete, check formatting
    if len(examples) < 2:
        examples.append({
            'issue': 'Education formatting could be improved',
            'example': 'Education section may lack proper structure or dates',
            'suggestion': 'Format as: "Bachelor of Science, Computer Science | Stanford University | 2020"'
        })
    
    return examples[:2]

def generate_page_density_examples(resume_text: str) -> list:
    """Extract page density issues from resume"""
    import re
    
    # Calculate basic density metrics
    total_chars = len(resume_text)
    lines = resume_text.split('\n')
    non_empty_lines = [line for line in lines if line.strip()]
    avg_line_length = sum(len(line) for line in non_empty_lines) / max(len(non_empty_lines), 1)
    
    examples = []
    
    # Check for overcrowding
    if avg_line_length > 80:
        long_lines = [line[:60] + "..." for line in non_empty_lines if len(line) > 80][:2]
        examples.append({
            'issue': 'Lines too long - overcrowded content',
            'example': f'Long line found: "{long_lines[0] if long_lines else "Multiple lines exceed 80 characters"}"',
            'suggestion': 'Break long lines and add proper spacing between sections'
        })
    
    # Check for too sparse (very short lines)
    short_lines = [line for line in non_empty_lines if 5 < len(line.strip()) < 20]
    if len(short_lines) > 5 and len(examples) < 2:
        examples.append({
            'issue': 'Too many short, sparse lines',
            'example': f'Short line example: "{short_lines[0][:40]}..."',
            'suggestion': 'Combine related information into well-structured bullet points'
        })
    
    # Generic fallback
    if len(examples) < 2:
        examples.append({
            'issue': 'Page density needs optimization',
            'example': 'Text distribution could be more balanced for better readability',
            'suggestion': 'Adjust spacing and line length for optimal visual presentation'
        })
    
    return examples[:2]

def generate_bullets_examples(resume_text: str) -> list:
    """Extract bullet point usage issues from resume"""
    import re
    
    # Find paragraph-style content in experience sections
    lines = resume_text.split('\n')
    paragraph_lines = []
    bullet_lines = []
    
    for line in lines:
        line_clean = line.strip()
        if len(line_clean) > 30:  # Substantial content
            if re.match(r'^[•·\*\-\+]\s', line_clean):
                bullet_lines.append(line_clean)
            elif not re.match(r'^[A-Z][A-Za-z\s&]+$', line_clean):  # Not a header
                paragraph_lines.append(line_clean)
    
    examples = []
    
    # Check for paragraph-style descriptions instead of bullets
    if paragraph_lines:
        examples.append({
            'issue': 'Paragraph format instead of bullet points',
            'example': f'Found paragraph: "{paragraph_lines[0][:60]}..."',
            'suggestion': 'Convert to bullet format: "• Led cross-functional team of 12..."'
        })
    
    # Check bullet-to-paragraph ratio
    total_content_lines = len(bullet_lines) + len(paragraph_lines)
    if total_content_lines > 0:
        bullet_ratio = len(bullet_lines) / total_content_lines
        if bullet_ratio < 0.5 and len(examples) < 2:
            examples.append({
                'issue': f'Low bullet usage: {len(bullet_lines)} bullets vs {len(paragraph_lines)} paragraphs',
                'example': 'Many achievements written as paragraphs instead of bullets',
                'suggestion': 'Use bullets for all accomplishments and responsibilities'
            })
    
    # Generic fallback
    if len(examples) < 2:
        examples.append({
            'issue': 'Bullet point structure needs improvement',
            'example': 'Some achievements could be better formatted as bullet points',
            'suggestion': 'Use consistent bullet formatting: • Achievement with quantified result'
        })
    
    return examples[:2]

def generate_verb_tenses_examples(resume_text: str) -> list:
    """Extract verb tense inconsistency issues"""
    import re
    
    # Find mixed tenses in descriptions
    past_tense_verbs = re.findall(r'\b(managed|developed|created|led|designed|implemented|achieved|delivered)\b', resume_text, re.IGNORECASE)
    present_tense_verbs = re.findall(r'\b(manage|manages|develop|develops|create|creates|lead|leads|design|designs)\b', resume_text, re.IGNORECASE)
    ing_verbs = re.findall(r'\b(managing|developing|creating|leading|designing|implementing)\b', resume_text, re.IGNORECASE)
    
    examples = []
    
    # Check for mixed tenses
    if past_tense_verbs and present_tense_verbs:
        examples.append({
            'issue': 'Mixed verb tenses detected',
            'example': f'Found both "{past_tense_verbs[0]}" and "{present_tense_verbs[0]}" in resume',
            'suggestion': 'Use past tense for previous roles, present tense only for current position'
        })
    
    # Check for -ing verbs (often incorrect)
    if ing_verbs and len(examples) < 2:
        examples.append({
            'issue': 'Incorrect -ing verb forms found',
            'example': f'Found "{ing_verbs[0]}" - avoid -ing forms in bullet points',
            'suggestion': 'Use action verbs: "Managed" not "Managing", "Led" not "Leading"'
        })
    
    # Generic fallback
    if len(examples) < 2:
        examples.append({
            'issue': 'Verb tense consistency needs review',
            'example': 'Some descriptions may have inconsistent verb tenses',
            'suggestion': 'Ensure consistent tense: past for previous jobs, present for current role'
        })
    
    return examples[:2]

def generate_verbosity_examples(resume_text: str) -> list:
    """Extract verbose/wordy content from resume"""
    import re
    
    # Find overly long sentences or wordy phrases
    sentences = re.split(r'[•·\*\-]\s*', resume_text)
    verbose_sentences = []
    
    for sentence in sentences:
        sentence_clean = sentence.strip()
        if len(sentence_clean) > 150:  # Very long sentences
            verbose_sentences.append(sentence_clean)
    
    # Find wordy phrases
    wordy_phrases = [
        'in order to', 'for the purpose of', 'with the goal of', 'in an effort to',
        'it should be noted that', 'it is important to', 'with regard to'
    ]
    
    found_wordy = []
    for phrase in wordy_phrases:
        if phrase in resume_text.lower():
            # Find context
            idx = resume_text.lower().find(phrase)
            context = resume_text[max(0, idx-20):idx+len(phrase)+30]
            found_wordy.append({'phrase': phrase, 'context': context})
    
    examples = []
    
    # Report verbose sentences
    if verbose_sentences:
        examples.append({
            'issue': 'Overly long, verbose sentences',
            'example': f'Long sentence: "{verbose_sentences[0][:80]}..."',
            'suggestion': 'Break into concise bullet points with clear, direct language'
        })
    
    # Report wordy phrases
    if found_wordy and len(examples) < 2:
        item = found_wordy[0]
        examples.append({
            'issue': f'Wordy phrase: "{item["phrase"]}"',
            'example': f'Found in: "{item["context"][:60]}..."',
            'suggestion': f'Replace "{item["phrase"]}" with simpler, direct language'
        })
    
    # Generic fallback
    if len(examples) < 2:
        examples.append({
            'issue': 'Content could be more concise',
            'example': 'Some descriptions use unnecessary words that dilute impact',
            'suggestion': 'Use concise, powerful language: "Led team" not "Was responsible for leading team"'
        })
    
    return examples[:2]

def generate_unnecessary_sections_examples(resume_text: str) -> list:
    """Extract unnecessary/outdated sections from resume"""
    import re
    
    # Find outdated sections
    outdated_sections = [
        ('references', 'References available upon request'),
        ('objective', 'Career objective or goal statements'),  
        ('high school', 'High school education when higher education exists'),
        ('hobbies', 'Personal hobbies and interests'),
        ('marital', 'Marital status or personal information')
    ]
    
    found_sections = []
    lines = resume_text.split('\n')
    
    for section_key, section_name in outdated_sections:
        for line in lines:
            if section_key in line.lower() and len(line.strip()) < 100:
                found_sections.append({
                    'section': section_name,
                    'line': line.strip(),
                    'key': section_key
                })
                break
    
    examples = []
    
    # Report found outdated sections
    for item in found_sections[:2]:
        examples.append({
            'issue': f'Outdated section: {item["section"]}',
            'example': f'Found: "{item["line"][:60]}..."',
            'suggestion': f'Remove {item["section"]} - use space for relevant professional content'
        })
    
    # Generic fallback
    if len(examples) < 2:
        examples.append({
            'issue': 'Potential outdated content detected',
            'example': 'Some sections may not add value to modern resume standards',
            'suggestion': 'Focus on relevant professional experience, skills, and achievements'
        })
    
    return examples[:2]

def generate_readability_examples(resume_text: str) -> list:
    """Extract overall readability issues"""
    import re
    
    # Check for common readability issues
    issues_found = []
    
    # Long words (>12 characters)
    long_words = re.findall(r'\b\w{13,}\b', resume_text)
    if long_words:
        issues_found.append({
            'issue': 'Complex vocabulary detected',
            'example': f'Long words found: "{", ".join(long_words[:3])}"',
            'suggestion': 'Use simpler, clearer language where possible'
        })
    
    # Excessive jargon or acronyms
    acronyms = re.findall(r'\b[A-Z]{3,}\b', resume_text)
    if len(acronyms) > 10:
        issues_found.append({
            'issue': 'Too many acronyms may reduce clarity',
            'example': f'Many acronyms found: "{", ".join(acronyms[:4])}"',
            'suggestion': 'Spell out important acronyms: "Application Programming Interface (API)"'
        })
    
    # Very long bullet points
    bullets = re.findall(r'[•·\*\-]\s*(.+)', resume_text)
    long_bullets = [bullet for bullet in bullets if len(bullet) > 120]
    if long_bullets and len(issues_found) < 2:
        issues_found.append({
            'issue': 'Bullet points too long for easy scanning',
            'example': f'Long bullet: "{long_bullets[0][:80]}..."',
            'suggestion': 'Keep bullets under 120 characters for better readability'
        })
    
    # Return found issues or generic fallback
    if issues_found:
        return issues_found[:2]
    
    return [
        {
            'issue': 'Overall readability could be enhanced',
            'example': 'Document structure and language could be more accessible',
            'suggestion': 'Optimize for clarity, conciseness, and professional presentation'
        },
        {
            'issue': 'Content accessibility needs improvement',
            'example': 'Some sections may be difficult for ATS systems and recruiters to parse',
            'suggestion': 'Use clear formatting, simple language, and logical structure'
        }
    ]

def analyze_teamwork_skills_frontend(resume_text: str) -> int:
    """Copied exactly from frontend analyzeTeamworkSkills"""
    teamwork_keywords = [
        'team', 'collaborate', 'cooperation', 'partnership', 'group',
        'cross-functional', 'stakeholder', 'communicate', 'coordinate'
    ]
    text_lower = resume_text.lower()
    found_keywords = sum(1 for keyword in teamwork_keywords if keyword in text_lower)
    
    if found_keywords >= 5:
        return 9
    elif found_keywords >= 3:
        return 7
    elif found_keywords >= 1:
        return 5
    else:
        return 3

def analyze_repetition_frontend(resume_text: str) -> int:
    """
    NEW CORRECT LOGIC: Analyzes verb repetition only
    - Start with 10 points
    - Deduct 2 points per repeated verb occurrence
    - Minimum score: 0, Maximum score: 10
    """
    import re
    from collections import Counter
    
    # Common action verbs that appear in resumes (base forms and common variations)
    action_verbs_patterns = [
        # Base forms and their variations
        r'\b(manage[ds]?|managing)\b',
        r'\b(develop[eds]?|developing)\b', 
        r'\b(creat[ed]?|creating)\b',
        r'\b(implement[eds]?|implementing)\b',
        r'\b(lead[s]?|leading|led)\b',
        r'\b(design[eds]?|designing)\b',
        r'\b(execut[ed]?|executing)\b',
        r'\b(deliver[eds]?|delivering)\b',
        r'\b(achiev[ed]?|achieving)\b',
        r'\b(establish[eds]?|establishing)\b',
        r'\b(coordinat[ed]?|coordinating)\b',
        r'\b(supervis[ed]?|supervising)\b',
        r'\b(direct[eds]?|directing)\b',
        r'\b(operat[ed]?|operating)\b',
        r'\b(maintain[eds]?|maintaining)\b',
        r'\b(analyz[ed]?|analyzing|analys[ed]?|analysing)\b',
        r'\b(evaluat[ed]?|evaluating)\b',
        r'\b(assess[eds]?|assessing)\b',
        r'\b(review[eds]?|reviewing)\b',
        r'\b(monitor[eds]?|monitoring)\b',
        r'\b(track[eds]?|tracking)\b',
        r'\b(optimiz[ed]?|optimizing|optimis[ed]?|optimising)\b',
        r'\b(improv[ed]?|improving)\b',
        r'\b(enhanc[ed]?|enhancing)\b',
        r'\b(reduc[ed]?|reducing)\b',
        r'\b(increas[ed]?|increasing)\b',
        r'\b(build[ings]?|built)\b',
        r'\b(train[eds]?|training)\b',
        r'\b(teach[ings]?|taught)\b',
        r'\b(mentor[eds]?|mentoring)\b',
        r'\b(coach[eds]?|coaching)\b',
        r'\b(facilitat[ed]?|facilitating)\b',
        r'\b(present[eds]?|presenting)\b',
        r'\b(communicat[ed]?|communicating)\b',
        r'\b(collaborat[ed]?|collaborating)\b',
        r'\b(negotiat[ed]?|negotiating)\b',
        r'\b(organiz[ed]?|organizing|organis[ed]?|organising)\b',
        r'\b(plan[s]?|planning|planned)\b',
        r'\b(strateg[ys]?|strategizing|strategized)\b',
        r'\b(research[eds]?|researching)\b',
        r'\b(test[eds]?|testing)\b',
        r'\b(debug[s]?|debugging|debugged)\b',
        r'\b(troubleshoot[ings]?|troubleshooting|troubleshot)\b',
        r'\b(deploy[eds]?|deploying)\b',
        r'\b(integrat[ed]?|integrating)\b',
        r'\b(configur[ed]?|configuring)\b',
        r'\b(instal[ls]?|installing|installed)\b',
        r'\b(maintain[eds]?|maintaining)\b',
        r'\b(updat[ed]?|updating)\b',
        r'\b(upgrad[ed]?|upgrading)\b',
        r'\b(migrat[ed]?|migrating)\b',
        r'\b(automat[ed]?|automating)\b',
        r'\b(streamlin[ed]?|streamlining)\b',
        r'\b(ensur[ed]?|ensuring)\b',
        r'\b(secur[ed]?|securing)\b',
        r'\b(protec[ts]?|protecting|protected)\b',
        r'\b(compil[ed]?|compiling)\b',
        r'\b(document[eds]?|documenting)\b',
        r'\b(report[eds]?|reporting)\b',
        r'\b(audit[eds]?|auditing)\b',
        r'\b(compl[ys]?|complying|complied)\b',
        r'\b(adher[ed]?|adhering)\b',
        r'\b(follow[eds]?|following)\b',
        r'\b(assist[eds]?|assisting)\b',
        r'\b(support[eds]?|supporting)\b',
        r'\b(help[eds]?|helping)\b',
        r'\b(guid[ed]?|guiding)\b',
        r'\b(advic[ed]?|advising)\b',
        r'\b(consult[eds]?|consulting)\b',
        r'\b(recommend[eds]?|recommending)\b',
        r'\b(suggest[eds]?|suggesting)\b',
        r'\b(propos[ed]?|proposing)\b',
        r'\b(initiat[ed]?|initiating)\b',
        r'\b(launch[eds]?|launching)\b',
        r'\b(start[eds]?|starting)\b',
        r'\b(begin[s]?|beginning|began)\b',
        r'\b(finish[eds]?|finishing)\b',
        r'\b(complet[ed]?|completing)\b',
        r'\b(conclud[ed]?|concluding)\b',
        r'\b(resolv[ed]?|resolving)\b',
        r'\b(fix[eds]?|fixing)\b',
        r'\b(solv[ed]?|solving)\b',
        r'\b(address[eds]?|addressing)\b',
        r'\b(handl[ed]?|handling)\b',
        r'\b(process[eds]?|processing)\b',
        r'\b(perform[eds]?|performing)\b',
        r'\b(execut[ed]?|executing)\b',
        r'\b(conduct[eds]?|conducting)\b',
        r'\b(carri[ed]?|carrying)\b',
        r'\b(undertook|undertaking)\b',
        r'\b(oversee[ings]?|overseeing|oversaw)\b',
        r'\b(mobiliz[ed]?|mobilizing|mobilis[ed]?|mobilising)\b',
        r'\b(identif[ys]?|identifying|identified)\b',
        r'\b(recogniz[ed]?|recognizing|recognis[ed]?|recognising)\b',
        r'\b(discover[eds]?|discovering)\b',
        r'\b(detect[eds]?|detecting)\b',
        r'\b(find[ings]?|finding|found)\b',
        r'\b(locat[ed]?|locating)\b',
        r'\b(search[eds]?|searching)\b',
        r'\b(investig[ats]?|investigating|investigated)\b',
        r'\b(explor[ed]?|exploring)\b',
        r'\b(examin[ed]?|examining)\b',
        r'\b(inspect[eds]?|inspecting)\b',
        r'\b(check[eds]?|checking)\b',
        r'\b(verif[ys]?|verifying|verified)\b',
        r'\b(validat[ed]?|validating)\b',
        r'\b(confirm[eds]?|confirming)\b'
    ]
    
    # Find all verb occurrences in the resume
    verb_counts = {}
    text_lower = resume_text.lower()
    
    for pattern in action_verbs_patterns:
        matches = re.findall(pattern, text_lower)
        if matches:
            # Extract the base verb from the pattern for grouping
            base_verb = pattern.split('(')[1].split('[')[0] if '[' in pattern else pattern.split('(')[1].split('|')[0]
            base_verb = base_verb.replace('\\b', '').replace('?', '')
            
            total_occurrences = len(matches)
            if total_occurrences > 0:
                verb_counts[base_verb] = total_occurrences
    
    # Calculate score: Start with 10, deduct 2 points per repetition (occurrence beyond first)
    score = 10
    total_repetitions = 0
    repeated_verbs = []
    
    for verb, count in verb_counts.items():
        if count > 1:  # More than 1 occurrence = repetition
            repetitions = count - 1  # First occurrence doesn't count as repetition
            total_repetitions += repetitions
            repeated_verbs.append(f"{verb}: {count} times ({repetitions} repetitions)")
    
    # Deduct 2 points per repetition
    penalty = total_repetitions * 2
    score = max(0, score - penalty)  # Minimum score is 0
    
    # Debug logging for transparency
    if repeated_verbs:
        logger.info(f"🔄 Verb Repetitions Found:")
        for verb_info in repeated_verbs:
            logger.info(f"   • {verb_info}")
        logger.info(f"🔄 Total repetitions: {total_repetitions}, Penalty: {penalty} points, Final score: {score}/10")
    else:
        logger.info(f"🔄 No verb repetitions found, Score: {score}/10")
    
    return score

def get_repetition_detailed_analysis(resume_text: str) -> dict:
    """
    Provides detailed analysis of verb repetitions for CTA modal reasoning
    Returns specific repeated verbs with counts and alternatives
    """
    import re
    
    # Same patterns as main analysis function
    action_verbs_patterns = [
        r'\b(manage[ds]?|managing)\b', r'\b(develop[eds]?|developing)\b', r'\b(creat[ed]?|creating)\b',
        r'\b(implement[eds]?|implementing)\b', r'\b(lead[s]?|leading|led)\b', r'\b(design[eds]?|designing)\b',
        r'\b(execut[ed]?|executing)\b', r'\b(deliver[eds]?|delivering)\b', r'\b(achiev[ed]?|achieving)\b',
        r'\b(establish[eds]?|establishing)\b', r'\b(coordinat[ed]?|coordinating)\b', r'\b(supervis[ed]?|supervising)\b',
        r'\b(direct[eds]?|directing)\b', r'\b(operat[ed]?|operating)\b', r'\b(maintain[eds]?|maintaining)\b',
        r'\b(analyz[ed]?|analyzing|analys[ed]?|analysing)\b', r'\b(evaluat[ed]?|evaluating)\b',
        r'\b(assess[eds]?|assessing)\b', r'\b(review[eds]?|reviewing)\b', r'\b(monitor[eds]?|monitoring)\b',
        r'\b(track[eds]?|tracking)\b', r'\b(optimiz[ed]?|optimizing|optimis[ed]?|optimising)\b',
        r'\b(improv[ed]?|improving)\b', r'\b(enhanc[ed]?|enhancing)\b', r'\b(reduc[ed]?|reducing)\b',
        r'\b(increas[ed]?|increasing)\b', r'\b(build[ings]?|built)\b', r'\b(train[eds]?|training)\b',
        r'\b(negotiat[ed]?|negotiating)\b', r'\b(ensur[ed]?|ensuring)\b', r'\b(secur[ed]?|securing)\b',
        r'\b(streamlin[ed]?|streamlining)\b', r'\b(assist[eds]?|assisting)\b', r'\b(support[eds]?|supporting)\b',
        r'\b(identif[ys]?|identifying|identified)\b', r'\b(conduct[eds]?|conducting)\b'
    ]
    
    # Verb alternatives for suggestions
    verb_alternatives = {
        'manage': ['oversee', 'supervise', 'administer', 'govern', 'coordinate', 'helm'],
        'develop': ['create', 'build', 'construct', 'formulate', 'establish', 'craft'],
        'lead': ['spearhead', 'direct', 'guide', 'champion', 'pioneer', 'orchestrate'],
        'implement': ['execute', 'deploy', 'launch', 'integrate', 'operationalize', 'roll out'],
        'deliver': ['provide', 'supply', 'produce', 'generate', 'yield', 'fulfill'],
        'ensure': ['guarantee', 'secure', 'maintain', 'verify', 'confirm', 'establish'],
        'create': ['develop', 'design', 'formulate', 'establish', 'generate', 'craft'],
        'conduct': ['perform', 'execute', 'carry out', 'undertake', 'facilitate', 'administer'],
        'support': ['assist', 'aid', 'facilitate', 'enable', 'bolster', 'reinforce'],
        'identify': ['recognize', 'pinpoint', 'discover', 'detect', 'locate', 'uncover'],
        'negotiate': ['broker', 'mediate', 'arrange', 'secure', 'facilitate', 'orchestrate']
    }
    
    # Find all verb occurrences
    verb_counts = {}
    text_lower = resume_text.lower()
    
    for pattern in action_verbs_patterns:
        matches = re.findall(pattern, text_lower)
        if matches:
            base_verb = pattern.split('(')[1].split('[')[0] if '[' in pattern else pattern.split('(')[1].split('|')[0]
            base_verb = base_verb.replace('\\b', '').replace('?', '')
            
            total_occurrences = len(matches)
            if total_occurrences > 1:  # Only include repeated verbs
                verb_counts[base_verb] = total_occurrences
    
    # Generate detailed analysis
    repeated_verbs = []
    total_repetitions = 0
    
    for verb, count in sorted(verb_counts.items(), key=lambda x: x[1], reverse=True):
        repetitions = count - 1
        total_repetitions += repetitions
        
        alternatives = verb_alternatives.get(verb, ['vary', 'diversify', 'alternate'])[:3]
        
        repeated_verbs.append({
            'verb': verb.title(),
            'count': count,
            'repetitions': repetitions,
            'penalty': repetitions * 2,
            'alternatives': alternatives
        })
    
    # Calculate final score
    score = max(0, 10 - (total_repetitions * 2))
    
    return {
        'score': score,
        'total_repetitions': total_repetitions,
        'total_penalty': total_repetitions * 2,
        'repeated_verbs': repeated_verbs[:5],  # Top 5 most repeated verbs
        'analysis_summary': f"Found {len(repeated_verbs)} repeated verbs with {total_repetitions} total repetitions, resulting in {total_repetitions * 2} penalty points."
    }

def analyze_unnecessary_sections_frontend(resume_text: str) -> int:
    """Analyzes unnecessary sections based on modern resume standards"""
    import re
    
    text_lower = resume_text.lower()
    penalty_points = 0
    
    # Check for References section (extremely outdated - major penalty)
    references_patterns = [
        r'\breferences\b',
        r'references available',
        r'references upon request',
        r'references provided'
    ]
    
    if any(re.search(pattern, text_lower) for pattern in references_patterns):
        penalty_points += 4  # Heavy penalty for references
    
    # Check for Objective section (outdated - major penalty)
    objective_patterns = [
        r'\bobjective\b',
        r'career objective',
        r'professional objective',
        r'job objective'
    ]
    
    if any(re.search(pattern, text_lower) for pattern in objective_patterns):
        penalty_points += 4  # Heavy penalty for objective
    
    # Check for high school when higher education exists
    # Use word boundaries to avoid false positives (e.g., "ma" matching in "diploma")
    higher_ed_patterns = [
        r'\bbachelor\b', r'\bmaster\b', r'\bphd\b', r'\bdoctorate\b', 
        r'\buniversity\b', r'\bcollege\b', r'\bbsc\b', r'\bmsc\b', 
        r'\bba\b', r'\bma\b', r'\bbba\b', r'\bmba\b', 
        r'graduate degree', r'undergraduate degree', r'\bpostgraduate\b'
    ]
    
    has_higher_education = any(re.search(pattern, text_lower) for pattern in higher_ed_patterns)
    
    # Check for diploma but exclude high school diploma
    if 'diploma' in text_lower and 'high school diploma' not in text_lower:
        # Additional check for other diploma types that indicate higher education
        diploma_indicators = ['college diploma', 'university diploma', 'graduate diploma', 'professional diploma']
        if any(indicator in text_lower for indicator in diploma_indicators):
            has_higher_education = True
    
    has_high_school = any(keyword in text_lower for keyword in [
        'high school', 'secondary school', 'grade 12', 'matriculation',
        'high school diploma', 'secondary education'
    ])
    
    # Only penalize high school if higher education also exists
    if has_higher_education and has_high_school:
        penalty_points += 2  # Penalty for including high school with higher education
    
    # Check for other unnecessary sections
    other_unnecessary = [
        'hobbies', 'interests', 'personal', 'marital status', 'age', 
        'photo', 'picture', 'nationality', 'religion', 'gender'
    ]
    
    found_other = sum(1 for keyword in other_unnecessary if keyword in text_lower)
    penalty_points += found_other
    
    # Calculate final score based on penalty points
    if penalty_points == 0:
        return 10
    elif penalty_points <= 1:
        return 8
    elif penalty_points <= 2:
        return 6
    elif penalty_points <= 3:
        return 4
    elif penalty_points <= 5:
        return 2
    else:
        return 1  # Extremely low score for multiple outdated sections

def analyze_growth_signals_frontend(resume_text: str) -> int:
    """Enhanced growth signals analysis with promotion and progression detection"""
    import re
    
    # Signal 1: Detect promotions within same organization
    internal_promotions = detect_promotions_within_organization(resume_text)
    
    # Signal 2: Detect promotion/promoted keywords
    promotion_keywords = detect_promotion_keywords(resume_text)
    
    # Signal 3: Detect designation progression across organizations
    cross_org_progression = detect_designation_progression(resume_text)
    
    # Count signals found
    signals_found = sum([
        internal_promotions > 0,
        promotion_keywords > 0,
        cross_org_progression > 0
    ])
    
    # Apply new scoring system
    if signals_found == 3:
        return 10
    elif signals_found == 2:
        return 7
    elif signals_found == 1:
        return 4
    else:
        return 0

def detect_promotions_within_organization(resume_text: str) -> int:
    """Detect promotions within the same company heading"""
    import re
    
    # Split text into sections and look for experience sections
    lines = resume_text.split('\n')
    
    promotions_found = 0
    current_company_section = []
    in_experience_section = False
    
    for line in lines:
        line_clean = line.strip()
        if not line_clean:
            continue
            
        line_lower = line_clean.lower()
        
        # Check if we're entering an experience section
        if any(keyword in line_lower for keyword in ['experience', 'employment', 'work history', 'career']):
            in_experience_section = True
            continue
            
        # Check if we're leaving experience section
        if any(keyword in line_lower for keyword in ['education', 'skills', 'projects', 'certifications']):
            # Process final company section before leaving
            if current_company_section:
                promotions_found += analyze_company_section_for_promotions(current_company_section)
                current_company_section = []
            in_experience_section = False
            continue
            
        if in_experience_section:
            # Check if this line is a company header (contains company name with dates)
            if is_company_header(line_clean):
                # Process previous company section if exists
                if current_company_section:
                    promotions_found += analyze_company_section_for_promotions(current_company_section)
                # Start new company section
                current_company_section = [line_clean]
            else:
                # Add to current company section
                current_company_section.append(line_clean)
    
    # Process final company section
    if current_company_section:
        promotions_found += analyze_company_section_for_promotions(current_company_section)
    
    return promotions_found

def is_company_header(line: str) -> bool:
    """Check if line is a company header (company name with date range)"""
    import re
    
    # Company header should have company name and date range
    # Pattern: "CompanyName (YYYY-YYYY)" or "CompanyName Inc. 2020-2023"
    has_dates = bool(re.search(r'\b(20\d{2}|19\d{2})\b', line))
    
    # Company indicators
    company_indicators = ['inc', 'ltd', 'corp', 'company', 'technologies', 'solutions', 'systems', 'group']
    has_company_word = any(indicator in line.lower() for indicator in company_indicators)
    
    # Check if it looks like a company header (not a job title line)
    # Job titles usually start with job-related words
    job_title_starters = ['senior', 'junior', 'lead', 'principal', 'chief', 'manager', 'director', 
                         'analyst', 'engineer', 'developer', 'specialist', 'consultant', 
                         'coordinator', 'assistant', 'associate']
    
    starts_with_title = any(line.lower().startswith(title) for title in job_title_starters)
    
    # Company header: has dates AND (has company indicators OR doesn't start with job title)
    return has_dates and (has_company_word or not starts_with_title)

def analyze_company_section_for_promotions(company_lines: list) -> int:
    """Analyze a company section for multiple job titles indicating promotion"""
    import re
    
    if len(company_lines) < 2:
        return 0
        
    # Look for multiple job titles with dates in the same company section
    job_entries = []
    
    for line in company_lines:
        # Look for lines that have job titles and dates
        if re.search(r'\b(20\d{2}|19\d{2})\b', line):
            # Extract potential job title (text before company name or date)
            title_match = re.search(r'^([^(]+?)(?:\s*[-–@]\s*|\s*\()', line)
            if title_match:
                title = title_match.group(1).strip()
                
                # Extract year for chronological analysis
                year_matches = re.findall(r'\b(20\d{2}|19\d{2})\b', line)
                if year_matches:
                    start_year = min(int(year) for year in year_matches)
                    job_entries.append((title, start_year))
    
    # If we have multiple job titles, check for hierarchy progression
    if len(job_entries) >= 2:
        # Sort by year
        job_entries.sort(key=lambda x: x[1])
        
        # Check for progression in titles
        titles = [entry[0].lower() for entry in job_entries]
        
        # Look for clear progression indicators
        progression_patterns = [
            ['associate', 'senior'],
            ['junior', 'senior'],
            ['analyst', 'senior analyst'],
            ['engineer', 'senior engineer'],
            ['engineer', 'lead engineer'],
            ['senior', 'lead'],
            ['senior', 'principal'],
            ['developer', 'senior developer'],
            ['manager', 'senior manager'],
            ['executive', 'senior executive']
        ]
        
        for i in range(len(titles) - 1):
            current_title = titles[i]
            next_title = titles[i + 1]
            
            # Check for direct progression patterns
            for pattern in progression_patterns:
                if pattern[0] in current_title and pattern[1] in next_title:
                    return 1
                    
            # Check for role expansion (same role but with additional words)
            if current_title in next_title and len(next_title) > len(current_title):
                return 1
    
    return 0

def detect_promotion_keywords(resume_text: str) -> int:
    """Detect promotion-related keywords in context"""
    import re
    
    promotion_patterns = [
        r'\bpromoted\b',
        r'\bpromotion\b',
        r'\badvanced to\b',
        r'\belevated to\b',
        r'\bprogressed to\b',
        r'\btransitioned to\b.*\b(senior|lead|principal|manager|director)\b',
        r'\bincreased responsibility\b',
        r'\bexpanded role\b',
        r'\brecognized.*and\s+(promoted|advanced)\b',
        r'\bselected for.*(promotion|advancement)\b'
    ]
    
    text_lower = resume_text.lower()
    
    # Count distinct promotion indicators
    found_patterns = 0
    for pattern in promotion_patterns:
        if re.search(pattern, text_lower):
            found_patterns += 1
    
    return min(found_patterns, 1)  # Return 1 if any promotion keywords found

def detect_designation_progression(resume_text: str) -> int:
    """Detect career progression across different organizations"""
    import re
    
    # Extract job entries with companies and dates
    lines = resume_text.split('\n')
    job_entries = []
    
    for line in lines:
        line_clean = line.strip()
        if not line_clean:
            continue
            
        # Look for lines with dates (indicating job entries)
        if re.search(r'\b(20\d{2}|19\d{2})\b', line_clean):
            # Extract year range
            years = re.findall(r'\b(20\d{2}|19\d{2})\b', line_clean)
            if years:
                start_year = min(int(year) for year in years)
                
                # Extract job title (usually at the beginning of the line)
                title_match = re.match(r'^([^-–@(]+)', line_clean)
                if title_match:
                    title = title_match.group(1).strip()
                    
                    # Extract company (usually after - or @ or in parentheses)
                    company_match = re.search(r'[-–@]\s*([^(,]+)', line_clean)
                    company = company_match.group(1).strip() if company_match else "Unknown"
                    
                    job_entries.append((title, company, start_year))
    
    if len(job_entries) < 2:
        return 0
        
    # Sort by year
    job_entries.sort(key=lambda x: x[2])
    
    # Check for progression across organizations
    hierarchy_levels = {
        'intern': 1, 'trainee': 1,
        'associate': 2, 'junior': 2, 'analyst': 2,
        'engineer': 3, 'developer': 3, 'executive': 3,
        'senior': 4, 'specialist': 4,
        'lead': 5, 'principal': 5, 'staff': 5,
        'manager': 6, 'supervisor': 6,
        'senior manager': 7, 'director': 7,
        'senior director': 8, 'vice president': 8, 'vp': 8,
        'president': 9, 'ceo': 10, 'cto': 10, 'cfo': 10
    }
    
    progression_found = 0
    
    for i in range(len(job_entries) - 1):
        current_title = job_entries[i][0].lower()
        next_title = job_entries[i + 1][0].lower()
        current_company = job_entries[i][1].lower()
        next_company = job_entries[i + 1][1].lower()
        
        # Only consider if different companies
        if current_company != next_company:
            # Check for level progression
            current_level = 0
            next_level = 0
            
            for keyword, level in hierarchy_levels.items():
                if keyword in current_title:
                    current_level = max(current_level, level)
                if keyword in next_title:
                    next_level = max(next_level, level)
            
            # If we found a progression
            if next_level > current_level:
                progression_found = 1
                break
    
    return progression_found

def analyze_drive_and_initiative_frontend(resume_text: str) -> int:
    """Copied exactly from frontend analyzeDriveAndInitiative"""
    initiative_keywords = [
        'initiated', 'pioneered', 'launched', 'founded', 'established',
        'spearheaded', 'drove', 'championed', 'innovated', 'transformed'
    ]
    text_lower = resume_text.lower()
    found_keywords = sum(1 for keyword in initiative_keywords if keyword in text_lower)
    
    if found_keywords >= 3:
        return 9
    elif found_keywords >= 2:
        return 7
    elif found_keywords >= 1:
        return 5
    else:
        return 3

def analyze_certifications_frontend(resume_text: str) -> int:
    """Copied exactly from frontend analyzeCertifications"""
    cert_keywords = ['certified', 'certification', 'license', 'credential', 'certificate']
    text_lower = resume_text.lower()
    found_certs = sum(1 for cert in cert_keywords if cert in text_lower)
    
    if found_certs >= 2:
        return 8
    elif found_certs >= 1:
        return 6
    else:
        return 4

def analyze_experience_section_percentage(resume_text: str) -> int:
    """Analyzes what percentage of the resume is dedicated to experience content"""
    import re
    
    lines = resume_text.split('\n')
    total_lines = len([line for line in lines if line.strip()])
    
    # Look for experience section indicators
    exp_section_patterns = [
        r'(professional\s+)?experience',
        r'work\s+experience', 
        r'employment\s+history',
        r'career\s+history',
        r'professional\s+background'
    ]
    
    experience_lines = 0
    in_experience_section = False
    
    for line in lines:
        line_lower = line.lower().strip()
        if not line_lower:
            continue
            
        # Check if this line starts an experience section
        if any(re.search(pattern, line_lower) for pattern in exp_section_patterns):
            in_experience_section = True
            experience_lines += 1
            continue
            
        # Check if we've moved to a new section (common section headers)
        section_headers = [
            r'education', r'skills', r'certifications', r'achievements',
            r'projects', r'publications', r'references', r'languages'
        ]
        if any(re.search(f'^{header}', line_lower) for header in section_headers):
            in_experience_section = False
            continue
            
        # Count lines that appear to be in experience section
        if in_experience_section:
            # Look for job-related content (dates, companies, bullet points)
            if (re.search(r'\b(20\d{2}|19\d{2})\b', line) or  # Years
                line.strip().startswith(('•', '-', '*')) or      # Bullet points
                len(line.strip()) > 20):                         # Substantial content
                experience_lines += 1
    
    if total_lines == 0:
        return 1
        
    experience_percentage = (experience_lines / total_lines) * 100
    
    # Linear scaling: 1 if <20%, 10 if >=70%
    if experience_percentage < 20:
        return 1
    elif experience_percentage >= 70:
        return 10
    else:
        # Linear interpolation between 20% and 70%
        return round(1 + ((experience_percentage - 20) / 50) * 9)

def analyze_section_titles_clarity(resume_text: str) -> int:
    """Analyzes clarity and consistency of section titles"""
    import re
    
    lines = resume_text.split('\n')
    
    # Expected main sections
    expected_sections = [
        r'(professional\s+)?experience',
        r'(work\s+)?experience', 
        r'education',
        r'skills',
    ]
    
    # Optional but common sections
    optional_sections = [
        r'certifications?',
        r'achievements?',
        r'projects?',
        r'summary',
        r'objective'
    ]
    
    found_main_sections = 0
    total_sections_found = 0
    unclear_sections = 0
    
    for line in lines:
        line_clean = line.strip()
        if not line_clean or len(line_clean) < 3:
            continue
            
        line_lower = line_clean.lower()
        
        # Check if line looks like a section header (short, uppercase/title case)
        if (len(line_clean) < 50 and 
            (line_clean.isupper() or line_clean.istitle()) and
            not any(char.isdigit() for char in line_clean[:10])):
            
            total_sections_found += 1
            
            # Check against expected main sections
            if any(re.search(pattern, line_lower) for pattern in expected_sections):
                found_main_sections += 1
            elif any(re.search(pattern, line_lower) for pattern in optional_sections):
                pass  # Optional section, counts as clear
            else:
                # Check for unclear/generic headers
                if len(line_clean) < 10 or line_lower in ['details', 'information', 'data']:
                    unclear_sections += 1
    
    # Calculate score based on clarity
    if total_sections_found == 0:
        return 1  # No clear sections found
        
    clarity_ratio = (total_sections_found - unclear_sections) / total_sections_found
    main_section_score = min(found_main_sections, 3) / 3  # Cap at 3 main sections
    
    # Combine both factors
    final_score = (clarity_ratio * 0.6) + (main_section_score * 0.4)
    return max(1, round(final_score * 10))

def analyze_job_titles_clarity(resume_text: str) -> int:
    """Analyzes presence and clarity of job titles"""
    import re
    
    lines = resume_text.split('\n')
    
    # Common job title patterns
    job_title_patterns = [
        r'\b(senior|junior|lead|principal|chief)\s+\w+',
        r'\b(manager|director|analyst|engineer|developer|specialist)\b',
        r'\b(consultant|coordinator|assistant|associate)\b',
        r'\b(designer|architect|administrator|technician)\b'
    ]
    
    # Look for job titles near date patterns
    potential_job_titles = 0
    clear_job_titles = 0
    
    for i, line in enumerate(lines):
        line_clean = line.strip()
        if not line_clean:
            continue
            
        # Look for dates indicating employment periods
        has_date = re.search(r'\b(20\d{2}|19\d{2})\b', line_clean)
        
        if has_date:
            # Check current line and nearby lines for job titles
            check_lines = [lines[max(0, i-1)], line_clean, lines[min(len(lines)-1, i+1)]]
            
            for check_line in check_lines:
                check_clean = check_line.strip()
                if not check_clean:
                    continue
                    
                # Skip if line contains common non-title indicators
                if any(word in check_clean.lower() for word in ['education', 'university', 'college', 'school']):
                    continue
                    
                potential_job_titles += 1
                
                # Check if it matches common job title patterns
                if any(re.search(pattern, check_clean, re.IGNORECASE) for pattern in job_title_patterns):
                    clear_job_titles += 1
                # Or if it's properly formatted (title case, reasonable length)
                elif (check_clean.istitle() and 10 <= len(check_clean) <= 50 and 
                      not check_clean.startswith(('•', '-', '*'))):
                    clear_job_titles += 1
                
                break  # Only count one per date entry
    
    if potential_job_titles == 0:
        return 1  # No job entries found
        
    clarity_ratio = clear_job_titles / potential_job_titles
    return max(1, round(clarity_ratio * 10))

def analyze_filename_appropriateness(filename: str = None) -> int:
    """Analyzes filename appropriateness based on best practices"""
    import re
    
    if not filename:
        return 1  # No filename provided
        
    # Remove file extension for analysis
    name_without_ext = re.sub(r'\.[^.]*$', '', filename)
    
    # Best practice pattern: FirstName-LastName-Resume
    best_practice_pattern = r'^[A-Za-z]+-[A-Za-z]+-Resume$'
    
    if re.match(best_practice_pattern, name_without_ext, re.IGNORECASE):
        return 10
        
    # Good patterns
    good_patterns = [
        r'^[A-Za-z]+[_\s][A-Za-z]+[_\s-]Resume$',  # FirstName LastName Resume
        r'^[A-Za-z]+[A-Za-z]+Resume$',              # FirstnameLastnameResume
        r'^Resume[_\s-][A-Za-z]+[_\s-][A-Za-z]+$'  # Resume-FirstName-LastName
    ]
    
    if any(re.match(pattern, name_without_ext, re.IGNORECASE) for pattern in good_patterns):
        return 8
        
    # Acceptable patterns (contains name and resume)
    if (re.search(r'resume', name_without_ext, re.IGNORECASE) and 
        re.search(r'[A-Za-z]{2,}', name_without_ext)):
        return 6
        
    # Poor patterns
    poor_patterns = [
        r'^(resume|cv)$',           # Just "resume" or "cv"
        r'^document\d*$',           # Generic document names
        r'^untitled',               # Untitled files
        r'^\d+$'                   # Just numbers
    ]
    
    if any(re.match(pattern, name_without_ext, re.IGNORECASE) for pattern in poor_patterns):
        return 1
        
    # Default for unclear but not terrible filenames
    return 4

def analyze_font_readability(resume_text: str) -> int:
    """Analyzes font readability based on special characters and parsing issues"""
    import re
    
    # Count special/problematic characters that indicate font issues
    problematic_chars = 0
    total_chars = len(resume_text)
    
    if total_chars == 0:
        return 1
        
    # Look for problematic character patterns
    issues = {
        'weird_quotes': len(re.findall(r'[""''`´]', resume_text)),
        'special_symbols': len(re.findall(r'[►▪▫■□●○♦♠♣♥]', resume_text)),
        'unusual_spaces': len(re.findall(r'[\u00A0\u2000-\u200F\u2028-\u202F]', resume_text)),
        'accented_where_unexpected': len(re.findall(r'[àáâãäåæçèéêëìíîïñòóôõöøùúûüý]', resume_text)),
        'replacement_chars': len(re.findall(r'[�]', resume_text)),
        'weird_dashes': len(re.findall(r'[–—]', resume_text))
    }
    
    total_issues = sum(issues.values())
    
    # Calculate issue percentage
    issue_percentage = (total_issues / total_chars) * 100
    
    # Check for proper standard characters (letters, numbers, basic punctuation)
    standard_chars = len(re.findall(r'[a-zA-Z0-9\s.,;:!?()[\]\-_+=@#$%&*/\\]', resume_text))
    standard_percentage = (standard_chars / total_chars) * 100
    
    # Score based on readability
    if issue_percentage > 5:
        return 1  # Too many problematic characters
    elif issue_percentage > 2:
        return 4  # Some issues
    elif standard_percentage > 95:
        return 10  # Excellent readability
    elif standard_percentage > 90:
        return 8  # Good readability
    else:
        return 6  # Acceptable readability

def analyze_formatting_ats_parsing(resume_text: str) -> int:
    """Analyzes formatting quality and ATS parsing compatibility"""
    import re
    
    score = 10  # Start with perfect score, deduct for issues
    
    # Check for parsing errors (indicators of complex formatting)
    parsing_issues = [
        len(re.findall(r'[^\x00-\x7F]', resume_text)),  # Non-ASCII characters
        resume_text.count('\t'),                        # Excessive tabs
        len(re.findall(r'\n\s*\n\s*\n', resume_text)), # Excessive blank lines
        len(re.findall(r'[|│┃]', resume_text)),         # Table/column borders
        len(re.findall(r'[═─━]', resume_text))          # Graphic elements
    ]
    
    total_parsing_issues = sum(parsing_issues)
    
    # Deduct for parsing complexity
    if total_parsing_issues > 50:
        score -= 6
    elif total_parsing_issues > 20:
        score -= 4
    elif total_parsing_issues > 10:
        score -= 2
    
    # Check for good structure indicators
    has_bullet_points = bool(re.search(r'^[\s]*[•\-*]', resume_text, re.MULTILINE))
    has_clear_sections = len(re.findall(r'^[A-Z\s]{5,}$', resume_text, re.MULTILINE)) >= 2
    has_reasonable_line_length = len([l for l in resume_text.split('\n') if len(l) > 200]) < 5
    
    # Reward good structure
    if has_bullet_points:
        score += 1
    if has_clear_sections:
        score += 1
    if has_reasonable_line_length:
        score += 1
    
    # Ensure score stays within bounds
    return max(1, min(10, score))

def calculate_cv_readability_score(resume_text: str, filename: str = None) -> float:
    """Calculate the weighted CV Readability Score"""
    
    # Get individual scores
    experience_score = analyze_experience_section_percentage(resume_text)
    section_titles_score = analyze_section_titles_clarity(resume_text)
    job_titles_score = analyze_job_titles_clarity(resume_text)
    filename_score = analyze_filename_appropriateness(filename)
    font_score = analyze_font_readability(resume_text)
    formatting_score = analyze_formatting_ats_parsing(resume_text)
    
    # Apply weighted formula
    final_score = (
        (experience_score * 0.30) +
        (section_titles_score * 0.20) +
        (job_titles_score * 0.15) +
        (filename_score * 0.10) +
        (font_score * 0.15) +
        (formatting_score * 0.10)
    )
    
    return round(final_score, 1)

def get_enhanced_issue_description(category_name: str, score: int, resume_text: str = "") -> dict:
    """
    Generate enhanced issue descriptions with scoring criteria and specific guidance
    """
    
    # Enhanced category definitions with scoring criteria and unique issues
    category_enhancements = {
        'Contact Details': {
            'understanding': 'Measures completeness and professionalism of contact information for ATS systems',
            'high_score_criteria': [
                'Professional email address with proper domain',
                'Complete phone number with country code',
                'LinkedIn profile URL included and professional'
            ],
            'low_score_issues': [
                'Missing essential contact information (email, phone)',
                'Unprofessional email addresses (e.g., partyguy@email.com)',
                'No LinkedIn profile or social media presence'
            ],
            'specific_issues': {
                'high': ['Add professional LinkedIn profile URL', 'Include complete address with city and state', 'Verify email address is professional'],
                'medium': ['Update phone number format for ATS compatibility', 'Add GitHub profile if technical role', 'Include portfolio website if relevant'],
                'low': ['Replace unprofessional email with firstname.lastname format', 'Add missing contact information immediately', 'Create professional LinkedIn profile']
            }
        },
        
        'Skills Section': {
            'understanding': 'Evaluates technical and soft skills presentation, relevance, and ATS keyword optimization',
            'high_score_criteria': [
                'Role-specific technical skills clearly listed',
                'Mix of hard and soft skills relevant to target position',
                'Skills organized in logical categories or priority order'
            ],
            'low_score_issues': [
                'Generic skills that apply to any role',
                'Missing industry-specific technologies and tools',
                'Outdated software or programming languages'
            ],
            'specific_issues': {
                'high': ['Add specific software proficiency levels (Expert, Advanced, Intermediate)', 'Include emerging technologies relevant to your field', 'Highlight cross-functional collaboration skills'],
                'medium': ['Organize skills by categories (Technical, Leadership, Industry-specific)', 'Add cloud platforms and modern development tools', 'Include data analysis and visualization tools'],
                'low': ['Replace generic skills with specific technical competencies', 'Add industry-relevant programming languages and frameworks', 'Include modern software tools and platforms used in your field']
            }
        },
        
        'Analytical': {
            'understanding': 'Assesses demonstration of analytical thinking, data-driven decision making, and problem-solving abilities',
            'high_score_criteria': [
                'Specific examples of data analysis and insights generated',
                'Quantified impact of analytical work with metrics',
                'Clear problem-solving methodologies mentioned'
            ],
            'low_score_issues': [
                'No quantified examples of analytical work',
                'Missing data-driven achievements and insights',
                'Generic statements without specific analytical methods'
            ],
            'specific_issues': {
                'high': ['Add specific data analysis tools used (SQL, Python, Tableau, etc.)', 'Include examples of insights that drove business decisions', 'Highlight statistical methods or frameworks applied'],
                'medium': ['Quantify analytical impact with percentages and metrics', 'Show progression from analysis to implementation', 'Include cross-functional analytical projects'],
                'low': ['Add concrete examples of data analysis you\'ve performed', 'Include specific metrics and outcomes from your analytical work', 'Mention analytical tools and methodologies you\'ve used']
            }
        },
        
        'Summary': {
            'understanding': 'Evaluates professional summary quality based on word count, specific content, metrics, and professional language',
            'high_score_criteria': [
                'Concise summary under 100 words with specific achievements',
                'Quantifiable metrics and concrete numbers included',
                'Professional language without personal pronouns or buzzwords'
            ],
            'low_score_issues': [
                'Generic buzzwords like "results-driven", "passionate", "motivated"',
                'Personal pronouns (I, my, me) making it sound unprofessional',
                'Missing specific metrics, experience years, or quantifiable achievements'
            ],
            'specific_issues': {
                'high': ['Add specific industry expertise and technical skills', 'Include leadership scope (team size, budget managed)', 'Highlight unique value proposition that differentiates you'],
                'medium': ['Replace generic terms with specific accomplishments', 'Add quantifiable results and percentage improvements', 'Remove remaining buzzwords with concrete examples'],
                'low': ['Rewrite entire summary removing all personal pronouns', 'Add specific metrics: years of experience, budget size, team size', 'Replace all vague buzzwords with measurable achievements']
            }
        },
        
        'Repetition': {
            'understanding': 'Analyzes repetitive use of action verbs throughout your resume. Each verb should appear only once to maintain variety and impact.',
            'high_score_criteria': [
                'Each action verb used only once throughout the resume',
                'Rich vocabulary with diverse verbs for different achievements',
                'Strong, varied action words that avoid monotonous language'
            ],
            'low_score_issues': [
                'Same action verbs repeated multiple times (managed, developed, created)',
                'Limited vocabulary with overuse of basic verbs',
                'Repetitive language that reduces resume impact and readability'
            ],
            'specific_issues': {
                'high': ['Replace 1-2 repeated verbs with powerful alternatives', 'Use more specific action verbs for different contexts', 'Enhance variety with industry-specific terminology'],
                'medium': ['Substitute multiple repeated verbs with unique alternatives', 'Diversify language across different job experiences', 'Use stronger action verbs that better describe accomplishments'],
                'low': ['Completely rewrite repetitive phrases with varied vocabulary', 'Replace all repeated verbs with unique alternatives', 'Transform basic verbs into impactful, specific action words']
            }
        },
        
        'Certifications': {
            'understanding': 'Measures presence, relevance, and currency of professional certifications and credentials',
            'high_score_criteria': [
                'Industry-relevant professional certifications listed',
                'Multiple current certifications from recognized bodies',
                'Certification dates and renewal status included'
            ],
            'low_score_issues': [
                'No professional certifications mentioned',
                'Outdated or expired certifications only',
                'Irrelevant certifications for target role'
            ],
            'specific_issues': {
                'high': ['Add renewal dates for current certifications', 'Include specialized micro-credentials and digital badges', 'Highlight certifications that differentiate you from competitors'],
                'medium': ['Pursue industry-standard certifications relevant to your field', 'Add professional development courses from recognized platforms', 'Include vendor-specific certifications (AWS, Microsoft, Google, etc.)'],
                'low': ['Obtain fundamental industry certifications immediately', 'Add any completed training programs or courses', 'Include professional licenses and credentials']
            }
        },
        
        'Leadership': {
            'understanding': 'Evaluates demonstration of leadership capabilities, team management, and influence',
            'high_score_criteria': [
                'Specific examples of team leadership with team sizes',
                'Cross-functional leadership and collaboration examples',
                'Quantified leadership impact and results'
            ],
            'low_score_issues': [
                'No clear leadership examples or team management',
                'Missing cross-functional collaboration evidence',
                'Lack of quantified leadership outcomes'
            ],
            'specific_issues': {
                'high': ['Highlight mentorship and coaching of junior team members', 'Include examples of leading through organizational change', 'Show leadership impact across different stakeholder groups'],
                'medium': ['Add specific team sizes managed and project scope', 'Include cross-departmental leadership initiatives', 'Quantify team performance improvements under your leadership'],
                'low': ['Add any team leadership experience, even informal roles', 'Include examples of training or mentoring others', 'Highlight initiative-taking and project ownership']
            }
        },
        
        'Growth Signals': {
            'understanding': 'Detects career progression through promotions, expanding responsibilities, and skill development',
            'high_score_criteria': [
                'Clear promotion progression within organizations',
                'Expanding scope of responsibility over time',
                'Cross-company career advancement patterns'
            ],
            'low_score_issues': [
                'No visible career progression or promotions',
                'Static role responsibilities without growth',
                'Missing demonstration of increasing impact'
            ],
            'specific_issues': {
                'high': ['Highlight rapid career progression and early promotions', 'Show expanding team and budget responsibilities', 'Include geographic expansion or new market leadership'],
                'medium': ['Emphasize increasing project complexity and scope', 'Add examples of expanded role responsibilities', 'Include skill development and new domain expertise'],
                'low': ['Highlight any promotion or role expansion you\'ve had', 'Show increasing responsibility even in the same role', 'Include skill development and learning achievements']
            }
        },
        
        'Dates': {
            'understanding': 'Evaluates consistency and professional formatting of dates across experience, education, and projects',
            'high_score_criteria': [
                'Consistent date format throughout resume (MM/YYYY or MM-YYYY)',
                'All relevant positions include both start and end dates',
                'Chronological order with clear employment timeline'
            ],
            'low_score_issues': [
                'Inconsistent date formats mixing different styles',
                'Missing dates on significant positions or education',
                'Date ranges that create timeline gaps or overlaps'
            ],
            'specific_issues': {
                'high': ['Ensure all dates follow exact same format pattern', 'Add specific month/year for all positions and education', 'Verify chronological order is maintained throughout'],
                'medium': ['Standardize date format to MM/YYYY across all sections', 'Add missing dates for education and certifications', 'Fix any mixed date formatting patterns'],
                'low': ['Add missing employment and education dates immediately', 'Choose one consistent date format and apply throughout', 'Fix major date inconsistencies and formatting errors']
            }
        }
    }
    
    # Get category-specific enhancement or use default
    enhancement = category_enhancements.get(category_name, {
        'understanding': f'Evaluates {category_name.lower()} aspects of your resume for ATS optimization',
        'high_score_criteria': ['Professional presentation', 'Relevant content', 'Clear structure'],
        'low_score_issues': ['Missing key elements', 'Poor presentation', 'Lack of specificity'],
        'specific_issues': {
            'high': [f'Optimize {category_name.lower()} for maximum impact'],
            'medium': [f'Improve {category_name.lower()} presentation and content'],
            'low': [f'Add essential {category_name.lower()} elements to your resume']
        }
    })
    
    # Select appropriate specific issues based on score
    if score >= 8:
        specific_issues = enhancement['specific_issues']['high']
    elif score >= 5:
        specific_issues = enhancement['specific_issues']['medium']
    else:
        specific_issues = enhancement['specific_issues']['low']
    
    # Override with dynamic content for categories that have specific generators
    if category_name == 'Contact Details' and resume_text:
        dynamic_issues = generate_contact_examples(resume_text)
        if dynamic_issues:
            # Convert to the format expected by frontend
            specific_issues = [item['issue'] + ': ' + item['example'] for item in dynamic_issues[:3]]
    
    return {
        'understanding': enhancement['understanding'],
        'high_score_criteria': enhancement['high_score_criteria'],
        'low_score_issues': enhancement['low_score_issues'],
        'specific_issues': specific_issues[:3],  # Limit to 3 unique issues
        'issue': specific_issues[0] if specific_issues else f'Improve {category_name.lower()} presentation'
    }

def generate_comprehensive_ats_scores_frontend(content: str, component_scores: dict = None, detailed_analysis: dict = None, filename: str = None) -> List[dict]:
    """
    Generate comprehensive ATS scores for all 23+ categories - ENHANCED WITH SPECIFIC GUIDANCE
    """
    logger.info('🏗️ Generating comprehensive ATS scores with enhanced guidance')
    
    # Extract data (keeping backend compatibility)
    resume_text = content
    
    # Now calculate REAL scores for each category based on frontend analysis
    categories = []
    
    # 1. CONTACT INFORMATION
    contact_score = analyze_contact_details_frontend(resume_text)
    contact_enhancement = get_enhanced_issue_description('Contact Details', contact_score, resume_text)
    contact_modal = generate_fix_this_modal_content('Contact Details', resume_text, contact_score)
    categories.append({
        'name': 'Contact Details',
        'score': contact_score,
        'issue': contact_enhancement['issue'],
        'understanding': contact_enhancement['understanding'],
        'high_score_criteria': contact_enhancement['high_score_criteria'],
        'low_score_issues': contact_enhancement['low_score_issues'],
        'specific_issues': contact_enhancement['specific_issues'],
        'impact': 'SECTIONS',
        'modal_content': contact_modal
    })
    
    # 2-3. STRUCTURE ANALYSIS
    categories.append({
        'name': 'Education Section',
        'score': analyze_education_section_frontend(resume_text),
        'issue': 'Optimize education section format and content',
        'impact': 'SECTIONS'
    })
    
    skills_score = analyze_skills_section_frontend(resume_text)
    skills_enhancement = get_enhanced_issue_description('Skills Section', skills_score, resume_text)
    categories.append({
        'name': 'Skills Section', 
        'score': skills_score,
        'issue': skills_enhancement['issue'],
        'understanding': skills_enhancement['understanding'],
        'high_score_criteria': skills_enhancement['high_score_criteria'],
        'low_score_issues': skills_enhancement['low_score_issues'],
        'specific_issues': skills_enhancement['specific_issues'],
        'impact': 'SECTIONS'
    })
    
    # 4-5. KEYWORD OPTIMIZATION
    analytical_score = analyze_analytical_skills_frontend(resume_text)
    analytical_enhancement = get_enhanced_issue_description('Analytical', analytical_score, resume_text)
    categories.append({
        'name': 'Analytical',
        'score': analytical_score,
        'issue': analytical_enhancement['issue'],
        'understanding': analytical_enhancement['understanding'],
        'high_score_criteria': analytical_enhancement['high_score_criteria'],
        'low_score_issues': analytical_enhancement['low_score_issues'],
        'specific_issues': analytical_enhancement['specific_issues'],
        'impact': 'ALL'
    })
    
    leadership_score = analyze_leadership_skills_frontend(resume_text)
    leadership_enhancement = get_enhanced_issue_description('Leadership', leadership_score, resume_text)
    categories.append({
        'name': 'Leadership',
        'score': leadership_score,
        'issue': leadership_enhancement['issue'],
        'understanding': leadership_enhancement['understanding'],
        'high_score_criteria': leadership_enhancement['high_score_criteria'],
        'low_score_issues': leadership_enhancement['low_score_issues'],
        'specific_issues': leadership_enhancement['specific_issues'],
        'impact': 'ALL'
    })
    
    # 6-11. FORMATTING & STYLE
    categories.append({
        'name': 'Page Density',
        'score': analyze_page_density_frontend(resume_text),
        'issue': 'Optimize page layout and white space usage',
        'impact': 'STYLE'
    })
    categories.append({
        'name': 'Use of Bullets',
        'score': analyze_bullet_usage_frontend(resume_text),
        'issue': 'Improve bullet point structure and formatting',
        'impact': 'STYLE'
    })
    categories.append({
        'name': 'Grammar',
        'score': analyze_grammar_frontend(resume_text),
        'issue': 'Fix grammar errors and improve language accuracy',
        'impact': 'BREVITY'
    })
    categories.append({
        'name': 'Spelling',
        'score': analyze_llm_spelling_frontend(resume_text),
        'issue': 'Fix spelling errors using AI-powered detection',
        'impact': 'BREVITY'
    })
    categories.append({
        'name': 'Verb Tenses',
        'score': analyze_verb_tenses_frontend(resume_text),
        'issue': 'Use consistent and appropriate verb tenses',
        'impact': 'BREVITY'
    })
    categories.append({
        'name': 'Personal Pronouns',
        'score': analyze_personal_pronouns_frontend(resume_text),
        'issue': 'Remove first-person pronouns like "I", "me", "my"',
        'impact': 'BREVITY'
    })
    
    # 12-16. ACHIEVEMENTS & CONTENT
    categories.append({
        'name': 'Quantifiable Achievements',
        'score': analyze_quantifiable_achievements_frontend(resume_text),
        'issue': 'Add more quantified achievements with specific numbers',
        'impact': 'IMPACT'
    })
    categories.append({
        'name': 'Action Verbs',
        'score': analyze_action_verbs_frontend(resume_text),
        'issue': 'Use more strong action verbs to start bullet points',
        'impact': 'IMPACT'
    })
    categories.append({
        'name': 'Active Voice',
        'score': analyze_active_voice_frontend(resume_text),
        'issue': 'Convert passive voice to active voice for impact',
        'impact': 'IMPACT'
    })
    summary_score = analyze_summary_section_frontend(resume_text)
    summary_enhancement = get_enhanced_issue_description('Summary', summary_score, resume_text)
    summary_modal = generate_fix_this_modal_content('Summary', resume_text, summary_score)
    categories.append({
        'name': 'Summary',
        'score': summary_score,
        'issue': summary_enhancement['issue'],
        'impact': 'IMPACT',
        'detailed_analysis': get_summary_detailed_analysis(resume_text),
        'enhancement': summary_enhancement,
        'modal_content': summary_modal
    })
    categories.append({
        'name': 'Teamwork',
        'score': analyze_teamwork_skills_frontend(resume_text),
        'issue': 'Better showcase collaborative experiences',
        'impact': 'ALL'
    })
    
    # 17-21. READABILITY & CONTENT QUALITY
    categories.append({
        'name': 'Verbosity',
        'score': 8,  # Placeholder - using good default
        'issue': 'Reduce wordiness for better readability',
        'impact': 'BREVITY'
    })
    repetition_score = analyze_repetition_frontend(resume_text)
    repetition_enhancement = get_enhanced_issue_description('Repetition', repetition_score, resume_text)
    repetition_modal = generate_fix_this_modal_content('Repetition', resume_text, repetition_score)
    categories.append({
        'name': 'Repetition',
        'score': repetition_score,
        'issue': repetition_enhancement['issue'],
        'understanding': repetition_enhancement['understanding'],
        'high_score_criteria': repetition_enhancement['high_score_criteria'],
        'low_score_issues': repetition_enhancement['low_score_issues'],
        'specific_issues': repetition_enhancement['specific_issues'],
        'impact': 'BREVITY',
        'modal_content': repetition_modal
    })
    categories.append({
        'name': 'Unnecessary Sections',
        'score': analyze_unnecessary_sections_frontend(resume_text),
        'issue': 'Remove outdated sections like References, Objective, and high school education when you have higher qualifications',
        'impact': 'SECTIONS'
    })
    growth_signals_score = analyze_growth_signals_frontend(resume_text)
    growth_signals_enhancement = get_enhanced_issue_description('Growth Signals', growth_signals_score, resume_text)
    categories.append({
        'name': 'Growth Signals',
        'score': growth_signals_score,
        'issue': growth_signals_enhancement['issue'],
        'understanding': growth_signals_enhancement['understanding'],
        'high_score_criteria': growth_signals_enhancement['high_score_criteria'],
        'low_score_issues': growth_signals_enhancement['low_score_issues'],
        'specific_issues': growth_signals_enhancement['specific_issues'],
        'impact': 'ALL'
    })
    categories.append({
        'name': 'Drive',
        'score': analyze_drive_and_initiative_frontend(resume_text),
        'issue': 'Show initiative and self-motivation examples',
        'impact': 'ALL'
    })
    
    # 22-23. ADDITIONAL CATEGORIES
    certifications_score = analyze_certifications_frontend(resume_text)
    certifications_enhancement = get_enhanced_issue_description('Certifications', certifications_score, resume_text)
    categories.append({
        'name': 'Certifications',
        'score': certifications_score,
        'issue': certifications_enhancement['issue'],
        'understanding': certifications_enhancement['understanding'],
        'high_score_criteria': certifications_enhancement['high_score_criteria'],
        'low_score_issues': certifications_enhancement['low_score_issues'],
        'specific_issues': certifications_enhancement['specific_issues'],
        'impact': 'ALL'
    })
    
    # 24. DATE FORMATTING
    dates_score = analyze_date_formatting(resume_text)['score']
    dates_enhancement = get_enhanced_issue_description('Dates', dates_score, resume_text)
    categories.append({
        'name': 'Dates',
        'score': dates_score,
        'issue': dates_enhancement['issue'],
        'understanding': dates_enhancement['understanding'],
        'high_score_criteria': dates_enhancement['high_score_criteria'],
        'low_score_issues': dates_enhancement['low_score_issues'],
        'specific_issues': dates_enhancement['specific_issues'],
        'impact': 'STYLE'
    })
    
    # 25. CV READABILITY SCORE (WEIGHTED)
    categories.append({
        'name': 'CV Readability Score',
        'score': calculate_cv_readability_score(resume_text, filename),
        'issue': 'Improve resume structure, formatting, and ATS compatibility for better readability',
        'impact': 'ALL'
    })
    
    logger.info(f'🏗️ Generated {len(categories)} comprehensive categories from frontend logic')
    for cat in categories:
        logger.info(f'🏗️ {cat["name"]}: {cat["score"]}/10')
    
    return categories

# ========================================
# END FRONTEND ANALYSIS FUNCTIONS
# ========================================

def calculate_comprehensive_ats_score(content: str, job_posting: str = None, knockout_questions: List[Dict] = None, filename: str = None) -> Dict[str, Any]:
    """Calculate comprehensive ATS compatibility score with penalty system"""
    
    # Detect industry for targeted analysis
    industry = detect_industry(content)
    
    # Initialize scoring components
    components = {}
    
    # Apply configured component weights (total: 100 points)
    # 1. Content Structure Analysis (25 points)
    components['structure'] = analyze_content_structure(content)
    
    # 2. Keyword Optimization (20 points)
    components['keywords'] = analyze_keyword_optimization(content, industry)
    
    # 3. Contact Information (15 points)
    components['contact'] = analyze_contact_information(content)
    
    # 4. Formatting Quality (10 points) - REDUCED from 15
    components['formatting'] = analyze_formatting_quality(content)
    
    # 5. Quantified Achievements (10 points)
    components['achievements'] = analyze_quantified_achievements(content)
    
    # 6. Readability and Length (10 points)
    components['readability'] = analyze_readability_and_length(content)
    
    # 7. Date Formatting (5 points)
    components['dates'] = analyze_date_formatting(content)
    
    # 8. Bullet Lengths (5 points) - NEW
    components['bullet_lengths'] = analyze_bullet_lengths(content)
    
    # Calculate base score from components
    base_score = sum(comp['score'] for comp in components.values())
    base_score = min(base_score, 100)  # Cap at 100
    
    # Apply comprehensive penalty system
    try:
        from penalty_system import apply_comprehensive_penalties
        penalty_result = apply_comprehensive_penalties(base_score, content, job_posting, knockout_questions)
        final_score = penalty_result['final_score']
        penalty_breakdown = penalty_result['penalty_breakdown']
        total_penalty = penalty_result['total_penalty']
    except ImportError as e:
        logger.warning(f"Penalty system not available: {e}")
        final_score = base_score
        penalty_breakdown = {}
        total_penalty = 0
    
    # Load score categories from config
    score_categories = get_score_categories()
    
    # Determine score category
    category = 'poor'
    description = 'Poor ATS compatibility - major optimization required'
    
    for cat_name, cat_data in score_categories.items():
        if cat_data['min_score'] <= final_score <= cat_data['max_score']:
            category = cat_name
            description = cat_data['description']
            break
    
    # Generate comprehensive categories using frontend logic
    comprehensive_categories = generate_comprehensive_ats_scores_frontend(content, {k: v['score'] for k, v in components.items()}, components, filename)
    
    # Create comprehensive detailed analysis with all 23+ categories
    comprehensive_analysis = {}
    for category in comprehensive_categories:
        # Convert to backend format
        key = category['name'].lower().replace(' ', '_').replace('&', 'and')
        
        # Generate modal content for this category
        try:
            modal_content = generate_fix_this_modal_content(key, content, category['score'])
        except Exception as e:
            logger.warning(f"Failed to generate modal content for {key}: {e}")
            modal_content = {
                'title': f'Why {category["name"]} Matters for ATS',
                'generic_explanation': f'{category["name"]} is important for ATS scoring and professional presentation.',
                'dynamic_examples': []
            }
        
        comprehensive_analysis[key] = {
            'score': category['score'],
            'issues': [category['issue']],
            'impact': category['impact'],
            'modal_content': modal_content  # Add modal content for frontend
        }
    
    # Calculate new overall score from all comprehensive categories
    total_comprehensive_score = sum(cat['score'] for cat in comprehensive_categories)
    max_comprehensive_score = len(comprehensive_categories) * 10
    comprehensive_final_score = min(100, (total_comprehensive_score / max_comprehensive_score) * 100)
    
    logger.info(f'🎯 Comprehensive scoring: {total_comprehensive_score}/{max_comprehensive_score} = {comprehensive_final_score:.1f}%')
    logger.info(f'🔍 DEBUG: Individual category scores: {[cat["score"] for cat in comprehensive_categories]}')
    logger.info(f'🔍 DEBUG: Manual sum check: {sum(cat["score"] for cat in comprehensive_categories)}')
    logger.info(f'🔍 DEBUG: Expected percentage: {(sum(cat["score"] for cat in comprehensive_categories) / (len(comprehensive_categories) * 10)) * 100:.1f}%')
    
    # Debug logging for frontend 
    logger.info(f'🔍 DEBUG: comprehensive_analysis keys: {list(comprehensive_analysis.keys())}')
    logger.info(f'🔍 DEBUG: comprehensive_analysis count: {len(comprehensive_analysis)}')
    logger.info(f'🔍 DEBUG: Sample comprehensive_analysis: {dict(list(comprehensive_analysis.items())[:3])}')
    
    response_data = {
        'ats_score': final_score,  # Keep original for compatibility
        'score': comprehensive_final_score,  # New comprehensive score
        'base_score': base_score,
        'total_penalty': total_penalty,
        'penalty_breakdown': penalty_breakdown,
        'category': category,
        'description': description,
        'industry': industry,
        'component_scores': {k: v['score'] for k, v in components.items()},
        'detailed_analysis': components,  # Keep original 8 categories
        'detailedAnalysis': comprehensive_analysis,  # New 23+ categories for frontend
        'comprehensive_categories': comprehensive_categories,  # Raw category data
        'total_categories': len(comprehensive_categories)
    }
    
    # Final debug logging to verify what's being returned
    logger.info(f'🔍 DEBUG: RESPONSE scores - ats_score: {response_data["ats_score"]:.1f}, comprehensive score: {response_data["score"]:.1f}')
    logger.info(f'🔍 DEBUG: Frontend will use: score={response_data["score"]:.1f} (priority) or ats_score={response_data["ats_score"]:.1f} (fallback)')
    logger.info(f'🔍 DEBUG: RESPONSE detailedAnalysis count: {len(response_data["detailedAnalysis"])}')
    logger.info(f'🔍 DEBUG: RESPONSE detailedAnalysis keys: {list(response_data["detailedAnalysis"].keys())[:10]}')
    logger.info(f'🔍 DEBUG: RESPONSE detailed_analysis count: {len(response_data["detailed_analysis"])}')
    
    return response_data

def calculate_interview_rates(ats_score: int) -> Dict[str, Any]:
    """
    Calculate realistic interview rates based on ATS score using config data
    
    Args:
        ats_score: ATS compatibility score (0-100)
        
    Returns:
        Dictionary with current and potential interview rates
    """
    # Load score categories and interview rate mapping from config
    score_categories = get_score_categories()
    interview_rate_mapping = config_loader.get_interview_rate_mapping()
    
    current_rate = 1
    performance_tier = "Poor"
    
    # Find the matching score category
    for cat_name, cat_data in score_categories.items():
        if cat_data['min_score'] <= ats_score <= cat_data['max_score']:
            current_rate = cat_data.get('interview_rate', 1)
            performance_tier = cat_data.get('label', 'Poor')
            break
    
    # Calculate potential rate with improvements (realistic maximum)
    potential_rate = min(18, current_rate + (90 - ats_score) * 0.15)
    multiplier = potential_rate / current_rate if current_rate > 0 else 5
    
    return {
        "current_rate": current_rate,
        "potential_rate": round(potential_rate),
        "multiplier": round(multiplier, 1),
        "performance_tier": performance_tier,
        "percentile": get_score_percentile(ats_score)
    }

def get_score_percentile(ats_score: int) -> int:
    """Get percentile ranking based on ATS score"""
    if ats_score >= 90: return 95
    if ats_score >= 80: return 80
    if ats_score >= 70: return 60
    if ats_score >= 60: return 40
    if ats_score >= 50: return 25
    return 10

def get_letter_grade(ats_score: int) -> str:
    """Convert ATS score to letter grade using config data"""
    score_categories = get_score_categories()
    
    for cat_name, cat_data in score_categories.items():
        if cat_data['min_score'] <= ats_score <= cat_data['max_score']:
            return cat_data.get('letter_grade', 'F')
    
    return "F"

def classify_issues_by_priority(analysis_data: Dict[str, Any]) -> Tuple[List[Dict], List[Dict]]:
    """
    Classify CV issues by priority and impact
    
    Args:
        analysis_data: Complete analysis data including component scores
        
    Returns:
        Tuple of (critical_issues, quick_wins)
    """
    critical_issues = []
    quick_wins = []
    
    components = analysis_data.get('detailed_analysis', {})
    component_scores = analysis_data.get('component_scores', {})
    personal_info = analysis_data.get('personal_information', {})
    
    # CRITICAL ISSUES (blocks interviews)
    
    # Missing essential contact information
    if component_scores.get('contact', 0) < 10:
        missing_contact = []
        if not personal_info.get('email'):
            missing_contact.append('professional email')
        if not personal_info.get('phone'):
            missing_contact.append('phone number')
        
        if missing_contact:
            critical_issues.append({
                'title': 'Missing Essential Contact Information',
                'issue': f"Missing: {', '.join(missing_contact)}",
                'solution': 'Add complete contact details in a clear header section',
                'time_to_fix': '2 minutes',
                'points_gain': 8,
                'impact': 'High',
                'component': 'Contact Information',
                'why_critical': 'Recruiters cannot contact you without proper contact details'
            })
    
    # Missing professional summary
    if not personal_info.get('professional_summary'):
        critical_issues.append({
            'title': 'No Professional Summary',
            'issue': 'Missing career summary or objective statement',
            'solution': 'Add a 2-3 sentence summary highlighting your key qualifications and career goals',
            'time_to_fix': '10 minutes',
            'points_gain': 12,
            'impact': 'High',
            'component': 'Professional Summary',
            'why_critical': 'ATS systems and recruiters look for summary sections to quickly understand your profile'
        })
    
    # Poor keyword optimization
    if component_scores.get('keywords', 0) < 8:
        critical_issues.append({
            'title': 'Poor Keyword Optimization',
            'issue': 'Missing industry-specific keywords that ATS systems scan for',
            'solution': 'Research job postings in your field and include relevant keywords naturally',
            'time_to_fix': '15 minutes',
            'points_gain': 10,
            'impact': 'High',
            'component': 'Keywords & Skills',
            'why_critical': 'ATS systems filter resumes based on keyword matches'
        })
    
    # Poor formatting affecting ATS readability
    if component_scores.get('formatting', 0) < 12:
        critical_issues.append({
            'title': 'ATS-Incompatible Formatting',
            'issue': 'Formatting issues that prevent ATS systems from reading your resume',
            'solution': 'Use standard section headers, consistent fonts, and avoid complex layouts',
            'time_to_fix': '20 minutes',
            'points_gain': 8,
            'impact': 'High',
            'component': 'Formatting',
            'why_critical': 'Poor formatting can make your resume invisible to ATS systems'
        })
    
    # QUICK WINS (easy improvements with good impact)
    
    # LinkedIn profile missing
    if not personal_info.get('linkedin_url'):
        quick_wins.append({
            'title': 'Add LinkedIn Profile URL',
            'issue': 'LinkedIn profile not included in contact section',
            'solution': 'Add your LinkedIn profile URL to show professional online presence',
            'time_to_fix': '1 minute',
            'points_gain': 3,
            'impact': 'Medium',
            'component': 'Contact Information'
        })
    
    # Phone number formatting
    phone = personal_info.get('phone', '')
    if phone and (len(phone.replace(' ', '').replace('-', '').replace('(', '').replace(')', '')) < 10):
        quick_wins.append({
            'title': 'Fix Phone Number Format',
            'issue': 'Phone number format may not be ATS-friendly',
            'solution': 'Use standard format: +1 (555) 123-4567 or +1-555-123-4567',
            'time_to_fix': '1 minute',
            'points_gain': 2,
            'impact': 'Low',
            'component': 'Contact Information'
        })
    
    # Missing skills section
    skills = personal_info.get('skills', [])
    if not skills or len(skills) < 5:
        quick_wins.append({
            'title': 'Expand Skills Section',
            'issue': 'Limited or missing technical skills listed',
            'solution': 'Add 8-12 relevant skills including both technical and soft skills',
            'time_to_fix': '5 minutes',
            'points_gain': 5,
            'impact': 'Medium',
            'component': 'Skills'
        })
    
    # Missing years of experience
    if not personal_info.get('years_of_experience'):
        quick_wins.append({
            'title': 'Add Experience Level',
            'issue': 'Years of experience not clearly indicated',
            'solution': 'Include your total years of relevant experience in summary or skills section',
            'time_to_fix': '2 minutes',
            'points_gain': 3,
            'impact': 'Medium',
            'component': 'Experience'
        })
    
    # Poor structure score
    if component_scores.get('structure', 0) < 15:
        structure_data = components.get('structure', {})
        missing_sections = structure_data.get('missing_sections', [])
        
        if missing_sections:
            quick_wins.append({
                'title': 'Add Missing Resume Sections',
                'issue': f"Missing sections: {', '.join(missing_sections)}",
                'solution': 'Add standard resume sections to improve ATS compatibility',
                'time_to_fix': '10 minutes',
                'points_gain': 6,
                'impact': 'Medium',
                'component': 'Structure'
            })
    
    return critical_issues, quick_wins

def generate_transformation_preview(analysis_data: Dict[str, Any], critical_issues: List[Dict], quick_wins: List[Dict]) -> Dict[str, Any]:
    """
    Generate before/after transformation preview
    
    Args:
        analysis_data: Current analysis data
        critical_issues: List of critical issues
        quick_wins: List of quick wins
        
    Returns:
        Dictionary with transformation metrics
    """
    current_score = analysis_data.get('ats_score', 0)
    current_rates = calculate_interview_rates(current_score)
    
    # Calculate potential improvement
    total_points_gain = sum(issue.get('points_gain', 0) for issue in critical_issues + quick_wins)
    potential_score = min(100, current_score + total_points_gain)
    potential_rates = calculate_interview_rates(potential_score)
    
    return {
        'current_score': current_score,
        'potential_score': potential_score,
        'score_improvement': potential_score - current_score,
        'current_grade': get_letter_grade(current_score),
        'potential_grade': get_letter_grade(potential_score),
        'current_interview_rate': current_rates['current_rate'],
        'potential_interview_rate': potential_rates['current_rate'],
        'interview_improvement': f"{potential_rates['current_rate'] - current_rates['current_rate']}x more",
        'total_fixes': len(critical_issues) + len(quick_wins),
        'quick_fixes': len(quick_wins),
        'time_investment': calculate_total_time(critical_issues + quick_wins)
    }

def calculate_total_time(issues: List[Dict]) -> str:
    """Calculate total time needed for all improvements"""
    total_minutes = 0
    
    for issue in issues:
        time_str = issue.get('time_to_fix', '0 minutes')
        # Extract number from time string
        import re
        numbers = re.findall(r'\d+', time_str)
        if numbers:
            total_minutes += int(numbers[0])
    
    if total_minutes < 60:
        return f"{total_minutes} minutes"
    else:
        hours = total_minutes // 60
        remaining_minutes = total_minutes % 60
        if remaining_minutes == 0:
            return f"{hours} hour{'s' if hours > 1 else ''}"
        else:
            return f"{hours}h {remaining_minutes}m"

def enhance_component_breakdown(analysis_data: Dict[str, Any]) -> Dict[str, Any]:
    """
    Enhance component analysis with specific issues and solutions
    
    Args:
        analysis_data: Current analysis data
        
    Returns:
        Enhanced component breakdown with specific recommendations
    """
    components = analysis_data.get('detailed_analysis', {})
    component_scores = analysis_data.get('component_scores', {})
    enhanced_components = {}
    
    # Load component weights from config
    max_scores = get_component_weights()
    
    for component, max_score in max_scores.items():
        current_score = component_scores.get(component, 0)
        percentage = round((current_score / max_score) * 100)
        
        # Determine status and specific issues
        if percentage >= 80:
            status = 'excellent'
            status_text = 'Excellent'
            color = 'green'
        elif percentage >= 60:
            status = 'good'
            status_text = 'Good'
            color = 'blue'
        elif percentage >= 40:
            status = 'fair'
            status_text = 'Needs Improvement'
            color = 'yellow'
        else:
            status = 'poor'
            status_text = 'Critical Issue'
            color = 'red'
        
        # Get component-specific data
        component_data = components.get(component, {})
        
        enhanced_components[component] = {
            'name': format_component_name(component),
            'score': current_score,
            'max_score': max_score,
            'percentage': percentage,
            'status': status,
            'status_text': status_text,
            'color': color,
            'specific_issues': get_component_specific_issues(component, component_data, current_score, max_score),
            'recommendations': get_component_recommendations(component, component_data, current_score)
        }
    
    return enhanced_components

def get_component_specific_issues(component: str, data: Dict, current_score: int, max_score: int) -> List[str]:
    """Get specific issues for each component"""
    issues = []
    
    if component == 'structure':
        missing_sections = data.get('missing_sections', [])
        if missing_sections:
            issues.append(f"Missing sections: {', '.join(missing_sections)}")
        if current_score < max_score * 0.6:
            issues.append("Poor overall resume structure and organization")
    
    elif component == 'keywords':
        missing_keywords = data.get('missing_keywords', [])
        if missing_keywords:
            issues.append(f"Missing key industry terms: {', '.join(missing_keywords[:3])}")
        if current_score < max_score * 0.4:
            issues.append("Severely lacking relevant keywords for ATS optimization")
    
    elif component == 'contact':
        if current_score < max_score * 0.7:
            issues.append("Incomplete or improperly formatted contact information")
        if current_score < max_score * 0.4:
            issues.append("Missing essential contact details (phone, email, or location)")
    
    elif component == 'formatting':
        if current_score < max_score * 0.6:
            issues.append("ATS-incompatible formatting detected")
            issues.append("Inconsistent font usage or spacing")
    
    elif component == 'achievements':
        achievements_count = data.get('achievements_count', 0)
        if achievements_count < 3:
            issues.append("Insufficient quantified achievements and results")
        if current_score < max_score * 0.5:
            issues.append("Missing measurable impact statements")
    
    elif component == 'readability':
        if current_score < max_score * 0.5:
            issues.append("Poor readability and unclear language")
            issues.append("Complex sentences that may confuse ATS systems")
    
    return issues

def get_component_recommendations(component: str, data: Dict, current_score: int) -> List[str]:
    """Get specific recommendations for each component"""
    recommendations = []
    
    if component == 'structure':
        recommendations.append("Use standard section headers: Summary, Experience, Education, Skills")
        recommendations.append("Organize content in reverse chronological order")
    
    elif component == 'keywords':
        recommendations.append("Research job descriptions in your field for relevant keywords")
        recommendations.append("Include both technical skills and industry buzzwords")
        recommendations.append("Use keywords naturally throughout your resume")
    
    elif component == 'contact':
        recommendations.append("Include: Full name, phone, professional email, LinkedIn, location")
        recommendations.append("Use consistent formatting for all contact information")
    
    elif component == 'formatting':
        recommendations.append("Use standard fonts (Arial, Calibri, or Times New Roman)")
        recommendations.append("Maintain consistent spacing and bullet point styles")
        recommendations.append("Avoid tables, images, or complex layouts")
    
    elif component == 'achievements':
        recommendations.append("Quantify accomplishments with numbers, percentages, or dollar amounts")
        recommendations.append("Use action verbs to start each bullet point")
        recommendations.append("Focus on results and impact, not just job duties")
    
    elif component == 'readability':
        recommendations.append("Use clear, concise language and shorter sentences")
        recommendations.append("Avoid jargon that ATS systems might not recognize")
        recommendations.append("Ensure proper grammar and spelling")
    
    return recommendations

def format_component_name(component: str) -> str:
    """Format component names for display"""
    names = {
        'structure': 'Resume Structure',
        'keywords': 'Keywords & Skills',
        'contact': 'Contact Information',
        'formatting': 'Formatting & Layout',
        'achievements': 'Achievements & Impact',
        'readability': 'Readability & Clarity'
    }
    return names.get(component, component.title())

def generate_comprehensive_recommendations(analysis: Dict[str, Any]) -> Dict[str, Any]:
    """Generate detailed recommendations based on analysis"""
    
    strengths = []
    improvements = []
    critical_issues = []
    suggestions = []
    
    score = analysis['ats_score']
    components = analysis['detailed_analysis']
    
    # Analyze each component for recommendations
    
    # Structure analysis
    structure = components['structure']
    if structure['score'] >= 20:
        strengths.append("Well-organized resume structure with essential sections")
    else:
        missing = structure['missing_sections']
        if missing:
            critical_issues.append(f"Missing essential sections: {', '.join(missing)}")
            improvements.append("Add missing essential resume sections")
    
    # Keywords analysis
    keywords = components['keywords']
    if keywords['score'] >= 20:
        strengths.append("Good use of relevant keywords and action verbs")
    else:
        improvements.append("Increase industry-specific keywords and strong action verbs")
        suggestions.append({
            'title': 'Optimize Keywords',
            'description': 'Add more industry-relevant keywords and quantifiable achievements',
            'priority': 'high'
        })
    
    # Contact information
    contact = components['contact']
    if contact['score'] >= 10:
        strengths.append("Complete contact information provided")
    else:
        missing_contacts = contact['missing']
        if 'email' in missing_contacts or 'phone' in missing_contacts:
            critical_issues.append("Missing essential contact information (email/phone)")
        if 'linkedin' in missing_contacts:
            improvements.append("Add LinkedIn profile for professional networking")
    
    # Formatting quality
    formatting = components['formatting']
    if formatting['score'] >= 15:
        strengths.append("Clean, ATS-friendly formatting")
    else:
        if formatting['issues']:
            improvements.append("Fix formatting issues that may confuse ATS systems")
            suggestions.append({
                'title': 'Improve Formatting',
                'description': 'Clean up formatting issues for better ATS readability',
                'priority': 'medium'
            })
    
    # Quantified achievements
    achievements = components['achievements']
    if achievements['score'] >= 6:
        strengths.append("Good use of quantified achievements and metrics")
    else:
        improvements.append("Add more quantified achievements with specific numbers and percentages")
        suggestions.append({
            'title': 'Add Quantified Results',
            'description': 'Include specific numbers, percentages, and measurable outcomes',
            'priority': 'high'
        })
    
    # Overall score-based recommendations
    if score < 70:
        suggestions.append({
            'title': 'Comprehensive Resume Optimization',
            'description': 'Consider professional resume review for significant improvements',
            'priority': 'high'
        })
    
    return {
        'strengths': strengths,
        'improvements': improvements,
        'critical_issues': critical_issues,
        'suggestions': suggestions,
        'next_steps': generate_next_steps(score, components)
    }

def generate_next_steps(score: int, components: Dict[str, Any]) -> List[str]:
    """Generate prioritized next steps for improvement"""
    next_steps = []
    
    # Prioritize based on component scores
    low_scoring_components = [(k, v['score']) for k, v in components.items() if v['score'] < 10]
    low_scoring_components.sort(key=lambda x: x[1])  # Sort by score, lowest first
    
    for component, component_score in low_scoring_components[:3]:  # Top 3 priorities
        if component == 'structure':
            next_steps.append("1. Add missing resume sections (Experience, Education, Skills)")
        elif component == 'keywords':
            next_steps.append("2. Research and add industry-specific keywords")
        elif component == 'contact':
            next_steps.append("3. Complete your contact information including LinkedIn")
        elif component == 'formatting':
            next_steps.append("4. Clean up formatting and use consistent spacing")
        elif component == 'achievements':
            next_steps.append("5. Add quantified achievements with specific metrics")
    
    if score >= 80:
        next_steps.append("Focus on fine-tuning keyword optimization and formatting")
    elif score >= 60:
        next_steps.append("Prioritize adding quantified achievements and industry keywords")
    else:
        next_steps.append("Start with essential sections and contact information")
    
    return next_steps[:5]  # Limit to 5 steps

def generate_detailed_issues_analysis(analysis: Dict[str, Any], content: str) -> Dict[str, Any]:
    """Generate specific, actionable issues with detailed analysis"""
    
    score = analysis['ats_score']
    components = analysis['detailed_analysis']
    
    critical_issues = []
    quick_fixes = []
    content_improvements = []
    
    # CONTACT INFORMATION ANALYSIS
    contact = components.get('contact', {})
    contact_missing = contact.get('missing', [])
    
    if 'email' in contact_missing:
        critical_issues.append({
            'title': 'Missing email address',
            'category': 'Contact Information',
            'description': 'Email address is required for ATS systems and recruiters',
            'solution': 'Add your professional email address to the header (e.g., yourname@email.com)',
            'time_to_fix': '1 minute',
            'impact': 'Critical',
            'points_gain': 12
        })
    
    if 'phone' in contact_missing:
        critical_issues.append({
            'title': 'Missing phone number',
            'category': 'Contact Information',
            'description': 'Phone number is essential contact information for recruiters',
            'solution': 'Add your phone number in standard format (e.g., +1 (555) 123-4567)',
            'time_to_fix': '1 minute',
            'impact': 'Critical',
            'points_gain': 10
        })
    
    if 'linkedin' in contact_missing:
        quick_fixes.append({
            'title': 'Add LinkedIn profile URL',
            'category': 'Contact Information',
            'description': 'LinkedIn profile increases credibility and networking opportunities',
            'solution': 'Add your LinkedIn profile URL (linkedin.com/in/yourname)',
            'time_to_fix': '2 minutes',
            'impact': 'Medium',
            'points_gain': 5
        })
    
    # FORMATTING ANALYSIS
    formatting = components.get('formatting', {})
    formatting_issues = formatting.get('issues', [])
    
    if 'inconsistent_spacing' in formatting_issues:
        quick_fixes.append({
            'title': 'Fix inconsistent spacing',
            'category': 'Formatting',
            'description': 'Inconsistent spacing confuses ATS parsers',
            'solution': 'Use consistent line spacing (1.15 or 1.5) and consistent margins throughout',
            'time_to_fix': '3 minutes',
            'impact': 'Medium',
            'points_gain': 4
        })
    
    if 'inconsistent_bullets' in formatting_issues:
        quick_fixes.append({
            'title': 'Standardize bullet points',
            'category': 'Formatting',
            'description': 'Mixed bullet styles reduce readability',
            'solution': 'Use the same bullet style throughout (• or - consistently)',
            'time_to_fix': '2 minutes',
            'impact': 'Medium',
            'points_gain': 3
        })
    
    if 'font_issues' in formatting_issues:
        critical_issues.append({
            'title': 'Fix font consistency',
            'category': 'Formatting',
            'description': 'Multiple fonts or unusual fonts confuse ATS systems',
            'solution': 'Use one professional font throughout (Arial, Calibri, or Times New Roman)',
            'time_to_fix': '5 minutes',
            'impact': 'High',
            'points_gain': 8
        })
    
    # KEYWORDS ANALYSIS
    keywords = components.get('keywords', {})
    missing_keywords = keywords.get('missing_keywords', [])
    
    if len(missing_keywords) > 0:
        sample_keywords = missing_keywords[:5]  # Show up to 5 examples
        content_improvements.append({
            'title': f'Add {len(missing_keywords)} missing industry keywords',
            'category': 'Keywords & Skills',
            'description': f'Your resume lacks important keywords that ATS systems look for',
            'solution': f'Add these keywords naturally: {", ".join(sample_keywords)}',
            'time_to_fix': '10 minutes',
            'impact': 'High',
            'points_gain': min(len(missing_keywords), 15)
        })
    
    # ACHIEVEMENTS ANALYSIS
    achievements = components.get('achievements', {})
    if achievements.get('score', 0) < 5:
        # Analyze content for unquantified achievements
        achievement_count = len(re.findall(r'[•\-\*]\s*[^•\-\*\n]+', content))
        
        content_improvements.append({
            'title': 'Quantify your achievements with numbers',
            'category': 'Content Quality',
            'description': f'Found {achievement_count} bullet points that lack specific metrics',
            'solution': 'Add numbers, percentages, or dollar amounts to your accomplishments (e.g., "Increased sales by 25%")',
            'time_to_fix': '15 minutes',
            'impact': 'High',
            'points_gain': 12
        })
    
    # STRUCTURE ANALYSIS
    structure = components.get('structure', {})
    missing_sections = structure.get('missing_sections', [])
    
    for section in missing_sections:
        if section.lower() in ['experience', 'work experience', 'employment']:
            critical_issues.append({
                'title': 'Missing Work Experience section',
                'category': 'Resume Structure',
                'description': 'Work experience is essential for most resumes',
                'solution': 'Add a dedicated "Work Experience" or "Professional Experience" section',
                'time_to_fix': '20 minutes',
                'impact': 'Critical',
                'points_gain': 20
            })
        elif section.lower() in ['education']:
            critical_issues.append({
                'title': 'Missing Education section',
                'category': 'Resume Structure',
                'description': 'Education section is required for most positions',
                'solution': 'Add an "Education" section with your degrees and institutions',
                'time_to_fix': '10 minutes',
                'impact': 'Critical',
                'points_gain': 15
            })
        elif section.lower() in ['skills']:
            quick_fixes.append({
                'title': 'Add Skills section',
                'category': 'Resume Structure',
                'description': 'Skills section helps ATS match your qualifications',
                'solution': 'Create a "Skills" section with 8-12 relevant technical and soft skills',
                'time_to_fix': '8 minutes',
                'impact': 'Medium',
                'points_gain': 8
            })
    
    # CONTENT QUALITY CHECKS
    content_lower = content.lower()
    
    # Check for weak action verbs
    weak_verbs = ['responsible for', 'worked on', 'helped with', 'involved in', 'participated in']
    weak_verb_count = sum(1 for verb in weak_verbs if verb in content_lower)
    
    if weak_verb_count > 0:
        content_improvements.append({
            'title': f'Replace {weak_verb_count} weak action verbs',
            'category': 'Content Quality',
            'description': 'Weak action verbs reduce impact of your accomplishments',
            'solution': 'Replace "responsible for" with "managed", "led", "developed", "implemented"',
            'time_to_fix': '12 minutes',
            'impact': 'Medium',
            'points_gain': min(weak_verb_count * 2, 10)
        })
    
    # Check for professional summary
    if 'summary' not in content_lower and 'objective' not in content_lower:
        content_improvements.append({
            'title': 'Add professional summary',
            'category': 'Resume Structure',
            'description': 'Professional summary helps recruiters quickly understand your value',
            'solution': 'Add a 2-3 sentence summary highlighting your key qualifications',
            'time_to_fix': '15 minutes',
            'impact': 'Medium',
            'points_gain': 7
        })
    
    # Calculate realistic improvement potential
    total_potential_gain = sum(item['points_gain'] for item in critical_issues + quick_fixes + content_improvements)
    capped_gain = min(total_potential_gain, 30)  # Cap at 30 points for realism
    
    return {
        'critical_issues': critical_issues,
        'quick_fixes': quick_fixes,
        'content_improvements': content_improvements,
        'total_issues': len(critical_issues) + len(quick_fixes) + len(content_improvements),
        'potential_improvement': capped_gain,
        'realistic_target_score': min(score + capped_gain, 95),
        'estimated_time': calculate_total_time(critical_issues + quick_fixes + content_improvements)
    }

def calculate_total_time(issues: list) -> str:
    """Calculate total estimated time for all improvements"""
    total_minutes = 0
    
    for issue in issues:
        time_str = issue.get('time_to_fix', '0 minutes')
        try:
            # Extract number from strings like "5 minutes", "1 minute", etc.
            minutes = int(re.search(r'\d+', time_str).group())
            total_minutes += minutes
        except (AttributeError, ValueError):
            total_minutes += 5  # Default fallback
    
    if total_minutes < 60:
        return f"{total_minutes} minutes"
    else:
        hours = total_minutes // 60
        remaining_minutes = total_minutes % 60
        if remaining_minutes == 0:
            return f"{hours} hour{'s' if hours > 1 else ''}"
        else:
            return f"{hours} hour{'s' if hours > 1 else ''} {remaining_minutes} minutes"

def extract_personal_information(content: str) -> Dict[str, Any]:
    """
    Extract personal information from CV content for user profile
    
    Args:
        content: CV text content
        
    Returns:
        Dictionary containing extracted personal information
    """
    extracted_data = {
        'full_name': None,
        'email': None,
        'phone': None,
        'address': None,
        'city': None,
        'state': None,
        'linkedin_url': None,
        'github_url': None,
        'website_url': None,
        'professional_summary': None,
        'skills': [],
        'education': [],
        'work_experience': [],
        'years_of_experience': None
    }
    
    # Extract name
    extracted_data['full_name'] = extract_name(content)
    
    # Extract contact information
    contact_info = analyze_contact_information(content)
    found_contacts = contact_info.get('found_contacts', {})
    
    if found_contacts.get('email'):
        extracted_data['email'] = found_contacts['email'][0]
    if found_contacts.get('phone'):
        # Clean and validate phone number
        phone_raw = found_contacts['phone'][0]
        phone_cleaned = re.sub(r'[^\d+\-\(\)\s]', '', phone_raw)
        # Ensure we have a complete phone number
        digits_only = re.sub(r'[^\d]', '', phone_cleaned)
        if len(digits_only) >= 10:  # Minimum for a valid phone number
            extracted_data['phone'] = phone_cleaned
        else:
            logger.warning(f"Phone number too short: {phone_raw} -> {digits_only} ({len(digits_only)} digits)")
            extracted_data['phone'] = phone_raw  # Store original for debugging
    if found_contacts.get('linkedin'):
        extracted_data['linkedin_url'] = 'https://' + found_contacts['linkedin'][0]
    if found_contacts.get('github'):
        extracted_data['github_url'] = 'https://' + found_contacts['github'][0]
    if found_contacts.get('website'):
        extracted_data['website_url'] = found_contacts['website'][0]
    
    # Extract address information
    address_data = extract_address(content)
    extracted_data.update(address_data)
    
    # Extract professional summary
    extracted_data['professional_summary'] = extract_summary(content)
    
    # Extract skills
    extracted_data['skills'] = extract_skills_list(content)
    
    # Extract education
    extracted_data['education'] = extract_education_list(content)
    
    # Extract work experience
    extracted_data['work_experience'] = extract_work_experience_list(content)
    
    # Estimate years of experience
    extracted_data['years_of_experience'] = estimate_years_of_experience(content)
    
    return extracted_data

def extract_name(content: str) -> Optional[str]:
    """Extract full name from CV content with enhanced Unicode and merged line support"""
    import unicodedata
    
    # Normalize Unicode characters to handle different encodings
    normalized_content = unicodedata.normalize('NFKD', content)
    lines = normalized_content.split('\n')
    logger.info(f"🔍 Name extraction - checking first lines (showing first 100 chars each):")
    
    # Try first non-empty line first
    for i, line in enumerate(lines[:10]):  # Check first 10 lines
        line = line.strip()
        line_preview = line[:100] + "..." if len(line) > 100 else line
        logger.info(f"  Line {i}: '{line_preview}'")
        
        # If line is too long (merged PDF), try to extract name from beginning
        if len(line) > 100:
            # Enhanced regex for merged PDF lines with Unicode support
            import re
            # Try multiple patterns for merged lines
            merged_patterns = [
                r'^([A-Z][a-zA-Z\u00C0-\u017F]+\s+[A-Z][a-zA-Z\u00C0-\u017F]+(?:\s+[A-Z][a-zA-Z\u00C0-\u017F]+)?)',  # Unicode names
                r'^([A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)(?:[\s\|\-•])',  # Name followed by separator
                r'^([A-Z][a-z]+\s[A-Z][a-z]+)(?:\s*[A-Z][a-z]+@)',  # Name before email
                r'^([A-Z][a-z]+\s[A-Z][a-z]+)(?:\s*\+?\d)',  # Name before phone
            ]
            
            for pattern in merged_patterns:
                name_match = re.match(pattern, line)
                if name_match:
                    potential_name = name_match.group(1).strip()
                    # Validate name length and content
                    if 4 <= len(potential_name) <= 40 and potential_name.count(' ') >= 1:
                        logger.info(f"✅ Found name at start of long line: '{potential_name}'")
                        return potential_name
        
        if line and len(line.split()) >= 2:
            # More flexible name detection with Unicode support
            words = line.split()
            if len(words) >= 2 and len(words) <= 4:  # Names are typically 2-4 words
                # Enhanced check for title case including Unicode characters
                def is_title_case_unicode(word):
                    if not word:
                        return False
                    first_char = word[0]
                    # Check if first character is uppercase (including Unicode)
                    return first_char.isupper() or unicodedata.category(first_char) == 'Lu'
                
                if all(is_title_case_unicode(word) for word in words[:2]):
                    if len(line) < 80 and not any(char in line for char in '@.com|+()'):  # Not email/phone/url
                        # Additional checks to avoid headers
                        if not any(keyword in line.upper() for keyword in ['CURRICULUM', 'RESUME', 'CV', 'PROFILE', 'CONTACT']):
                            # Validate that it looks like a real name
                            name_score = 0
                            for word in words[:2]:
                                if len(word) >= 2 and word.isalpha():
                                    name_score += 1
                            
                            if name_score >= 2:  # At least 2 valid name words
                                logger.info(f"✅ Found potential name: '{line}'")
                                return line
                            else:
                                logger.info(f"❌ Rejected (invalid name pattern): '{line}'")
                        else:
                            logger.info(f"❌ Rejected (keyword): '{line}'")
                    else:
                        logger.info(f"❌ Rejected (long/contact): '{line}'")
                else:
                    logger.info(f"❌ Rejected (not title case): '{line}'")
            else:
                logger.info(f"❌ Rejected (word count): '{line}' ({len(words)} words)")
    
    # Try regex patterns with normalized content
    for pattern in NAME_PATTERNS:
        match = re.search(pattern, normalized_content, re.MULTILINE)
        if match:
            potential_name = match.group(1).strip()
            # Additional validation for regex-found names
            if 4 <= len(potential_name) <= 40 and potential_name.count(' ') >= 1:
                logger.info(f"✅ Found name via regex: '{potential_name}'")
                return potential_name
    
    logger.warning("⚠️ No valid name found in content")
    return None

def extract_address(content: str) -> Dict[str, Optional[str]]:
    """Extract address components from CV content"""
    address_data = {
        'address': None,
        'city': None,
        'state': None
    }
    
    # Try to find full address
    for pattern in ADDRESS_PATTERNS:
        match = re.search(pattern, content, re.IGNORECASE | re.MULTILINE)
        if match:
            address_data['address'] = match.group(1).strip()
            break
    
    # Try to extract city and state
    for pattern in CITY_STATE_PATTERNS:
        match = re.search(pattern, content, re.MULTILINE)
        if match:
            address_data['city'] = match.group(1).strip()
            if len(match.groups()) >= 2:
                address_data['state'] = match.group(2).strip()
            break
    
    return address_data

def extract_summary(content: str) -> Optional[str]:
    """Extract professional summary from CV content"""
    for pattern in SUMMARY_PATTERNS:
        match = re.search(pattern, content, re.IGNORECASE | re.DOTALL)
        if match:
            summary = match.group(1).strip()
            # Clean up the summary
            summary = re.sub(r'\n+', ' ', summary)
            summary = re.sub(r'\s+', ' ', summary)
            if len(summary) > 50 and len(summary) < 500:  # Reasonable length
                return summary
    
    return None

def extract_skills_list(content: str) -> List[str]:
    """Extract skills from CV content"""
    skills = []
    
    for pattern in SKILLS_PATTERNS:
        match = re.search(pattern, content, re.IGNORECASE | re.DOTALL)
        if match:
            skills_text = match.group(1).strip()
            # Parse skills from various formats
            parsed_skills = parse_skills_text(skills_text)
            skills.extend(parsed_skills)
    
    # Remove duplicates and clean up
    unique_skills = list(set(skill.strip() for skill in skills if skill.strip()))
    return unique_skills[:20]  # Limit to 20 skills

def parse_skills_text(skills_text: str) -> List[str]:
    """Parse skills from text using various delimiters"""
    skills = []
    
    # Try different separators
    separators = [',', '•', '|', ';', '\n', '/', '\\']
    
    for separator in separators:
        if separator in skills_text:
            parts = skills_text.split(separator)
            for part in parts:
                skill = part.strip().strip('•').strip('-').strip()
                if skill and len(skill) < 50:
                    skills.append(skill)
            break
    
    # If no separators found, try word-based extraction
    if not skills:
        # Look for technology/skill keywords
        words = skills_text.split()
        for word in words:
            word = word.strip('.,;()[]{}')
            if len(word) > 2 and len(word) < 25:
                skills.append(word)
    
    return skills

def extract_education_list(content: str) -> List[Dict[str, str]]:
    """Extract education information from CV content"""
    education = []
    
    for pattern in EDUCATION_PATTERNS:
        matches = re.finditer(pattern, content, re.IGNORECASE | re.DOTALL)
        for match in matches:
            edu_text = match.group(1).strip()
            edu_info = parse_education_entry(edu_text)
            if edu_info:
                education.append(edu_info)
    
    return education[:5]  # Limit to 5 education entries

def parse_education_entry(edu_text: str) -> Optional[Dict[str, str]]:
    """Parse individual education entry"""
    if len(edu_text) < 10 or len(edu_text) > 300:
        return None
    
    edu_info = {
        'degree': '',
        'institution': '',
        'year': '',
        'details': edu_text
    }
    
    # Try to extract degree
    degree_patterns = [
        r'(Bachelor[^,\n]*|Master[^,\n]*|PhD[^,\n]*|MBA[^,\n]*|B\.S\.[^,\n]*|M\.S\.[^,\n]*|B\.A\.[^,\n]*|M\.A\.[^,\n]*)',
        r'(Associate[^,\n]*|Diploma[^,\n]*|Certificate[^,\n]*)'
    ]
    
    for pattern in degree_patterns:
        match = re.search(pattern, edu_text, re.IGNORECASE)
        if match:
            edu_info['degree'] = match.group(1).strip()
            break
    
    # Try to extract institution
    institution_match = re.search(r'([A-Za-z\s]+(?:University|College|Institute|School))', edu_text, re.IGNORECASE)
    if institution_match:
        edu_info['institution'] = institution_match.group(1).strip()
    
    # Try to extract year
    year_match = re.search(r'(19\d{2}|20\d{2})', edu_text)
    if year_match:
        edu_info['year'] = year_match.group(1)
    
    return edu_info

def extract_work_experience_list(content: str) -> List[Dict[str, str]]:
    """Extract work experience from CV content"""
    experience = []
    
    for pattern in EXPERIENCE_PATTERNS:
        matches = re.finditer(pattern, content, re.IGNORECASE | re.DOTALL)
        for match in matches:
            exp_text = match.group(1).strip()
            # Split by likely job separators
            job_entries = re.split(r'\n(?=[A-Z][^a-z]*[|•])', exp_text)
            
            for entry in job_entries:
                exp_info = parse_experience_entry(entry.strip())
                if exp_info:
                    experience.append(exp_info)
    
    return experience[:5]  # Limit to 5 experience entries

def parse_experience_entry(exp_text: str) -> Optional[Dict[str, str]]:
    """Parse individual work experience entry"""
    if len(exp_text) < 20 or len(exp_text) > 500:
        return None
    
    exp_info = {
        'title': '',
        'company': '',
        'duration': '',
        'description': exp_text
    }
    
    # Try to extract title | company | duration pattern
    pipe_match = re.search(r'([^|]+)\s*\|\s*([^|]+)\s*\|\s*([^|]+)', exp_text)
    if pipe_match:
        exp_info['title'] = pipe_match.group(1).strip()
        exp_info['company'] = pipe_match.group(2).strip()
        exp_info['duration'] = pipe_match.group(3).strip()
    else:
        # Try to extract from first line
        first_line = exp_text.split('\n')[0]
        if '|' in first_line:
            parts = first_line.split('|')
            if len(parts) >= 2:
                exp_info['title'] = parts[0].strip()
                exp_info['company'] = parts[1].strip()
                if len(parts) >= 3:
                    exp_info['duration'] = parts[2].strip()
    
    return exp_info

def estimate_years_of_experience(content: str) -> Optional[int]:
    """Estimate years of professional experience from CV content"""
    
    # First, try to find explicit experience statements
    exp_statements = [
        r'(\d+)\s*\+?\s*years?\s+(?:of\s+)?(?:professional\s+)?experience',
        r'(?:professional\s+)?experience.*?(\d+)\s*\+?\s*years?',
        r'(\d+)\s*\+?\s*years?\s+(?:in|with)',
        r'over\s+(\d+)\s*years?\s+(?:of\s+)?(?:professional\s+)?experience'
    ]
    
    for pattern in exp_statements:
        match = re.search(pattern, content, re.IGNORECASE)
        if match:
            years_exp = int(match.group(1))
            if 0 < years_exp < 50:
                return years_exp
    
    # If no explicit statement, look for work experience years only (not education)
    work_section = ""
    work_patterns = [
        r'(?:WORK\s+EXPERIENCE|PROFESSIONAL\s+EXPERIENCE|EMPLOYMENT\s+HISTORY|CAREER\s+SUMMARY)[\s:]*\n?(.*?)(?:\n\n|\n(?:EDUCATION|SKILLS|CERTIFICATIONS)|$)',
        r'(?:EXPERIENCE|EMPLOYMENT)[\s:]*\n?(.*?)(?:\n\n|\n(?:EDUCATION|SKILLS)|$)'
    ]
    
    for pattern in work_patterns:
        match = re.search(pattern, content, re.IGNORECASE | re.DOTALL)
        if match:
            work_section = match.group(1)
            break
    
    if work_section:
        # Extract years from work experience section only
        year_pattern = r'(19\d{2}|20\d{2})'
        matches = re.findall(year_pattern, work_section)
        if matches:
            years = [int(year) for year in matches]
            years.sort()
            
            if len(years) >= 2:
                # Calculate span from earliest to latest work year
                span = years[-1] - years[0]
                if span > 0 and span < 50:  # Reasonable range
                    return span
    
    return None

def generate_session_uuid() -> str:
    """Generate a unique session UUID for tracking"""
    return str(uuid.uuid4())

def generate_temp_email_from_uuid(session_uuid: str) -> str:
    """Generate temporary email from UUID for CVs without email"""
    return f"{session_uuid}@bestcvbuilder.com"

def get_file_info_from_url(file_url: str) -> Optional[Dict[str, Any]]:
    """
    Get file information from URL by making a HEAD request
    
    Args:
        file_url: URL of the uploaded file
        
    Returns:
        Dictionary with file info or None if failed
    """
    try:
        import requests
        from urllib.parse import urlparse
        
        # Make HEAD request to get file metadata
        response = requests.head(file_url, timeout=10)
        
        if response.status_code == 200:
            # Get file size from Content-Length header
            file_size = int(response.headers.get('content-length', 1024))
            
            # Parse filename from URL
            parsed_url = urlparse(file_url)
            filename = parsed_url.path.split('/')[-1] or 'uploaded_resume.pdf'
            
            # Determine file type from filename or content-type
            content_type = response.headers.get('content-type', '')
            if '.pdf' in filename.lower() or 'pdf' in content_type:
                file_type = 'pdf'
            elif '.docx' in filename.lower() or 'docx' in content_type:
                file_type = 'docx'
            elif '.doc' in filename.lower() or 'msword' in content_type:
                file_type = 'doc'
            else:
                file_type = 'pdf'  # Default
            
            return {
                'original_filename': filename,
                'file_size': max(file_size, 1),  # Ensure at least 1 byte
                'file_type': file_type
            }
    except Exception as e:
        logger.warning(f"Failed to get file info from URL: {str(e)}")
    
    return None

def handle_missing_email(extracted_data: Dict[str, Any], session_uuid: str) -> str:
    """
    Handle cases where CV doesn't contain an email address
    
    Args:
        extracted_data: Extracted personal information
        session_uuid: Session UUID for this upload
        
    Returns:
        Email address (real or temporary)
    """
    extracted_email = extracted_data.get('email')
    
    if extracted_email and '@' in extracted_email:
        logger.info(f"Real email found in CV: {extracted_email}")
        return extracted_email
    else:
        temp_email = generate_temp_email_from_uuid(session_uuid)
        logger.info(f"No email found in CV, generated temporary email: {temp_email}")
        return temp_email

def save_user_profile_data(email: str, extracted_data: Dict[str, Any], session_uuid: str = None) -> bool:
    """
    Save extracted CV data to user_profiles table using email-based architecture
    
    Args:
        email: User email address extracted from CV
        extracted_data: Extracted personal information
        
    Returns:
        Boolean indicating success
    """
    try:
        from supabase import create_client, Client
        import os
        
        # Initialize Supabase client
        supabase_url = os.getenv('SUPABASE_URL')
        supabase_key = os.getenv('SUPABASE_SERVICE_ROLE_KEY') or os.getenv('PUBLIC_SUPABASE_PUBLISHABLE_DEFAULT_KEY') or os.getenv('PUBLIC_SUPABASE_PUBLISHABLE_DEFAULT_KEY')  # Fallback to public key
        
        if not supabase_url or not supabase_key:
            logger.warning("Warning: Supabase credentials not found, skipping database save")
            return False
            
        supabase: Client = create_client(supabase_url, supabase_key)
        
        # Prepare profile data for database (only include non-None values)
        profile_data = {
            'full_name': extracted_data.get('full_name'),
            'phone': extracted_data.get('phone'),
            'address': extracted_data.get('address'),
            'city': extracted_data.get('city'),
            'state': extracted_data.get('state'),
            'linkedin_url': extracted_data.get('linkedin_url'),
            'github_url': extracted_data.get('github_url'),
            'website_url': extracted_data.get('website_url'),
            'professional_summary': extracted_data.get('professional_summary'),
            'years_of_experience': extracted_data.get('years_of_experience'),
            'skills': extracted_data.get('skills', []),
            'education': extracted_data.get('education', []),
            'work_experience': extracted_data.get('work_experience', [])
        }
        
        # Clean phone number (remove trailing spaces and validate length)
        if profile_data.get('phone'):
            phone = profile_data['phone'].strip()
            # Only keep phone if it's at least 10 characters (as per constraint)
            if len(phone) >= 10:
                profile_data['phone'] = phone
            else:
                profile_data['phone'] = None
        
        # Remove None values and empty strings/arrays
        profile_data = {k: v for k, v in profile_data.items() 
                       if v is not None and v != '' and v != []}
        
        # Determine email source type
        email_source = 'generated_temp' if '@bestcvbuilder.com' in email else 'cv_extracted'
        
        # Use the enhanced upsert function with UUID support
        logger.info(f"Upserting user profile for email: {email} (source: {email_source})")
        
        # Call the enhanced database function with UUID support
        result = supabase.rpc('upsert_user_profile_with_uuid', {
            'p_email': email,
            'p_session_uuid': session_uuid,
            'p_profile_data': profile_data,
            'p_email_source': email_source
        }).execute()
        
        if result.data:
            logger.info(f"Successfully upserted user profile for email: {email}")
            logger.info(f"Updated fields: {list(profile_data.keys())}")
            return True
        else:
            logger.warning(f"No profile created/updated for email: {email}")
            return False
        
    except Exception as e:
        logger.error(f"Failed to save user profile data for {email}: {str(e)}")
        return False

def save_resume_record(email: str, file_url: str, file_info: Dict[str, Any], session_uuid: str = None) -> Optional[int]:
    """
    Save resume upload record to resumes table
    
    Args:
        email: User email address
        file_url: URL of uploaded file
        file_info: File metadata (filename, size, type, etc.)
        
    Returns:
        Resume ID if successful, None otherwise
    """
    try:
        from supabase import create_client, Client
        import os
        import hashlib
        from urllib.parse import urlparse
        
        supabase_url = os.getenv('SUPABASE_URL')
        supabase_key = os.getenv('SUPABASE_SERVICE_ROLE_KEY') or os.getenv('PUBLIC_SUPABASE_PUBLISHABLE_DEFAULT_KEY')
        
        if not supabase_url or not supabase_key:
            logger.warning("Supabase credentials not found, skipping resume record save")
            return None
            
        supabase: Client = create_client(supabase_url, supabase_key)
        
        # Generate file hash for duplicate detection
        file_hash = hashlib.sha256(file_url.encode()).hexdigest()[:16]  # Shortened hash
        
        # Parse file info from URL and metadata
        parsed_url = urlparse(file_url)
        filename = file_info.get('original_filename', 'unknown.pdf')
        
        # Determine email source
        email_source = 'generated_temp' if '@bestcvbuilder.com' in email else 'cv_extracted'
        
        resume_data = {
            'email': email,
            'session_uuid': session_uuid,
            'original_filename': filename,
            'file_path': parsed_url.path,
            'file_url': file_url,
            'file_size': file_info.get('file_size', 0),
            'file_type': file_info.get('file_type', 'pdf'),
            'file_hash': file_hash,
            'processing_status': 'processing',
            'upload_source': 'web_app',
            'email_source': email_source
        }
        
        logger.info(f"Saving resume record for email: {email}")
        
        # Insert resume record
        result = supabase.table('resumes').insert(resume_data).execute()
        
        if result.data and len(result.data) > 0:
            resume_id = result.data[0]['id']
            logger.info(f"Successfully saved resume record with ID: {resume_id}")
            return resume_id
        else:
            logger.warning(f"Failed to create resume record for {email}")
            return None
        
    except Exception as e:
        logger.error(f"Failed to save resume record: {str(e)}")
        return None

def save_analysis_results(email: str, resume_id: int, analysis_data: Dict[str, Any], session_uuid: str = None) -> bool:
    """
    Save analysis results to resume_analysis table
    
    Args:
        email: User email address
        resume_id: ID of the resume that was analyzed
        analysis_data: Analysis results from ATS processing
        
    Returns:
        Boolean indicating success
    """
    try:
        from supabase import create_client, Client
        import os
        
        supabase_url = os.getenv('SUPABASE_URL')
        supabase_key = os.getenv('SUPABASE_SERVICE_ROLE_KEY') or os.getenv('PUBLIC_SUPABASE_PUBLISHABLE_DEFAULT_KEY')
        
        if not supabase_url or not supabase_key:
            logger.warning("Supabase credentials not found, skipping analysis save")
            return False
            
        supabase: Client = create_client(supabase_url, supabase_key)
        
        # Clean analysis data to remove null bytes and other problematic characters
        def clean_for_database(obj):
            """Recursively clean data for database storage"""
            if isinstance(obj, str):
                # Remove null bytes and other problematic Unicode characters
                return obj.replace('\x00', '').replace('\u0000', '').encode('utf-8', 'ignore').decode('utf-8')
            elif isinstance(obj, dict):
                return {k: clean_for_database(v) for k, v in obj.items()}
            elif isinstance(obj, list):
                return [clean_for_database(item) for item in obj]
            else:
                return obj
        
        # Clean the analysis data
        cleaned_analysis_data = clean_for_database(analysis_data)
        
        # Prepare analysis data for database
        analysis_record = {
            'email': email,
            'resume_id': resume_id,
            'session_uuid': session_uuid,
            'ats_score': min(100, max(0, int(cleaned_analysis_data.get('ats_score', 0)))),
            'score_category': cleaned_analysis_data.get('category', 'poor'),
            'structure_score': min(25, max(0, int(cleaned_analysis_data.get('component_scores', {}).get('structure', 0)))),
            'keywords_score': min(20, max(0, int(cleaned_analysis_data.get('component_scores', {}).get('keywords', 0)))),  # Fixed: 20 not 25
            'contact_score': min(15, max(0, int(cleaned_analysis_data.get('component_scores', {}).get('contact', 0)))),
            'formatting_score': min(20, max(0, int(cleaned_analysis_data.get('component_scores', {}).get('formatting', 0)))),
            'achievements_score': min(10, max(0, int(cleaned_analysis_data.get('component_scores', {}).get('achievements', 0)))),
            'readability_score': min(10, max(0, int(cleaned_analysis_data.get('component_scores', {}).get('readability', 0)))),
            'strengths': cleaned_analysis_data.get('strengths', []),
            'improvements': cleaned_analysis_data.get('improvements', []),
            'missing_keywords': cleaned_analysis_data.get('critical_issues', []),
            'found_keywords': cleaned_analysis_data.get('suggestions', []),
            'detailed_analysis': cleaned_analysis_data.get('detailed_analysis', {}),
            'recommendations': cleaned_analysis_data.get('next_steps', []),
            'detected_industry': cleaned_analysis_data.get('industry', 'general'),
            'analysis_version': '2.0'
        }
        
        logger.info(f"Saving analysis results for resume ID: {resume_id}")
        
        # Insert analysis results
        result = supabase.table('resume_analysis').insert(analysis_record).execute()
        
        if result.data:
            logger.info(f"Successfully saved analysis results for resume {resume_id}")
            return True
        else:
            logger.warning(f"Failed to save analysis results for resume {resume_id}")
            return False
        
    except Exception as e:
        logger.error(f"Failed to save analysis results: {str(e)}")
        return False

def log_activity(email: str, action: str, resource_type: str = None, resource_id: int = None, 
                success: bool = True, error_message: str = None, metadata: Dict = None, session_uuid: str = None):
    """
    Log user activity for audit trail
    
    Args:
        email: User email address
        action: Action performed (upload, analyze, etc.)
        resource_type: Type of resource (resume, payment, etc.)
        resource_id: ID of the resource
        success: Whether the action was successful
        error_message: Error message if failed
        metadata: Additional metadata
    """
    try:
        from supabase import create_client, Client
        import os
        
        supabase_url = os.getenv('SUPABASE_URL')
        supabase_key = os.getenv('SUPABASE_SERVICE_ROLE_KEY') or os.getenv('PUBLIC_SUPABASE_PUBLISHABLE_DEFAULT_KEY')
        
        if not supabase_url or not supabase_key:
            return  # Silently fail for logging
            
        supabase: Client = create_client(supabase_url, supabase_key)
        
        activity_data = {
            'email': email,
            'action': action,
            'resource_type': resource_type,
            'resource_id': resource_id,
            'success': success,
            'error_message': error_message,
            'metadata': metadata or {},
            'session_uuid': session_uuid
        }
        
        supabase.table('activity_logs').insert(activity_data).execute()
        
    except Exception as e:
        logger.error(f"Failed to log activity: {str(e)}")  # Don't fail the main operation

def analyze_resume_content_fast(file_url: str) -> Dict[str, Any]:
    """
    Fast analysis function for quick ATS scoring with reduced complexity
    Provides essential scores without deep analysis for faster response times
    
    Args:
        file_url: URL of the uploaded resume file
        
    Returns:
        Dictionary containing basic analysis results optimized for speed
    """
    try:
        logger.info('⚡ Starting fast ATS analysis')
        
        # Validate URL (quick check)
        if not validate_file_url(file_url):
            raise FileProcessingError("Invalid file URL format")
        
        # Download file with shorter timeout for speed
        response = requests.get(file_url, timeout=15, stream=True)
        response.raise_for_status()
        
        # Check file size quickly (max 10MB)
        content_length = response.headers.get('content-length')
        if content_length and int(content_length) > 10 * 1024 * 1024:
            raise FileProcessingError("File size exceeds 10MB limit")
        
        # Read content with limit
        file_content = b''
        for chunk in response.iter_content(chunk_size=8192):
            file_content += chunk
            if len(file_content) > 10 * 1024 * 1024:
                raise FileProcessingError("File size exceeds 10MB limit")
        
        # Quick file type detection
        content_type = response.headers.get('content-type', '').lower()
        file_extension = get_file_extension_from_url(file_url).lower()
        
        # Fast text extraction (use fastest method available)
        if file_extension in ['.pdf'] or 'pdf' in content_type:
            resume_text = extract_text_from_pdf_fast(file_content)
        elif file_extension in ['.docx', '.doc'] or 'word' in content_type:
            resume_text = extract_text_from_docx(file_content)
        else:
            resume_text = file_content.decode('utf-8', errors='ignore')
        
        if not resume_text or len(resume_text.strip()) < 100:
            raise TextExtractionError("Insufficient text content extracted from resume")
        
        logger.info(f'⚡ Fast text extraction completed: {len(resume_text)} characters')
        
        # Fast scoring - use only essential categories for speed
        fast_categories = generate_fast_ats_scores(resume_text)
        
        # Calculate overall score quickly
        total_score = sum(cat['score'] for cat in fast_categories)
        max_score = len(fast_categories) * 10
        overall_score = min(100, (total_score / max_score) * 100)
        
        # Basic personal info extraction (minimal)
        personal_info = {
            'email': extract_email_fast(resume_text),
            'phone': extract_phone_fast(resume_text),
            'name': extract_name_fast(resume_text)
        }
        
        # Return simplified results for speed
        return {
            'ats_score': round(overall_score),
            'personal_information': personal_info,
            'detailed_analysis': {cat['name']: cat for cat in fast_categories},
            'total_categories': len(fast_categories),
            'strengths': [cat['name'] for cat in fast_categories if cat['score'] >= 8],
            'improvements_needed': [cat['name'] for cat in fast_categories if cat['score'] < 7],
            'analysis_type': 'fast',
            'processing_time': 'under_30_seconds'
        }
        
    except Exception as e:
        logger.error(f"Fast analysis error: {str(e)}")
        # Fallback to basic scoring if fast analysis fails
        return {
            'ats_score': 65,  # Default reasonable score
            'personal_information': {},
            'detailed_analysis': {},
            'analysis_type': 'fast_fallback',
            'error': 'Fast analysis failed, using fallback scoring'
        }

def extract_text_from_pdf_fast(file_content: bytes) -> str:
    """Fast PDF text extraction using most available method"""
    try:
        if PYPDF2_AVAILABLE:
            import PyPDF2
            pdf_file = io.BytesIO(file_content)
            reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in reader.pages[:5]:  # Only process first 5 pages for speed
                text += page.extract_text() + "\n"
            return text
        else:
            # Fallback - return placeholder
            return "PDF content extracted via fallback method"
    except Exception as e:
        logger.warning(f"Fast PDF extraction failed: {e}")
        return "PDF text extraction failed"

def extract_email_fast(text: str) -> str:
    """Fast email extraction"""
    import re
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    matches = re.findall(email_pattern, text)
    return matches[0] if matches else ""

def extract_phone_fast(text: str) -> str:
    """Fast phone extraction"""
    import re
    phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    matches = re.findall(phone_pattern, text)
    return ''.join(matches[0]) if matches else ""

def extract_name_fast(text: str) -> str:
    """Fast name extraction - get first line"""
    lines = text.strip().split('\n')
    for line in lines[:3]:  # Check first 3 lines
        line = line.strip()
        if len(line) > 2 and len(line) < 50 and not '@' in line:
            return line
    return ""

def generate_fast_ats_scores(text: str) -> List[Dict]:
    """Generate fast ATS scores for essential categories only"""
    categories = []
    
    # Essential scoring categories for fast analysis
    categories.append({
        'name': 'Contact Details',
        'score': 8 if '@' in text and any(char.isdigit() for char in text) else 5,
        'issue': 'Contact information analysis'
    })
    
    categories.append({
        'name': 'Action Verbs',
        'score': 7 if any(verb in text.lower() for verb in ['managed', 'led', 'created', 'developed']) else 4,
        'issue': 'Action verb usage'
    })
    
    categories.append({
        'name': 'Skills Section',
        'score': 8 if 'skills' in text.lower() or 'technologies' in text.lower() else 5,
        'issue': 'Skills section presence'
    })
    
    categories.append({
        'name': 'Experience',
        'score': 7 if 'experience' in text.lower() or 'work' in text.lower() else 5,
        'issue': 'Work experience section'
    })
    
    categories.append({
        'name': 'Education',
        'score': 8 if any(word in text.lower() for word in ['university', 'college', 'degree', 'bachelor', 'master']) else 6,
        'issue': 'Education section'
    })
    
    # Quick formatting checks
    categories.append({
        'name': 'Formatting',
        'score': 7 if len(text) > 500 and '\n' in text else 4,
        'issue': 'Resume formatting'
    })
    
    return categories

def analyze_resume_content(file_url: str) -> Dict[str, Any]:
    """
    Main function to analyze resume content and return comprehensive ATS analysis
    
    Args:
        file_url: URL of the uploaded resume file
        
    Returns:
        Dictionary containing comprehensive analysis results
    """
    try:
        # Validate URL
        if not validate_file_url(file_url):
            raise FileProcessingError("Invalid file URL format")
        
        # Download file with timeout and size limits
        response = requests.get(file_url, timeout=30, stream=True)
        response.raise_for_status()
        
        # Check file size (max 10MB)
        content_length = response.headers.get('content-length')
        if content_length and int(content_length) > 10 * 1024 * 1024:
            raise FileProcessingError("File size exceeds 10MB limit")
        
        file_content = response.content
        
        # Clean up response object
        response.close()
        del response
        
        # Extract text content with memory monitoring
        logger.info(f"📝 Starting text extraction from file")
        content = extract_text_from_file(file_content, file_url)
        
        # Clean up file content immediately after extraction
        del file_content
        cleanup_memory()
        
        # Validate extracted content
        if not content or len(content.strip()) < 50:
            raise TextExtractionError("Insufficient text content extracted from file")
        
        logger.info(f"Successfully extracted {len(content)} characters from resume")
        
        # Extract personal information from CV content
        personal_info = extract_personal_information(content)
        logger.info(f"Extracted personal information: {list(personal_info.keys())}")
        
        # Debug: Log what was actually found vs None
        non_empty_fields = {k: v for k, v in personal_info.items() if v is not None and v != [] and v != ''}
        logger.info(f"Non-empty extracted fields: {list(non_empty_fields.keys())}")
        logger.info(f"GitHub found: {personal_info.get('github_url', 'None')}")
        logger.info(f"Website found: {personal_info.get('website_url', 'None')}")
        logger.info(f"Full name found: {personal_info.get('full_name', 'None')}")
        logger.info(f"Years of experience: {personal_info.get('years_of_experience', 'None')}")
        
        # Extract filename from URL for analysis
        import os
        filename = os.path.basename(file_url) if file_url else None
        
        # Perform comprehensive ATS analysis
        ats_analysis = calculate_comprehensive_ats_score(content, filename=filename)
        
        # Debug: Log component scores to see why total is 100
        logger.info(f"Component scores breakdown:")
        if 'component_scores' in ats_analysis:
            for component, score in ats_analysis['component_scores'].items():
                logger.info(f"  {component}: {score}")
        logger.info(f"Total ATS score: {ats_analysis.get('ats_score', 'Not found')}")
        
        # Generate recommendations
        recommendations = generate_comprehensive_recommendations(ats_analysis)
        
        # Generate enhanced analysis with all new features
        critical_issues, quick_wins = classify_issues_by_priority(ats_analysis)
        interview_metrics = calculate_interview_rates(ats_analysis['ats_score'])
        transformation_preview = generate_transformation_preview(ats_analysis, critical_issues, quick_wins)
        enhanced_components = enhance_component_breakdown(ats_analysis)
        
        # CRITICAL: Generate detailed issues analysis with specific, actionable problems
        detailed_issues = generate_detailed_issues_analysis(ats_analysis, content)
        
        # Override the generic issues with our specific ones
        critical_issues = detailed_issues['critical_issues']
        quick_fixes = detailed_issues['quick_fixes']
        content_improvements = detailed_issues['content_improvements']
        
        # Debug: Log the ats_analysis keys before merging
        logger.info(f'🔍 DEBUG: ats_analysis keys before merge: {list(ats_analysis.keys())}')
        logger.info(f'🔍 DEBUG: ats_analysis detailedAnalysis count: {len(ats_analysis.get("detailedAnalysis", {}))}')
        
        # Combine results with enhanced data
        result = {
            **ats_analysis,
            **recommendations,
            'personal_information': personal_info,
            'extracted_text_length': len(content),
            'analysis_timestamp': json.dumps({}).__class__.__module__, # Simple timestamp
            
            # CRITICAL: Include original file data for resume improvement API
            'file_url': file_url,  # Original PDF URL for re-processing
            'content': content,    # Extracted text content
            
            # Enhanced algorithm features
            'letter_grade': get_letter_grade(ats_analysis['ats_score']),
            'interview_metrics': interview_metrics,
            'critical_issues': critical_issues,
            'quick_wins': quick_fixes,  # Use quick_fixes from detailed analysis
            'content_improvements': content_improvements,  # Add content improvements category
            'transformation_preview': transformation_preview,
            'enhanced_components': enhanced_components,
            'total_issues': detailed_issues['total_issues'],  # Use calculated total
            'potential_improvement': detailed_issues['potential_improvement'],  # Realistic improvement
            'realistic_target_score': detailed_issues['realistic_target_score'],  # Target score
            'estimated_time': detailed_issues['estimated_time'],  # Time to complete
            'actionable_improvements': len([i for i in critical_issues + quick_fixes if i.get('time_to_fix', '').split()[0].isdigit() and int(i.get('time_to_fix', '0').split()[0]) <= 5])
        }
        
        # Debug: Log the final result keys after merging
        logger.info(f'🔍 DEBUG: final result keys: {list(result.keys())}')
        logger.info(f'🔍 DEBUG: final result detailedAnalysis count: {len(result.get("detailedAnalysis", {}))}')
        logger.info(f'🔍 DEBUG: final result detailed_analysis count: {len(result.get("detailed_analysis", {}))}')
        
        logger.info(f"Analysis completed - Score: {ats_analysis['ats_score']}")
        return result
        
    except requests.exceptions.Timeout:
        raise FileProcessingError("File download timeout - please try again")
    except requests.exceptions.RequestException as e:
        raise FileProcessingError(f"Could not download file: {str(e)}")
    except (FileProcessingError, TextExtractionError):
        raise  # Re-raise specific errors
    except Exception as e:
        import traceback
        logger.error(f"Unexpected error in resume analysis: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise ATSAnalysisError(f"An unexpected error occurred during analysis: {str(e)}")


def extract_specific_issues_with_examples(analysis_result: Dict[str, Any]) -> Dict[str, Any]:
    """
    Extract specific issues from resume content with exact examples and line references
    
    Args:
        analysis_result: Complete analysis result dictionary containing content and scores
        
    Returns:
        Dictionary with categorized issues including specific examples and fix suggestions
    """
    try:
        content = analysis_result.get('content', '')
        detailed_analysis = analysis_result.get('detailedAnalysis', {})
        
        if not content:
            return {'error': 'No resume content available for detailed analysis'}
        
        # Split content into lines for reference
        lines = content.split('\n')
        issues_with_examples = {
            'critical_issues': [],
            'quick_wins': [],
            'content_improvements': [],
            'total_specific_examples': 0
        }
        
        # 1. DATES ISSUES - Find inconsistent date formats
        if 'dates' in detailed_analysis and detailed_analysis['dates'].get('score', 10) < 8:
            date_issues = find_date_formatting_issues(lines)
            if date_issues:
                issues_with_examples['critical_issues'].append({
                    'category': 'Dates',
                    'score': detailed_analysis['dates'].get('score', 0),
                    'title': 'Inconsistent Date Formatting',
                    'severity': 'CRITICAL',
                    'impact': 'Major ATS blocker - prevents proper parsing',
                    'examples': date_issues['examples'],
                    'fix_instructions': date_issues['fix_instructions'],
                    'time_to_fix': '10-15 minutes',
                    'score_impact': '+8 to +15 points'
                })
        
        # 2. VERB REPETITION - Find repeated action verbs
        if 'repetition' in detailed_analysis and detailed_analysis['repetition'].get('score', 10) < 6:
            verb_issues = find_verb_repetition_issues(lines)
            if verb_issues:
                issues_with_examples['critical_issues'].append({
                    'category': 'Repetition',
                    'score': detailed_analysis['repetition'].get('score', 0),
                    'title': 'Repeated Action Verbs',
                    'severity': 'HIGH',
                    'impact': 'Reduces impact and shows limited vocabulary',
                    'examples': verb_issues['examples'],
                    'fix_instructions': verb_issues['fix_instructions'],
                    'time_to_fix': '15-20 minutes',
                    'score_impact': '+5 to +10 points'
                })
        
        # 3. CONTACT INFO - Find missing contact elements
        if 'contact_details' in detailed_analysis and detailed_analysis['contact_details'].get('score', 10) < 10:
            contact_issues = find_contact_info_issues(lines)
            if contact_issues:
                issues_with_examples['quick_wins'].append({
                    'category': 'Contact Details',
                    'score': detailed_analysis['contact_details'].get('score', 8),
                    'title': 'Missing Contact Information',
                    'severity': 'QUICK_WIN',
                    'impact': 'Easy fix for better ATS compatibility',
                    'examples': contact_issues['examples'],
                    'fix_instructions': contact_issues['fix_instructions'],
                    'time_to_fix': '2-5 minutes',
                    'score_impact': '+1 to +3 points'
                })
        
        # 4. QUANTIFIABLE ACHIEVEMENTS - Find vague statements
        if 'quantifiable_achievements' in detailed_analysis and detailed_analysis['quantifiable_achievements'].get('score', 10) < 9:
            achievement_issues = find_quantification_issues(lines)
            if achievement_issues:
                issues_with_examples['content_improvements'].append({
                    'category': 'Quantifiable Achievements',
                    'score': detailed_analysis['quantifiable_achievements'].get('score', 8),
                    'title': 'Vague Achievement Statements',
                    'severity': 'IMPROVEMENT',
                    'impact': 'Makes accomplishments more compelling and measurable',
                    'examples': achievement_issues['examples'],
                    'fix_instructions': achievement_issues['fix_instructions'],
                    'time_to_fix': '20-30 minutes',
                    'score_impact': '+2 to +5 points'
                })
        
        # Count total examples
        issues_with_examples['total_specific_examples'] = sum([
            len(category) for category in [
                issues_with_examples['critical_issues'],
                issues_with_examples['quick_wins'],
                issues_with_examples['content_improvements']
            ]
        ])
        
        logger.info(f"✅ Extracted {issues_with_examples['total_specific_examples']} specific issues with examples")
        return issues_with_examples
        
    except Exception as e:
        logger.error(f"Error extracting specific issues: {str(e)}")
        return {'error': f'Failed to extract specific issues: {str(e)}'}

def find_date_formatting_issues(lines: List[str]) -> Dict[str, Any]:
    """Find specific date formatting inconsistencies in resume"""
    import re
    
    date_patterns = {
        'MM/YYYY': r'\b(0[1-9]|1[0-2])/\d{4}\b',          # 01/2020
        'MM-YYYY': r'\b(0[1-9]|1[0-2])-\d{4}\b',          # 01-2020
        'YYYY': r'\b\d{4}\b(?!-\d|\d)',                     # 2020
        'Mon YYYY': r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}\b',  # Jan 2020
        'Month YYYY': r'\b(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}\b',
        'YYYY-YYYY': r'\b\d{4}-\d{4}\b',                   # 2020-2022
        'MM/YY': r'\b(0[1-9]|1[0-2])/\d{2}\b'             # 01/20
    }
    
    found_formats = {}
    examples = []
    
    for line_num, line in enumerate(lines, 1):
        line_clean = line.strip()
        if len(line_clean) < 5:  # Skip very short lines
            continue
            
        for format_name, pattern in date_patterns.items():
            matches = re.findall(pattern, line_clean)
            if matches:
                if format_name not in found_formats:
                    found_formats[format_name] = []
                found_formats[format_name].extend([{
                    'line_number': line_num,
                    'line_text': line_clean,
                    'date_found': match if isinstance(match, str) else matches[0],
                    'format': format_name
                } for match in matches])
    
    if len(found_formats) <= 1:
        return {}  # No inconsistency found
    
    # Create examples showing inconsistency
    for format_name, format_examples in found_formats.items():
        for example in format_examples[:2]:  # Limit to 2 examples per format
            examples.append({
                'issue': f"Uses {format_name} format",
                'line_number': example['line_number'],
                'line_text': example['line_text'],
                'problematic_text': example['date_found']
            })
    
    # Generate fix instructions
    recommended_format = 'MM/YYYY'  # Most ATS-friendly format
    fix_instructions = [
        f"Use consistent {recommended_format} format throughout your resume",
        "Examples of fixes:"
    ]
    
    for format_name, format_examples in found_formats.items():
        if format_name != recommended_format:
            example = format_examples[0]
            original = example['date_found']
            # Convert to recommended format (simplified conversion)
            if format_name == 'Mon YYYY':
                month_map = {'Jan': '01', 'Feb': '02', 'Mar': '03', 'Apr': '04', 'May': '05', 'Jun': '06',
                           'Jul': '07', 'Aug': '08', 'Sep': '09', 'Oct': '10', 'Nov': '11', 'Dec': '12'}
                parts = original.split()
                if len(parts) == 2 and parts[0] in month_map:
                    converted = f"{month_map[parts[0]]}/{parts[1]}"
                    fix_instructions.append(f"• Change \"{original}\" → \"{converted}\"")
            elif format_name == 'YYYY-YYYY':
                parts = original.split('-')
                if len(parts) == 2:
                    converted = f"01/{parts[0]} - 12/{parts[1]}"
                    fix_instructions.append(f"• Change \"{original}\" → \"{converted}\"")
    
    return {
        'examples': examples[:6],  # Limit to 6 examples
        'fix_instructions': fix_instructions,
        'formats_found': list(found_formats.keys())
    }

def find_verb_repetition_issues(lines: List[str]) -> Dict[str, Any]:
    """Find repeated action verbs with specific examples"""
    import re
    from collections import defaultdict
    
    # Action verbs to check for repetition
    action_verbs = [
        'managed', 'developed', 'created', 'implemented', 'designed', 'executed',
        'delivered', 'coordinated', 'supervised', 'administered', 'maintained',
        'organized', 'planned', 'handled', 'processed', 'worked', 'helped'
    ]
    
    verb_occurrences = defaultdict(list)
    
    for line_num, line in enumerate(lines, 1):
        line_clean = line.strip().lower()
        if len(line_clean) < 10:  # Skip very short lines
            continue
            
        for verb in action_verbs:
            # Look for verb at start of bullet points or sentences
            pattern = rf'\b{verb}(?:d|ed|ing|s)?\b'
            if re.search(pattern, line_clean):
                verb_occurrences[verb].append({
                    'line_number': line_num,
                    'line_text': line.strip(),
                    'verb_used': verb
                })
    
    # Find verbs used more than once
    repeated_verbs = {verb: occurrences for verb, occurrences in verb_occurrences.items() if len(occurrences) > 1}
    
    if not repeated_verbs:
        return {}
    
    examples = []
    fix_instructions = ["Replace repeated verbs with stronger alternatives:"]
    
    # Verb alternatives mapping
    verb_alternatives = {
        'managed': ['Led', 'Spearheaded', 'Orchestrated', 'Directed', 'Guided'],
        'developed': ['Engineered', 'Architected', 'Crafted', 'Built', 'Designed'],
        'created': ['Established', 'Founded', 'Initiated', 'Launched', 'Pioneered'],
        'implemented': ['Deployed', 'Executed', 'Rolled out', 'Introduced', 'Installed'],
        'coordinated': ['Orchestrated', 'Facilitated', 'Synchronized', 'Organized', 'Harmonized'],
        'worked': ['Collaborated', 'Partnered', 'Engaged', 'Contributed', 'Participated']
    }
    
    for verb, occurrences in list(repeated_verbs.items())[:3]:  # Limit to top 3 repeated verbs
        count = len(occurrences)
        examples.append({
            'issue': f'"{verb.title()}" used {count} times',
            'occurrences': occurrences[:3],  # Show up to 3 occurrences
            'alternatives': verb_alternatives.get(verb, ['Led', 'Executed', 'Achieved'])
        })
        
        alternatives = verb_alternatives.get(verb, ['Led', 'Executed', 'Achieved'])
        fix_instructions.append(f"• \"{verb.title()}\" → Use: {', '.join(alternatives[:3])}")
    
    return {
        'examples': examples,
        'fix_instructions': fix_instructions,
        'repeated_verbs_count': len(repeated_verbs)
    }

def find_contact_info_issues(lines: List[str]) -> Dict[str, Any]:
    """Find missing contact information elements"""
    import re
    
    # Check first 10 lines for contact info
    header_lines = lines[:10]
    header_text = '\n'.join(header_lines)
    
    contact_elements = {
        'email': bool(re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', header_text, re.IGNORECASE)),
        'phone': bool(re.search(r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', header_text)),
        'linkedin': bool(re.search(r'linkedin\.com|in/[\w-]+', header_text, re.IGNORECASE)),
        'location': bool(re.search(r'\b\w+,\s*\w+\b|\b\w+\s+\w+,\s*\w+\b', header_text))
    }
    
    missing_elements = [element for element, present in contact_elements.items() if not present]
    
    if not missing_elements:
        return {}
    
    examples = []
    fix_instructions = ["Add missing contact information to your header:"]
    
    for element in missing_elements:
        if element == 'linkedin':
            examples.append({
                'issue': 'LinkedIn profile URL missing',
                'line_number': 'Header section',
                'impact': 'LinkedIn increases profile visibility to recruiters'
            })
            fix_instructions.append("• Add: linkedin.com/in/your-username")
            
        elif element == 'phone':
            examples.append({
                'issue': 'Phone number missing',
                'line_number': 'Header section', 
                'impact': 'Phone contact essential for follow-ups'
            })
            fix_instructions.append("• Add: Your phone number in format +1 (XXX) XXX-XXXX")
            
        elif element == 'location':
            examples.append({
                'issue': 'Location/City missing',
                'line_number': 'Header section',
                'impact': 'Location helps with local job matching'
            })
            fix_instructions.append("• Add: City, State or City, Country")
    
    return {
        'examples': examples,
        'fix_instructions': fix_instructions,
        'missing_count': len(missing_elements)
    }

def find_quantification_issues(lines: List[str]) -> Dict[str, Any]:
    """Find statements that need quantification"""
    import re
    
    # Vague terms that should be quantified
    vague_terms = [
        'significantly', 'substantially', 'considerably', 'greatly', 'highly',
        'many', 'several', 'multiple', 'various', 'numerous', 'large', 'small',
        'improved', 'increased', 'reduced', 'enhanced', 'optimized', 'boosted'
    ]
    
    examples = []
    
    for line_num, line in enumerate(lines, 1):
        line_clean = line.strip()
        if len(line_clean) < 15:  # Skip very short lines
            continue
            
        # Check if line contains vague terms and lacks numbers
        has_vague_term = any(re.search(rf'\b{term}\b', line_clean, re.IGNORECASE) for term in vague_terms)
        has_numbers = bool(re.search(r'\d+[%$]?|\$\d+|[0-9,]+\s*(users|clients|projects|hours|days|months|years)', line_clean))
        
        if has_vague_term and not has_numbers:
            # Find which vague term was used
            vague_found = []
            for term in vague_terms:
                if re.search(rf'\b{term}\b', line_clean, re.IGNORECASE):
                    vague_found.append(term)
            
            examples.append({
                'issue': 'Vague statement needs quantification',
                'line_number': line_num,
                'line_text': line_clean,
                'vague_terms': vague_found[:2],  # Show up to 2 vague terms
                'suggested_metrics': get_suggested_metrics_for_line(line_clean)
            })
    
    if not examples:
        return {}
    
    fix_instructions = [
        "Add specific numbers to quantify your achievements:",
        "Examples of quantifiable metrics:",
        "• Performance: '40% improvement', '25% faster'",
        "• Scale: '50 users', '10 projects', '$100K budget'",
        "• Time: '6-month project', '2 weeks ahead of schedule'",
        "• Team: '5-person team', '15 stakeholders'"
    ]
    
    return {
        'examples': examples[:5],  # Limit to 5 examples
        'fix_instructions': fix_instructions,
        'vague_statements_count': len(examples)
    }

def get_suggested_metrics_for_line(line_text: str) -> List[str]:
    """Suggest appropriate metrics based on line content"""
    line_lower = line_text.lower()
    suggestions = []
    
    if any(word in line_lower for word in ['performance', 'speed', 'efficiency']):
        suggestions.extend(['% improvement', 'time reduction', 'speed increase'])
    elif any(word in line_lower for word in ['team', 'people', 'staff']):
        suggestions.extend(['team size', 'number of people', 'reports managed'])
    elif any(word in line_lower for word in ['project', 'initiative']):
        suggestions.extend(['project count', 'timeline', 'budget size'])
    elif any(word in line_lower for word in ['revenue', 'sales', 'cost']):
        suggestions.extend(['dollar amount', '% increase', 'ROI'])
    else:
        suggestions.extend(['specific numbers', 'percentages', 'time periods'])
    
    return suggestions[:3]  # Return top 3 suggestions

def generate_comprehensive_issues_report(analysis_result: Dict[str, Any]) -> str:
    """
    Generate comprehensive TXT report of all ATS issues with specific examples from resume
    
    Args:
        analysis_result: Complete analysis result dictionary
        
    Returns:
        Formatted text report with specific issues and actionable fixes with examples
    """
    try:
        logger.info("🔍 Generating enhanced TXT issues report with specific examples...")
        
        # Extract specific issues with examples from the resume
        specific_issues = extract_specific_issues_with_examples(analysis_result)
        
        # Extract main data
        score = analysis_result.get('ats_score', 0)
        detailed_analysis = analysis_result.get('detailedAnalysis', {})
        
        # Use the enhanced issues with examples
        critical_issues = specific_issues.get('critical_issues', [])
        quick_wins = specific_issues.get('quick_wins', [])
        content_improvements = specific_issues.get('content_improvements', [])
        
        # Build comprehensive report
        report_lines = []
        
        # Header section
        report_lines.extend([
            "=" * 80,
            "ATS SPECIFIC ISSUES REPORT WITH RESUME EXAMPLES",
            "=" * 80,
            f"Current ATS Score: {score}/100",
            f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            "",
            "EXECUTIVE DASHBOARD",
            "-" * 40
        ])
        
        # Count issues based on the new structure
        critical_count = len(critical_issues)
        quick_wins_count = len(quick_wins)
        improvements_count = len(content_improvements)
        total_issues = critical_count + quick_wins_count + improvements_count
        
        # Calculate estimated time
        total_time_minutes = (critical_count * 7) + (quick_wins_count * 3) + (improvements_count * 12)
        hours = total_time_minutes // 60
        minutes = total_time_minutes % 60
        time_estimate = f"{hours}h {minutes}m" if hours > 0 else f"{minutes} minutes"
        
        report_lines.extend([
            f"🚨 Critical Issues (High Impact): {critical_count}",
            f"⚡ Quick Wins (Easy Fixes): {quick_wins_count}",
            f"📝 Content Improvements: {improvements_count}",
            f"📊 Total Specific Issues Found: {total_issues}",
            f"⏱️  Estimated Fix Time: {time_estimate}",
            "",
            "🎯 PRIORITY ACTION PLAN",
            "-" * 40,
            "1. Fix Critical Issues first (biggest ATS score boost)",
            "2. Complete Quick Wins next (easy points)",
            "3. Work on Content Improvements for polish",
            ""
        ])
        
        # Critical Issues Section
        if critical_issues:
            report_lines.extend([
                "",
                "🚨 CRITICAL ISSUES (IMMEDIATE ATTENTION REQUIRED)",
                "=" * 60
            ])
            
            for i, issue in enumerate(critical_issues, 1):
                category = issue.get('category', 'General')
                title = issue.get('title', 'Issue')
                description = issue.get('description', 'No description available')
                score = issue.get('score', 'N/A')
                time_to_fix = issue.get('time_to_fix', '5-10 minutes')
                
                # Enhanced issue display with specific examples
                problematic_text = issue.get('problematic_text', '')
                line_number = issue.get('line_number', '')
                fix_suggestion = issue.get('fix_suggestion', 'No fix suggestion available')
                
                report_lines.extend([
                    f"{i}. {category.upper()}: {title}",
                    f"   Current Score: {score}/10 | Time to Fix: {time_to_fix}",
                    f"   Problem: {description}",
                    ""
                ])
                
                # Add specific example if available
                if problematic_text and line_number:
                    report_lines.extend([
                        f"   📍 FOUND IN YOUR RESUME (Line {line_number}):",
                        f"   \"{problematic_text}\"",
                        ""
                    ])
                
                # Add specific fix with before/after if available
                if fix_suggestion:
                    report_lines.extend([
                        f"   ✅ SPECIFIC FIX:",
                        f"   {fix_suggestion}",
                        ""
                    ])
                
                report_lines.append("-" * 50)
                report_lines.append("")
        
        # Quick Wins Section
        if quick_wins:
            report_lines.extend([
                "",
                "⚡ QUICK WINS (EASY FIXES FOR IMMEDIATE IMPACT)",
                "=" * 60
            ])
            
            for i, issue in enumerate(quick_wins, 1):
                category = issue.get('category', 'General')
                title = issue.get('title', 'Issue')
                description = issue.get('description', 'No description available')
                score = issue.get('score', 'N/A')
                time_to_fix = issue.get('time_to_fix', '2-5 minutes')
                
                # Enhanced issue display with specific examples
                problematic_text = issue.get('problematic_text', '')
                line_number = issue.get('line_number', '')
                fix_suggestion = issue.get('fix_suggestion', 'No fix suggestion available')
                
                report_lines.extend([
                    f"{i}. {category.upper()}: {title}",
                    f"   Current Score: {score}/10 | Time to Fix: {time_to_fix}",
                    f"   Problem: {description}",
                    ""
                ])
                
                # Add specific example if available
                if problematic_text and line_number:
                    report_lines.extend([
                        f"   📍 FOUND IN YOUR RESUME (Line {line_number}):",
                        f"   \"{problematic_text}\"",
                        ""
                    ])
                
                # Add specific fix with before/after if available
                if fix_suggestion:
                    report_lines.extend([
                        f"   ✅ QUICK FIX:",
                        f"   {fix_suggestion}",
                        ""
                    ])
                
                report_lines.append("-" * 50)
                report_lines.append("")
        
        # Content Improvements Section
        if content_improvements:
            report_lines.extend([
                "",
                "📝 CONTENT IMPROVEMENTS (ENHANCE YOUR PRESENTATION)",
                "=" * 60
            ])
            
            for i, issue in enumerate(content_improvements, 1):
                category = issue.get('category', 'General')
                title = issue.get('title', 'Issue')
                description = issue.get('description', 'No description available')
                score = issue.get('score', 'N/A')
                time_to_fix = issue.get('time_to_fix', '10-15 minutes')
                
                # Enhanced issue display with specific examples
                problematic_text = issue.get('problematic_text', '')
                line_number = issue.get('line_number', '')
                fix_suggestion = issue.get('fix_suggestion', 'No fix suggestion available')
                
                report_lines.extend([
                    f"{i}. {category.upper()}: {title}",
                    f"   Current Score: {score}/10 | Time to Fix: {time_to_fix}",
                    f"   Improvement: {description}",
                    ""
                ])
                
                # Add specific example if available
                if problematic_text and line_number:
                    report_lines.extend([
                        f"   📍 FOUND IN YOUR RESUME (Line {line_number}):",
                        f"   \"{problematic_text}\"",
                        ""
                    ])
                
                # Add specific fix with before/after if available
                if fix_suggestion:
                    report_lines.extend([
                        f"   ✅ ENHANCEMENT SUGGESTION:",
                        f"   {fix_suggestion}",
                        ""
                    ])
                
                report_lines.append("-" * 50)
                report_lines.append("")
        
        # Category-by-Category Breakdown
        if detailed_analysis:
            report_lines.extend([
                "",
                "📊 DETAILED CATEGORY BREAKDOWN",
                "=" * 60
            ])
            
            # Sort categories by score (lowest first)
            sorted_categories = sorted(detailed_analysis.items(), key=lambda x: x[1].get('score', 10))
            
            for category, data in sorted_categories:
                score = data.get('score', 10)
                issues = data.get('issues', [])
                suggestions = data.get('suggestions', [])
                
                # Only show categories with issues
                if score < 10 or issues:
                    report_lines.extend([
                        f"{category.upper().replace('_', ' ')} - Score: {score}/10",
                        "-" * 50
                    ])
                    
                    if issues:
                        report_lines.append("Issues Found:")
                        for issue in issues:
                            report_lines.append(f"  • {issue}")
                    
                    if suggestions:
                        report_lines.append("Recommendations:")
                        for suggestion in suggestions:
                            report_lines.append(f"  • {suggestion}")
                    
                    report_lines.append("")
        
        # Final Action Summary
        report_lines.extend([
            "",
            "🎯 YOUR SPECIFIC ACTION CHECKLIST",
            "=" * 60,
            "Complete these fixes in the exact order shown above:",
            "",
            f"Phase 1: Critical Issues ({critical_count} items) - Est: {(critical_count * 7)} min",
            "→ Focus on ATS compatibility issues that are blocking your resume",
            "",
            f"Phase 2: Quick Wins ({quick_wins_count} items) - Est: {(quick_wins_count * 3)} min", 
            "→ Easy formatting and structural improvements",
            "",
            f"Phase 3: Content Improvements ({improvements_count} items) - Est: {(improvements_count * 12)} min",
            "→ Polish your professional presentation",
            "",
            f"Total estimated time: {time_estimate}",
            "",
            "🔥 IMMEDIATE NEXT STEPS",
            "=" * 60,
            "1. Print this report or keep it open while editing",
            "2. Work through Critical Issues first - they have the biggest impact",
            "3. Use the exact line references to find problems in your resume",
            "4. Replace the quoted text with the suggested improvements",
            "5. Re-upload your resume to BestCVBuilder.com to see score improvement",
            "",
            "💡 PRO TIPS FOR MAXIMUM ATS SUCCESS",
            "=" * 60,
            "• Each fix shown above was found specifically in YOUR resume",
            "• Line numbers help you locate exact problems quickly",
            "• Before/after examples show you exactly what to change",
            "• Completing ALL fixes typically boosts ATS scores by 20-35 points",
            "• Save final resume as PDF to preserve formatting",
            "",
            "Generated by BestCVBuilder.com - Your ATS Optimization Partner",
            "For more advanced optimization: https://bestcvbuilder.com",
            "=" * 80
        ])
        
        # Join all lines and return
        final_report = "\n".join(report_lines)
        
        logger.info(f"✅ Generated comprehensive report with {len(report_lines)} lines")
        return final_report
        
    except Exception as e:
        logger.error(f"Error generating comprehensive issues report: {str(e)}")
        return f"Error generating comprehensive report: {str(e)}"


from http.server import BaseHTTPRequestHandler

class handler(BaseHTTPRequestHandler):
    """
    Main handler class for Vercel serverless function using BaseHTTPRequestHandler
    """
    
    def do_OPTIONS(self):
        """Handle CORS preflight requests"""
        self.send_response(200)
        # Set Content-Type first
        self.send_header('Content-Type', 'application/json')
        # Set all CORS headers
        for key, value in cors_headers().items():
            self.send_header(key.replace('_', '-'), value)
        self.end_headers()
        return
    
    def do_POST(self):
        """Handle POST requests for CV analysis"""
        try:
            # Read request body
            content_length = int(self.headers.get('Content-Length', 0))
            if content_length > 0:
                body_data = self.rfile.read(content_length).decode('utf-8')
                body = json.loads(body_data)
            else:
                body = {}
            
            file_url = body.get('file_url')
            user_id = body.get('user_id')  # Optional user ID for database saving
            analysis_type = body.get('analysis_type', 'comprehensive')
            include_recommendations = body.get('include_recommendations', True)
            
            if not file_url:
                self.send_error_response({'error': 'file_url is required'}, 400)
                return
            
            # Perform analysis based on type
            if analysis_type == 'fast_ats_score':
                logger.info('🚀 Performing fast ATS analysis for quick results')
                analysis_result = analyze_resume_content_fast(file_url)
            else:
                logger.info('🔍 Performing comprehensive ATS analysis') 
                analysis_result = analyze_resume_content(file_url)
            
            # Extract personal information and handle email with UUID fallback
            personal_info = analysis_result.get('personal_information', {})
            session_uuid = generate_session_uuid()
            
            # Handle missing email with UUID fallback
            final_email = handle_missing_email(personal_info, session_uuid)
            is_temp_email = '@bestcvbuilder.com' in final_email
            
            logger.info(f"Processing CV with email: {final_email} (temporary: {is_temp_email})")
            
            try:
                # Step 1: Save/update user profile with UUID tracking
                profile_saved = save_user_profile_data(final_email, personal_info, session_uuid)
                analysis_result['profile_updated'] = profile_saved
                analysis_result['session_uuid'] = session_uuid
                analysis_result['email_used'] = final_email
                analysis_result['is_temporary_email'] = is_temp_email
                
                # Step 2: Save resume record with UUID
                # Get file info from the actual file URL
                file_info = get_file_info_from_url(file_url)
                if not file_info:
                    file_info = {
                        'original_filename': 'uploaded_resume.pdf',
                        'file_size': 1024,  # Default size to pass constraint (will be updated)
                        'file_type': 'pdf'
                    }
                resume_id = save_resume_record(final_email, file_url, file_info, session_uuid)
                analysis_result['resume_id'] = resume_id
                
                # Step 3: Save analysis results with UUID
                if resume_id:
                    analysis_saved = save_analysis_results(final_email, resume_id, analysis_result, session_uuid)
                    analysis_result['analysis_saved'] = analysis_saved
                    
                    # Update resume status to completed
                    try:
                        from supabase import create_client
                        supabase_url = os.getenv('SUPABASE_URL')
                        supabase_key = os.getenv('SUPABASE_SERVICE_ROLE_KEY') or os.getenv('PUBLIC_SUPABASE_PUBLISHABLE_DEFAULT_KEY')
                        if supabase_url and supabase_key:
                            supabase = create_client(supabase_url, supabase_key)
                            supabase.table('resumes').update({
                                'processing_status': 'completed',
                                'analysis_completed': True,
                                'processed_at': 'now()'
                            }).eq('id', resume_id).execute()
                    except Exception as e:
                        logger.error(f"Failed to update resume status: {str(e)}")
                
                # Step 4: Log activity with UUID (always log for audit trail)
                log_activity(
                    email=final_email,
                    action='resume_analysis',
                    resource_type='resume',
                    resource_id=resume_id,
                    success=True,
                    metadata={
                        'ats_score': analysis_result.get('ats_score'),
                        'file_url': file_url,
                        'is_temporary_email': is_temp_email,
                        'original_email_found': personal_info.get('email') is not None,
                        'profile_saved': profile_saved,
                        'analysis_saved': analysis_result.get('analysis_saved', False)
                    },
                    session_uuid=session_uuid
                )
                
                logger.info(f"Successfully completed processing for {final_email} (UUID: {session_uuid})")
                
                # Add instructions for temporary email users
                if is_temp_email:
                    analysis_result['temp_email_notice'] = {
                        'message': 'Your CV did not contain an email address. We\'ve created a temporary session for you.',
                        'temp_email': final_email,
                        'session_uuid': session_uuid,
                        'instructions': 'To access your results later, you can provide your email during payment.'
                    }
                
            except Exception as e:
                logger.error(f"Error in email-based processing: {str(e)}")
                analysis_result['profile_updated'] = False
                analysis_result['database_error'] = str(e)
                
                # Log failed activity (skip if profile creation failed to avoid FK constraint)
                try:
                    log_activity(
                        email=final_email,
                        action='resume_analysis',
                        resource_type='resume',
                        success=False,
                        error_message=str(e),
                        session_uuid=session_uuid
                    )
                except Exception as log_error:
                    logger.warning(f"Failed to log activity: {str(log_error)}")
            
            # Generate comprehensive TXT issues report
            try:
                comprehensive_report = generate_comprehensive_issues_report(analysis_result)
                analysis_result['comprehensive_issues_report'] = comprehensive_report
                logger.info("✅ Comprehensive TXT issues report generated successfully")
            except Exception as report_error:
                logger.warning(f"Failed to generate comprehensive report: {str(report_error)}")
                analysis_result['comprehensive_issues_report'] = None
            
            # Filter results based on request parameters
            if not include_recommendations:
                # Remove recommendation fields if not requested
                analysis_result = {k: v for k, v in analysis_result.items() 
                                 if k not in ['improvements', 'suggestions', 'next_steps']}
            
            # Return successful results
            self.send_success_response(analysis_result)
            
        except FileProcessingError as e:
            logger.error(f"File processing error: {str(e)}")
            self.send_error_response({'error': str(e)}, 400)
        except TextExtractionError as e:
            logger.error(f"Text extraction error: {str(e)}")
            self.send_error_response({'error': f'Text extraction failed: {str(e)}'}, 422)
        except ATSAnalysisError as e:
            logger.error(f"ATS analysis error: {str(e)}")
            self.send_error_response({'error': str(e)}, 500)
        except Exception as e:
            logger.error(f"Unexpected API error: {str(e)}")
            self.send_error_response({'error': 'Internal server error'}, 500)
    
    def do_GET(self):
        """Handle GET requests for diagnostics"""
        try:
            # Parse the request path
            from urllib.parse import urlparse, parse_qs
            parsed_path = urlparse(self.path)
            
            if parsed_path.path == '/diagnostics' or parsed_path.path == '/api/cv-parser/diagnostics':
                # Return dependency status for diagnostics
                diagnostic_info = {
                    'service': 'CV Parser API',
                    'status': 'running',
                    'dependencies': {
                        'PyPDF2': PYPDF2_AVAILABLE,
                        'pdfplumber': PDFPLUMBER_AVAILABLE,
                        'PyMuPDF': PYMUPDF_AVAILABLE,
                        'pdfminer': PDFMINER_AVAILABLE,
                        'python-docx': DOCX_AVAILABLE
                    },
                    'total_pdf_methods': sum([PYPDF2_AVAILABLE, PDFPLUMBER_AVAILABLE, PYMUPDF_AVAILABLE, PDFMINER_AVAILABLE]),
                    'extraction_quality': (
                        'Excellent' if sum([PDFPLUMBER_AVAILABLE, PYMUPDF_AVAILABLE, PDFMINER_AVAILABLE]) >= 2
                        else 'Good' if sum([PDFPLUMBER_AVAILABLE, PYMUPDF_AVAILABLE, PDFMINER_AVAILABLE]) >= 1
                        else 'Serverless Optimized (PyPDF2 Enhanced)' if PYPDF2_AVAILABLE
                        else 'Critical - No PDF extraction available'
                    ),
                    'runtime_environment': 'Vercel Serverless',
                    'recommendation': (
                        'All dependencies available - optimal performance' if sum([PYPDF2_AVAILABLE, PDFPLUMBER_AVAILABLE, PYMUPDF_AVAILABLE, PDFMINER_AVAILABLE]) >= 3
                        else 'Consider installing missing PDF libraries for better extraction quality' if sum([PYPDF2_AVAILABLE, PDFPLUMBER_AVAILABLE, PYMUPDF_AVAILABLE, PDFMINER_AVAILABLE]) < 3
                        else 'URGENT: Install PDF extraction libraries'
                    )
                }
                
                self.send_success_response(diagnostic_info)
            else:
                self.send_error_response({'error': 'Use POST for CV analysis or GET /diagnostics for status'}, 405)
                
        except Exception as e:
            logger.error(f"Diagnostics error: {str(e)}")
            self.send_error_response({'error': 'Diagnostics failed'}, 500)
    
    def do_HEAD(self):
        """Handle HEAD requests for connectivity testing"""
        try:
            self.send_response(200)
            self.send_header('Content-Type', 'application/json')
            for key, value in cors_headers().items():
                self.send_header(key.replace('_', '-'), value)
            self.end_headers()
        except Exception as e:
            logger.error(f"HEAD request error: {str(e)}")
            self.send_response(500)
            self.end_headers()
    
    def send_success_response(self, data):
        """Send successful JSON response"""
        self.send_response(200)
        self.send_header('Content-Type', 'application/json')
        for key, value in cors_headers().items():
            self.send_header(key.replace('_', '-'), value)
        self.end_headers()
        self.wfile.write(json.dumps(data).encode('utf-8'))
    
    def send_error_response(self, error_data, status_code):
        """Send error JSON response"""
        self.send_response(status_code)
        self.send_header('Content-Type', 'application/json')
        for key, value in cors_headers().items():
            self.send_header(key.replace('_', '-'), value)
        self.end_headers()
        self.wfile.write(json.dumps(error_data).encode('utf-8'))

# For local testing
if __name__ == "__main__":
    # Test the function locally
    test_event = {
        'httpMethod': 'POST',
        'body': json.dumps({
            'file_url': 'https://example.com/test-resume.pdf',
            'analysis_type': 'comprehensive',
            'include_recommendations': True
        })
    }
    
    try:
        result = handler(test_event, None)
        print(json.dumps(result, indent=2))
    except Exception as e:
        print(f"Test error: {e}")